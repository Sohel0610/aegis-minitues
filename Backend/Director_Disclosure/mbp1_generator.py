"""
MBP-1 Form Generator — Production Grade
Python 3.12+  |  python-docx 1.2+

Usage:
    python mbp1_generator.py --input director_data.json [--output MBP1_<DIN>.docx]
    python mbp1_generator.py --input director_data.json --date "1st April, 2025"

Input JSON schema (all fields optional except `din` or `name`):
    {
      "din": "...",
      "name": "...",
      "din_status": "...",
      "gender": "...",
      "indian": "Yes",
      "dir3_kyc": "...",
      "approve_date": "...",
      "updated_at": "...",
      "association": [
        {
          "cin": "...",
          "com_name": "...",
          "designation": "...",
          "appointment": "YYYY-MM-DD",
          "entity_type": "company"
        }
      ],
      // Optional MBP-1 overrides:
      "father_name": "...",
      "address": "...",
      "signature_date": "...",
      "signature_place": "Ahmedabad",
      "relatives": {
        "huf": "...",
        "wife": "...",
        "father": "...",
        "mother": "...",
        "son": "...",
        "sons_wife": "...",
        "daughter": "...",
        "daughters_husband": "...",
        "brother": "...",
        "sister": "..."
      },
      "relative_bodies_corporate": [...],
      "relative_firms": "NIL",
      "shareholding_above_2pct": "NIL"
    }
"""

from __future__ import annotations

import argparse
import json
import re
import sys
import os
import time
import psycopg2
from psycopg2.extras import RealDictCursor
from dotenv import load_dotenv
from copy import deepcopy
from dataclasses import dataclass, field
from datetime import date, datetime
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.style import WD_STYLE_TYPE


# ---------------------------------------------------------------------------
# Constants / defaults
# ---------------------------------------------------------------------------

PLACEHOLDER_PREFIX = "{{"
PLACEHOLDER_SUFFIX = "}}"

DEFAULT_PLACE = "Ahmedabad"
DEFAULT_FONT  = "Adani"
BODY_SIZE     = Pt(10)
SMALL_SIZE    = Pt(9)
TITLE_SIZE    = Pt(12)

PROGRESS_FILE = os.path.join(os.path.dirname(__file__), "sync_progress_mbp1.json")

def report_progress(current, total, status="Generating Documents..."):
    try:
        with open(PROGRESS_FILE, "w") as f:
            json.dump({"current": current, "total": total, "status": status, "timestamp": time.time()}, f)
    except:
        pass

# ---------------------------------------------------------------------------
# XML Cleanup Utility
# ---------------------------------------------------------------------------
def clean_xml_string(s: Any) -> str:
    """Removes control characters that are illegal in XML."""
    if s is None:
        return ""
    s = str(s)
    # Remove NULL bytes and other common illegal XML control chars
    # XML 1.0 allows: #x9 | #xA | #xD | [#x20-#xD7FF] | [#xE000-#xFFFD] | [#x10000-#x10FFFF]
    return "".join(c for c in s if _is_valid_xml_char(c))

def _is_valid_xml_char(c: str) -> bool:
    codepoint = ord(c)
    return (
        codepoint == 0x9 or
        codepoint == 0xA or
        codepoint == 0xD or
        (codepoint >= 0x20 and codepoint <= 0xD7FF) or
        (codepoint >= 0xE000 and codepoint <= 0xFFFD) or
        (codepoint >= 0x10000 and codepoint <= 0x10FFFF)
    )

# CIN patterns
_PUB_CIN_RE  = re.compile(r'^[UL]\d{5}[A-Z]{2}\d{4}PLC\d{6}$')
_PRIV_CIN_RE = re.compile(r'^[UL]\d{5}[A-Z]{2}\d{4}PTC\d{6}$')


# ---------------------------------------------------------------------------
# Data model
# ---------------------------------------------------------------------------

@dataclass
class DirectorData:
    din: str = "{{DIRECTOR_DIN}}"
    name: str = "{{DIRECTOR_NAME}}"
    din_status: str = "{{DIRECTOR_DIN_STATUS}}"
    gender: str = "{{DIRECTOR_GENDER}}"
    nationality: str = "{{DIRECTOR_NATIONALITY}}"
    kyc_status: str = "{{DIRECTOR_KYC_STATUS}}"
    approval_date: str = "{{DIRECTOR_APPROVAL_DATE}}"
    last_updated: str = "{{DIRECTOR_LAST_UPDATED}}"

    father_name: str = "{{FATHER_NAME}}"
    address: str = "{{DIRECTOR_ADDRESS}}"
    signature_date: str = "{{SIGNATURE_DATE}}"
    signature_place: str = DEFAULT_PLACE

    # primary designation used in signature block
    primary_designation: str = "Director"

    associations: list[dict] = field(default_factory=list)

    # relatives
    rel_huf: str = "{{RELATIVE_HUF}}"
    rel_wife: str = "{{RELATIVE_WIFE}}"
    spouse_pan: str = ""
    rel_father: str = "{{RELATIVE_FATHER}}"
    rel_mother: str = "{{RELATIVE_MOTHER}}"
    rel_son: str = "{{RELATIVE_SON}}"
    rel_sons_wife: str = "{{RELATIVE_SONS_WIFE}}"
    rel_daughter: str = "{{RELATIVE_DAUGHTER}}"
    rel_daughters_husband: str = "{{RELATIVE_DAUGHTERS_HUSBAND}}"
    rel_brother: str = "{{RELATIVE_BROTHER}}"
    rel_sister: str = "{{RELATIVE_SISTER}}"

    rel_bodies_corporate: list[dict] = field(default_factory=list)
    rel_firms: str = "NIL"
    shareholding_above_2pct: str = "NIL"

    # derived
    public_companies: list[dict] = field(default_factory=list)
    private_subsidiary: list[dict] = field(default_factory=list)
    private_non_subsidiary: list[dict] = field(default_factory=list)
    primary_company: str = "{{PRIMARY_COMPANY_NAME}}"
    target_companies: list[str] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Database Utilities
# ---------------------------------------------------------------------------

def get_db_connection():
    # Load configuration from .env
    env_path = os.path.join(os.path.dirname(__file__), '..', 'aegis_backend', '.env')
    load_dotenv(env_path)
    
    try:
        conn = psycopg2.connect(
            host=os.getenv('POSTGRES_HOST'),
            port=os.getenv('POSTGRES_PORT', '5432'),
            user=os.getenv('POSTGRES_USER'),
            password=os.getenv('POSTGRES_PASSWORD'),
            dbname=os.getenv('POSTGRES_DATABASE_DIRECTOR', 'director_disclosure_system'),
            sslmode='require'
        )
        return conn
    except Exception as e:
        print(f"[DB_ERROR] Connection failed: {e}")
        return None

def fetch_full_data_from_db(din: str, cins: list[str], sig_date: Optional[str] = None):
    conn = get_db_connection()
    if not conn: return None
    
    try:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        
        # 0. Ensure Schema Integrity (Self-healing migration)
        cur.execute("ALTER TABLE family_information.director_family ADD COLUMN IF NOT EXISTS spouse_pan TEXT")
        cur.execute("ALTER TABLE family_information.director_family ADD COLUMN IF NOT EXISTS din TEXT")
        conn.commit()

        # 1. Fetch Director Info (Master + Profile + Family)
        cur.execute("""
            SELECT 
                d.name, d.din, d.din_status, d.gender,
                p.address, p.pan, p.date_of_birth, p.qualification, p.experience,
                f.father, f.mother, f.son, f.sons_wife, f.daughter, 
                f.daughters_husband, f.brother, f.sister,
                f.section_2_77_i as rel_huf_db,
                f.section_2_77_ii as rel_wife_db,
                f.spouse_pan as spouse_pan_db
            FROM directors_master.directors d
            LEFT JOIN directors_profile.directors_profile p ON TRIM(d.din) = TRIM(p.din)
            LEFT JOIN family_information.director_family f 
                ON TRIM(d.din) = TRIM(f.din) OR TRIM(UPPER(d.name)) = TRIM(UPPER(f.director_name))
            WHERE TRIM(d.din) = TRIM(%s)
        """, (din,))
        d_row = cur.fetchone()
        if not d_row:
            print(f"[WARN] No director found for DIN {din}")
            return None
            
        # 2. Fetch Relational Family Members (Aggregate multiples with commas)
        cur.execute("""
            SELECT relationship, full_name 
            FROM family_information.director_family_members 
            WHERE TRIM(din) = TRIM(%s)
            ORDER BY relationship, pairing_group
        """, (din,))
        m_rows = cur.fetchall()
        
        relational_map = {}
        for m in m_rows:
            rel = m['relationship']
            name = m['full_name']
            if rel not in relational_map: relational_map[rel] = []
            relational_map[rel].append(name)
            
        def _get_fam(key, legacy_val):
            names = relational_map.get(key, [])
            if not names: return legacy_val or "NIL"
            return ", ".join(names)

        f_row = d_row # Combined row
            
        # 3. Fetch Target Company Names
        target_co_names = []
        if cins:
            placeholders = ', '.join(['%s'] * len(cins))
            cur.execute(f"SELECT name FROM directors_data.companies WHERE cin IN ({placeholders})", tuple(cins))
            target_co_names = [r['name'] for r in cur.fetchall()]
        
        if not target_co_names:
            target_co_names = ["{{PRIMARY_COMPANY_NAME}}"]
            
        # 4. Fetch Associations (Other Companies) - Filter for Active Only
        cur.execute("""
            SELECT cin, company_name as com_name, designation, appointment_date as appointment, status
            FROM directors_master.external_board_members 
            WHERE din = %s
            AND (status IS NULL OR status = '' OR status = 'None' OR status ILIKE 'Active%%')
            ORDER BY appointment_date DESC
        """, (din,))
        assoc_rows = cur.fetchall()
        
        # Map to Objects
        dd = DirectorData()
        dd.din = d_row['din']
        dd.name = d_row['name']
        dd.din_status = d_row['din_status'] or "Approved"
        dd.gender = d_row['gender'] or "Male"
        dd.address = d_row['address'] or "{{DIRECTOR_ADDRESS}}"
        dd.father_name = f_row.get('father') or "{{FATHER_NAME}}"
        dd.signature_date = sig_date if sig_date else f"{date.today().strftime('%d')}th {date.today().strftime('%B, %Y')}"
        dd.primary_company = target_co_names[0]
        dd.target_companies = target_co_names
        
        # Family Mapping (Prefer Relational, Fallback to Legacy Columns)
        dd.rel_huf = f_row.get('rel_huf_db') or "NIL"
        dd.rel_wife = _get_fam('Spouse', f_row.get('rel_wife_db'))
        dd.spouse_pan = f_row.get('spouse_pan_db') or ""
        dd.rel_father = _get_fam('Father', f_row.get('father'))
        dd.rel_mother = _get_fam('Mother', f_row.get('mother'))
        dd.rel_son = _get_fam('Son', f_row.get('son'))
        dd.rel_sons_wife = _get_fam("Son's Wife", f_row.get('sons_wife'))
        dd.rel_daughter = _get_fam('Daughter', f_row.get('daughter'))
        dd.rel_daughters_husband = _get_fam("Daughter's Husband", f_row.get('daughters_husband'))
        dd.rel_brother = _get_fam('Brother', f_row.get('brother'))
        dd.rel_sister = _get_fam('Sister', f_row.get('sister'))

        # Associations
        dd.associations = [dict(a) for a in assoc_rows]
        
        # Derive designation
        desig_priority = {"Company Secretary": 0, "Director": 1, "Additional Director": 2}
        if dd.associations:
            best = sorted(dd.associations, key=lambda x: desig_priority.get(x.get("designation",""), 99))
            dd.primary_designation = best[0].get("designation", "Director")
            
        # Classify
        for a in dd.associations:
            acin = a.get("cin", "")
            if _PUB_CIN_RE.match(acin):
                dd.public_companies.append(a)
            elif _PRIV_CIN_RE.match(acin):
                dd.private_non_subsidiary.append(a)
            else:
                dd.public_companies.append(a)

        return dd
        
    except Exception as e:
        print(f"[ERROR] DB Fetch failed: {e}")
        return None
    finally:
        conn.close()

def get_all_pairs_from_db():
    conn = get_db_connection()
    if not conn: return []
    try:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        # Fetch all (din, cin) pairs, but restricted to Group Companies (those in our master companies table)
        cur.execute("""
            SELECT ea.din, ea.cin 
            FROM directors_master.external_board_members ea
            JOIN directors_data.companies c ON ea.cin = c.cin
            WHERE ea.din IS NOT NULL AND ea.cin IS NOT NULL
        """)
        pairs = cur.fetchall()
        return [(p['din'], p['cin']) for p in pairs]
    finally:
        conn.close()


def _p(value: Any, placeholder: str) -> str:
    """Return value if non-empty/truthy string, else return placeholder."""
    if value is None:
        return placeholder
    s = str(value).strip()
    if s in ("", "---", "-", "None", "null"):
        return placeholder
    return s


def _fmt_date(val: Any) -> str:
    """Convert date object or ISO string → DD/MM/YYYY. Return as-is if not parseable."""
    if isinstance(val, (date, datetime)):
        return val.strftime("%d/%m/%Y")
    if not val:
        return ""
    try:
        d = date.fromisoformat(str(val))
        return d.strftime("%d/%m/%Y")
    except Exception:
        return str(val)


def parse_input(raw: dict) -> DirectorData:
    """Map raw JSON dict → DirectorData, leaving unknown fields as placeholders."""
    dd = DirectorData()

    dd.din            = _p(raw.get("din"),         "{{DIRECTOR_DIN}}")
    dd.name           = _p(raw.get("name"),        "{{DIRECTOR_NAME}}")
    dd.din_status     = _p(raw.get("din_status"),  "{{DIRECTOR_DIN_STATUS}}")
    dd.gender         = _p(raw.get("gender"),      "{{DIRECTOR_GENDER}}")
    dd.nationality    = _p("Indian" if raw.get("indian") == "Yes" else raw.get("indian"), "{{DIRECTOR_NATIONALITY}}")
    dd.kyc_status     = _p(raw.get("dir3_kyc"),    "{{DIRECTOR_KYC_STATUS}}")
    dd.approval_date  = _p(raw.get("approve_date"), "{{DIRECTOR_APPROVAL_DATE}}")
    dd.last_updated   = _p(raw.get("updated_at"),  "{{DIRECTOR_LAST_UPDATED}}")

    dd.father_name      = _p(raw.get("father_name"),    "{{FATHER_NAME}}")
    dd.address          = _p(raw.get("address"),        "{{DIRECTOR_ADDRESS}}")
    dd.signature_date   = _p(raw.get("signature_date"), "{{SIGNATURE_DATE}}")
    dd.signature_place  = _p(raw.get("signature_place"), DEFAULT_PLACE)

    assoc = raw.get("association", [])
    # Filter for active only
    dd.associations = [
        a for a in assoc 
        if not a.get("status") or 
        a.get("status", "").strip() == "" or 
        str(a.get("status", "")).lower().startswith("active")
    ]

    # derive primary company (first in list)
    if assoc:
        dd.primary_company = assoc[0]["com_name"]

    # derive primary designation
    desig_priority = {"Company Secretary": 0, "Director": 1, "Additional Director": 2}
    if assoc:
        best = sorted(assoc, key=lambda x: desig_priority.get(x.get("designation",""), 99))
        dd.primary_designation = best[0].get("designation", "Director")

    # classify associations
    for a in assoc:
        cin = a.get("cin", "")
        if _PUB_CIN_RE.match(cin):
            dd.public_companies.append(a)
        elif _PRIV_CIN_RE.match(cin):
            dd.private_non_subsidiary.append(a)
        else:
            dd.public_companies.append(a)   # default to public if unclassified

    # relatives
    rels = raw.get("relatives", {})
    dd.rel_huf                = _p(rels.get("huf"),                  "{{RELATIVE_HUF}}")
    dd.rel_wife               = _p(rels.get("wife"),                 "{{RELATIVE_WIFE}}")
    dd.rel_father             = _p(rels.get("father"),               "{{RELATIVE_FATHER}}")
    dd.rel_mother             = _p(rels.get("mother"),               "{{RELATIVE_MOTHER}}")
    dd.rel_son                = _p(rels.get("son"),                  "{{RELATIVE_SON}}")
    dd.rel_sons_wife          = _p(rels.get("sons_wife"),            "{{RELATIVE_SONS_WIFE}}")
    dd.rel_daughter           = _p(rels.get("daughter"),             "{{RELATIVE_DAUGHTER}}")
    dd.rel_daughters_husband  = _p(rels.get("daughters_husband"),    "{{RELATIVE_DAUGHTERS_HUSBAND}}")
    dd.rel_brother            = _p(rels.get("brother"),              "{{RELATIVE_BROTHER}}")
    dd.rel_sister             = _p(rels.get("sister"),               "{{RELATIVE_SISTER}}")

    dd.rel_bodies_corporate   = raw.get("relative_bodies_corporate", [])
    dd.rel_firms              = _p(raw.get("relative_firms"),        "NIL")
    dd.shareholding_above_2pct = _p(raw.get("shareholding_above_2pct"), "NIL")

    return dd


# ---------------------------------------------------------------------------
# Formatting helpers
# ---------------------------------------------------------------------------

def _set_font(run, bold: bool = False, size: Pt = BODY_SIZE,
              font: str = DEFAULT_FONT, underline: bool = False,
              color: RGBColor | None = None) -> None:
    run.font.name   = font
    run.font.size   = size
    run.font.bold   = bold
    run.font.underline = underline
    if color:
        run.font.color.rgb = color


def _para_spacing(para, before: Pt = Pt(0), after: Pt = Pt(0),
                  line: float | None = None) -> None:
    pf = para.paragraph_format
    pf.space_before = before
    pf.space_after  = after
    if line is not None:
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing      = line


def _add_para(doc: Document, text: str = "", bold: bool = False,
              size: Pt = BODY_SIZE, align=WD_ALIGN_PARAGRAPH.LEFT,
              before: Pt = Pt(0), after: Pt = Pt(4),
              font: str = DEFAULT_FONT) -> Any:
    para = doc.add_paragraph()
    para.alignment = align
    _para_spacing(para, before=before, after=after)
    if text:
        run = para.add_run(clean_xml_string(text))
        _set_font(run, bold=bold, size=size, font=font)
    return para


def _add_mixed_para(doc: Document, parts: list[tuple[str, bool]],
                    align=WD_ALIGN_PARAGRAPH.LEFT,
                    before: Pt = Pt(0), after: Pt = Pt(4),
                    size: Pt = BODY_SIZE) -> Any:
    """parts = [(text, bold), ...]"""
    para = doc.add_paragraph()
    para.alignment = align
    _para_spacing(para, before=before, after=after)
    for text, bold in parts:
        run = para.add_run(clean_xml_string(text))
        _set_font(run, bold=bold, size=size)
    return para


def _set_cell_bg(cell, hex_color: str) -> None:
    """No background color — plain white throughout to match document style."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "FFFFFF")   # always white
    tcPr.append(shd)


def _set_cell_borders(cell, color: str = "000000") -> None:
    """Plain black single border on all sides — no color override."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")   # always black
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _cell_text(cell, text: str, bold: bool = False,
               size: Pt = BODY_SIZE, align=WD_ALIGN_PARAGRAPH.LEFT,
               color: RGBColor | None = None) -> None:
    para = cell.paragraphs[0]
    para.alignment = align
    _para_spacing(para, before=Pt(2), after=Pt(2))
    run = para.add_run(clean_xml_string(text))
    _set_font(run, bold=bold, size=size, color=color)


def _table_header_row(table, headers: list[str], col_widths_cm: list[float],
                      bg: str = "FFFFFF") -> None:
    """Plain header row: white background, black bold text, black borders."""
    row = table.rows[0]
    for i, (cell, header) in enumerate(zip(row.cells, headers)):
        _set_cell_bg(cell, "FFFFFF")
        _set_cell_borders(cell, "000000")
        cell.width = Cm(col_widths_cm[i])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _cell_text(cell, header, bold=True, size=SMALL_SIZE,
                   align=WD_ALIGN_PARAGRAPH.CENTER)


def _add_table_row(table, values: list[str],
                   col_widths_cm: list[float],
                   bg: str = "FFFFFF", row_idx: int = 0) -> None:
    """Plain data row: white background, black borders, normal text."""
    row = table.add_row()
    for i, (cell, val) in enumerate(zip(row.cells, values)):
        _set_cell_bg(cell, "FFFFFF")
        _set_cell_borders(cell, "000000")
        cell.width = Cm(col_widths_cm[i])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _cell_text(cell, val, size=SMALL_SIZE)


def _hr(doc: Document, color: str = "000000") -> None:
    """Horizontal rule via paragraph bottom border."""
    para = doc.add_paragraph()
    _para_spacing(para, before=Pt(2), after=Pt(2))
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _sig_table(doc: Document, dd: DirectorData) -> None:
    """Two-column signature block: Place/Date on left, Signature details on right."""
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"

    data = [
        ("", f"Signature   :  _______________"),
        (f"Place:  {dd.signature_place}",  f"Name         :  {dd.name}"),
        (f"Date:   {dd.signature_date}",   f"Designation :  {dd.primary_designation}"),
        ("",                               f"DIN             :  {dd.din}"),
    ]

    for r_idx, (left, right) in enumerate(data):
        row = table.rows[r_idx]
        for cell in row.cells:
            for border_side in ["top", "bottom", "left", "right"]:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement("w:tcBorders")
                for side in ["top", "bottom", "left", "right"]:
                    el = OxmlElement(f"w:{side}")
                    el.set(qn("w:val"), "none")
                    el.set(qn("w:sz"), "0")
                    el.set(qn("w:space"), "0")
                    el.set(qn("w:color"), "FFFFFF")
                    tcBorders.append(el)
                tcPr.append(tcBorders)

        _cell_text(row.cells[0], left,  size=SMALL_SIZE)
        _cell_text(row.cells[1], right, size=SMALL_SIZE)

    _para_spacing(doc.add_paragraph(), before=Pt(6), after=Pt(6))


# ---------------------------------------------------------------------------
# Document sections
# ---------------------------------------------------------------------------

def _section_title(doc: Document) -> None:
    """FORM MBP-1 title block."""
    p1 = _add_para(doc, "FORM MBP – 1", bold=True, size=Pt(14),
                   align=WD_ALIGN_PARAGRAPH.CENTER, after=Pt(2))
    p2 = _add_para(doc, "Notice of Interest by Director", bold=True,
                   size=Pt(12), align=WD_ALIGN_PARAGRAPH.CENTER, after=Pt(2))
    p3 = _add_para(
        doc,
        "[Pursuant to Section 184(1) of the Companies Act, 2013 and "
        "Rule 9(1) of the Companies (Meetings of Board and its Powers) Rules, 2014]",
        bold=False, size=SMALL_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER, after=Pt(8)
    )
    _hr(doc)


def _section_addressee(doc: Document, dd: DirectorData) -> None:
    _add_para(doc, "", after=Pt(4))
    _add_para(doc, "To,", size=BODY_SIZE, after=Pt(0))
    _add_para(doc, "The Board of Directors of", size=BODY_SIZE, after=Pt(0))

    if dd.target_companies:
        for i, company in enumerate(dd.target_companies):
            after = Pt(8) if i == len(dd.target_companies) - 1 else Pt(0)
            _add_para(doc, company.upper(), bold=True, size=BODY_SIZE, after=after)
    else:
        _add_para(doc, "{{PRIMARY_COMPANY_NAME}}", bold=True, size=BODY_SIZE, after=Pt(8))


def _section_opening(doc: Document, dd: DirectorData) -> None:
    _add_para(doc, "Dear Sir(s)", bold=False, size=BODY_SIZE, after=Pt(6))

    para = doc.add_paragraph()
    _para_spacing(para, after=Pt(8))
    parts = [
        (f"I, ", False),
        (dd.name, True),
        (f", son/daughter of ", False),
        (dd.father_name, True),
        (f", resident of {dd.address}, being a Director in the Company hereby "
         "give notice of my interest or concern in the following company or companies, "
         "bodies corporate, firms or other association of individuals:—", False),
    ]
    for text, bold in parts:
        run = para.add_run(text)
        _set_font(run, bold=bold, size=BODY_SIZE)


def _section_interest_table(doc: Document, dd: DirectorData) -> None:
    """Sr. No / Company / Designation / Shareholding / Date table."""
    headers   = ["Sr.\nNo.", "Names of Companies / Bodies Corporate /\nFirms / Association of Individuals",
                 "Nature of Interest\nor Concern", "Share-\nholding", "Date of\nInterest"]
    col_w_cm  = [1.2, 7.0, 3.2, 1.8, 2.2]

    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _table_header_row(table, headers, col_w_cm)

    for idx, assoc in enumerate(dd.associations):
        appt = assoc.get("appointment", "")
        date_str = _fmt_date(appt) if appt else "—"
        _add_table_row(
            table,
            [str(idx + 1), assoc.get("com_name", ""), assoc.get("designation", ""),
             "NIL", date_str],
            col_w_cm, row_idx=idx
        )

    if not dd.associations:
        _add_table_row(table, ["—", "{{ROW_COMPANY_NAME}}", "{{ROW_DESIGNATION}}", "NIL", "{{ROW_DATE}}"],
                       col_w_cm, row_idx=0)

    doc.add_paragraph()


def _section_signature(doc: Document, dd: DirectorData) -> None:
    _sig_table(doc, dd)


def _section_declaration_header(doc: Document, dd: DirectorData) -> None:
    _add_para(doc, "", after=Pt(6))
    _add_para(
        doc,
        f"Declaration by {dd.name} giving various details including the list of Relatives "
        "and their interest or concern in the companies, bodies corporate, firms or other "
        "association of individuals:",
        size=BODY_SIZE, after=Pt(8)
    )
    _add_para(doc, "I.\tList of Relatives under Section 2(77) of the Companies Act, 2013",
              bold=True, size=BODY_SIZE, after=Pt(6))


def _section_relatives(doc: Document, dd: DirectorData) -> None:
    relatives = [
        ("I) Section 2(77)(i)",    "H.U.F in which I am a Member:", dd.rel_huf),
        ("II) Section 2(77)(ii)",  "Wife:",                         dd.rel_wife),
        ("III) Section 2(77)(iii)", None, None),  # numbered sub-items
    ]

    # HUF
    _add_mixed_para(doc, [("I)\tSection 2(77)(i)\t\t:  H.U.F in which I am a Member:  ", False),
                          (dd.rel_huf, False)], size=BODY_SIZE, after=Pt(3))

    # Wife
    _add_mixed_para(doc, [("II)\tSection 2(77)(ii)\t:  Wife:  ", False),
                          (dd.rel_wife, False)], size=BODY_SIZE, after=Pt(3))

    # Section (iii) - numbered family
    _add_para(doc, "III)\tSection 2(77)(iii)\t:", bold=False, size=BODY_SIZE, after=Pt(3))

    fam = [
        ("Father",               dd.rel_father),
        ("Mother",               dd.rel_mother),
        ("Son",                  dd.rel_son),
        ("Son's Wife",           dd.rel_sons_wife),
        ("Daughter",             dd.rel_daughter),
        ("Daughter's Husband",   dd.rel_daughters_husband),
        ("Brother",              dd.rel_brother),
        ("Sister",               dd.rel_sister),
    ]

    # Create a borderless table for alignment
    table = doc.add_table(rows=len(fam), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    for i, (label, value) in enumerate(fam):
        row = table.rows[i]
        
        # Remove borders for each cell
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top', 'bottom', 'left', 'right']:
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'none')
                tcBorders.append(border)
            tcPr.append(tcBorders)

        # Set content
        cell_label = row.cells[0]
        cell_value = row.cells[1]
        
        # Relationship column
        p_label = cell_label.paragraphs[0]
        p_label.paragraph_format.left_indent = Inches(0.5)
        run_l = p_label.add_run(f"{i+1}.\t{label}\t:")
        _set_font(run_l, size=BODY_SIZE)
        
        # Name column
        p_val = cell_value.paragraphs[0]
        run_v = p_val.add_run(value)
        _set_font(run_v, size=BODY_SIZE)

    doc.add_paragraph()


def _section_relative_declarations(doc: Document, dd: DirectorData) -> None:
    _add_para(doc, "II.\tBodies Corporate in which my relatives are Directors:", bold=True, size=BODY_SIZE, after=Pt(4))

    if dd.rel_bodies_corporate:
        headers  = ["Name of the Relative(s)", "Name of the Company"]
        col_w_cm = [7.0, 8.6]
        table = doc.add_table(rows=1, cols=2)
        _table_header_row(table, headers, col_w_cm)
        for r_idx, item in enumerate(dd.rel_bodies_corporate):
            _add_table_row(table, [item.get("relative", ""), item.get("company", "")],
                           col_w_cm, row_idx=r_idx)
    else:
        _add_para(doc, "\t—  NIL  —", size=BODY_SIZE, after=Pt(4),
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()
    _add_para(doc, "III.\tFirms in which my relatives are partner [Section 2(76)(iii)]:",
              bold=True, size=BODY_SIZE, after=Pt(2))
    _add_para(doc, f"\t\t\t—————  {dd.rel_firms}  —————",
              align=WD_ALIGN_PARAGRAPH.CENTER, size=BODY_SIZE, after=Pt(6))

    _add_para(doc,
              "IV.\tPublic Companies in which I am a Director and hold together with my relatives "
              "more than 2% of the paid up share capital [Section 2(76)(v)]:",
              bold=True, size=BODY_SIZE, after=Pt(2))
    _add_para(doc, f"\t\t\t—————  {dd.shareholding_above_2pct}  —————",
              align=WD_ALIGN_PARAGRAPH.CENTER, size=BODY_SIZE, after=Pt(6))

    doc.add_paragraph()


def _section_compliance_letter(doc: Document, dd: DirectorData) -> None:
    """Section 164(2) compliance letter — second major section."""
    doc.add_page_break()

    _add_para(doc, dd.signature_date, size=BODY_SIZE, after=Pt(8))
    _add_para(doc, "To", size=BODY_SIZE)
    _add_para(doc, "The Board of Directors of", size=BODY_SIZE)

    all_companies = [a["com_name"] for a in dd.associations]
    # deduplicate while preserving order
    seen = set()
    unique_companies = []
    for c in all_companies:
        if c not in seen:
            seen.add(c)
            unique_companies.append(c)

    for company in unique_companies:
        _add_para(doc, company, bold=True, size=BODY_SIZE, after=Pt(1))

    doc.add_paragraph()
    _add_para(doc, "Dear Sir,", size=BODY_SIZE, after=Pt(6))
    _add_para(doc, "Sub: Compliance under Section 164(2) of the Companies Act, 2013",
              bold=True, size=BODY_SIZE, after=Pt(8))

    # (A) Public companies
    _add_para(doc, "(A) Public Limited Companies:", bold=True, size=BODY_SIZE, after=Pt(4))
    pub_companies = [a["com_name"] for a in dd.public_companies]
    seen_pub = set()
    for c in pub_companies:
        if c not in seen_pub:
            seen_pub.add(c)
            _add_para(doc, c, size=BODY_SIZE, after=Pt(1))

    doc.add_paragraph()

    # (B) Private subsidiary
    _add_para(doc, "(B) Private Limited Companies which are subsidiary(ies) of Public Companies:",
              bold=True, size=BODY_SIZE, after=Pt(4))
    if dd.private_subsidiary:
        for a in dd.private_subsidiary:
            _add_para(doc, a["com_name"], size=BODY_SIZE, after=Pt(1))
    else:
        _add_para(doc, "\t\t\tNIL", size=BODY_SIZE,
                  align=WD_ALIGN_PARAGRAPH.CENTER, after=Pt(4))

    doc.add_paragraph()

    # (C) Private non-subsidiary
    _add_para(doc, "(C) Private Limited Companies which are not subsidiary(ies) of Public Companies:",
              bold=True, size=BODY_SIZE, after=Pt(4))
    seen_priv = set()
    for a in dd.private_non_subsidiary:
        c = a["com_name"]
        if c not in seen_priv:
            seen_priv.add(c)
            _add_para(doc, c, size=BODY_SIZE, after=Pt(1))
    if not dd.private_non_subsidiary:
        _add_para(doc, "\t\t\tNIL", size=BODY_SIZE,
                  align=WD_ALIGN_PARAGRAPH.CENTER, after=Pt(4))

    doc.add_paragraph()

    # Certification text
    cert_text = (
        "I hereby certify that the aforesaid Companies have not:\n\n"
        "1.\tfailed to file financial statements or annual returns for any continuous period "
        "of three financial years; or\n\n"
        "2.\tfailed to repay the deposits accepted by it or pay interest thereon or redeem any "
        "debentures on due date or pay interest due thereon or pay any dividend declared and "
        "such failure to pay or redeem continues for one year or more."
    )

    _add_para(
        doc,
        "I hereby certify that the aforesaid Companies have not:",
        size=BODY_SIZE, after=Pt(4)
    )
    _add_para(
        doc,
        "1.\tfailed to file financial statements or annual returns for any continuous period "
        "of three financial years; or",
        size=BODY_SIZE, after=Pt(4)
    )
    _add_para(
        doc,
        "2.\tfailed to repay the deposits accepted by it or pay interest thereon or redeem any "
        "debentures on due date or pay interest due thereon or pay any dividend declared and "
        "such failure to pay or redeem continues for one year or more.",
        size=BODY_SIZE, after=Pt(8)
    )

    _add_para(
        doc,
        f"Accordingly, I am not disqualified as a Director under the provisions of Section 164(2) "
        f"of the Companies Act, 2013 as on 31st March, {_current_year(dd.signature_date)}.",
        size=BODY_SIZE, after=Pt(4)
    )
    _add_para(
        doc,
        f"Further, I also confirm that I have not incurred any other disqualification in terms of "
        f"the provisions of Section 164(1) of the Companies Act, 2013 as on 31st March, "
        f"{_current_year(dd.signature_date)}.",
        size=BODY_SIZE, after=Pt(8)
    )

    _add_para(doc, "Thanking You,", size=BODY_SIZE, after=Pt(4))
    doc.add_paragraph()
    _add_para(doc, "…………………………", size=BODY_SIZE, after=Pt(2))
    _add_para(doc, dd.name, bold=True, size=BODY_SIZE, after=Pt(2))
    _add_para(doc, f"DIN: {dd.din}", size=BODY_SIZE, after=Pt(2))


def _current_year(sig_date: str) -> str:
    """Extract year from signature date string, fallback to current year."""
    import re
    m = re.search(r"\b(20\d{2})\b", sig_date)
    if m:
        return m.group(1)
    return str(date.today().year)


def _section_profile_appendix(doc: Document, dd: DirectorData) -> None:
    """Optional appendix: Director profile summary from API."""
    doc.add_page_break()
    _add_para(doc, "ANNEXURE — Director Profile Summary", bold=True,
              size=Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER, after=Pt(6))
    _hr(doc)
    doc.add_paragraph()

    headers  = ["Field", "Details"]
    col_w_cm = [5.5, 10.1]
    rows = [
        ("DIN",             dd.din),
        ("Name",            dd.name),
        ("Status",          dd.din_status),
        ("Gender",          dd.gender),
        ("Nationality",     dd.nationality),
        ("DIR3 KYC",        dd.kyc_status),
        ("Approval Date",   dd.approval_date),
        ("Last Updated",    dd.last_updated),
        ("Total Companies", str(len(dd.associations))),
    ]

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _table_header_row(table, headers, col_w_cm)
    for i, (label, value) in enumerate(rows):
        _add_table_row(table, [label, value], col_w_cm, row_idx=i)

    doc.add_paragraph()

    # Company associations table
    _add_para(doc, "Company Associations", bold=True, size=Pt(10), after=Pt(4))
    headers2  = ["CIN", "Company Name", "Designation", "Appointment Date"]
    col_w2_cm = [4.0, 7.5, 2.8, 2.3]
    table2 = doc.add_table(rows=1, cols=4)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    _table_header_row(table2, headers2, col_w2_cm)
    for i, assoc in enumerate(dd.associations):
        _add_table_row(
            table2,
            [assoc.get("cin", ""), assoc.get("com_name", ""),
             assoc.get("designation", ""), _fmt_date(assoc.get("appointment", ""))],
            col_w2_cm, row_idx=i
        )


# ---------------------------------------------------------------------------
# Page setup
# ---------------------------------------------------------------------------

def _setup_page(doc: Document) -> None:
    """A4 page, 2.5 cm margins (standard for Indian corporate forms)."""
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    margin = Cm(2.5)
    section.top_margin    = margin
    section.bottom_margin = margin
    section.left_margin   = margin
    section.right_margin  = margin


def _setup_styles(doc: Document) -> None:
    """Override default Normal style font."""
    style = doc.styles["Normal"]
    font  = style.font
    font.name = DEFAULT_FONT
    font.size = BODY_SIZE


# ---------------------------------------------------------------------------
# Main builder
# ---------------------------------------------------------------------------

def build_document(dd: DirectorData, include_annexure: bool = True) -> Document:
    doc = Document()
    _setup_page(doc)
    _setup_styles(doc)

    # -- Page 1: MBP-1 Form --
    _section_title(doc)
    _section_addressee(doc, dd)
    _section_opening(doc, dd)
    _section_interest_table(doc, dd)
    _section_signature(doc, dd)
    _hr(doc)
    _section_declaration_header(doc, dd)
    _section_relatives(doc, dd)
    _section_relative_declarations(doc, dd)
    _section_signature(doc, dd)   # second signature block

    # -- Page 2: Section 164(2) compliance letter --
    _section_compliance_letter(doc, dd)

    # -- Page 3: Annexure (optional) --
    if include_annexure:
        _section_profile_appendix(doc, dd)

    return doc


# ---------------------------------------------------------------------------
# CLI entrypoint
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate Form MBP-1 Word documents from Azure PostgreSQL Database."
    )
    # DB Mode
    parser.add_argument("--din",    help="Director DIN")
    parser.add_argument("--cin",    help="Target Company CIN (comma separated for multiple)")
    parser.add_argument("--all",    action="store_true", help="Generate for all directors across all companies")
    parser.add_argument("--year",   default="2024-25",   help="Fiscal year for output folder")
    
    # Legacy / Overlay
    parser.add_argument("--input",  help="Path to director JSON file (legacy mode)")
    parser.add_argument("--date",   help="Signature date e.g. '1st April, 2025'")
    parser.add_argument("--no-annexure", action="store_true", help="Skip the profile summary annexure")

    args = parser.parse_args()

    pairs = []
    if args.all:
        print("[BATCH] Fetching all Director-Company pairs from DB...")
        pairs = get_all_pairs_from_db()
    elif args.din and args.cin:
        # Split CINs if multiple provided
        cins = [c.strip() for c in args.cin.split(",")]
        # For multi-company mode, we treat them as one set
        pairs = [(args.din, cins)]
    elif args.input:
        # Legacy mode
        input_path = Path(args.input)
        if not input_path.exists():
            print(f"[ERROR] File not found: {input_path}")
            return
        raw = json.loads(input_path.read_text(encoding="utf-8"))
        if args.date: raw["signature_date"] = args.date
        dd = parse_input(raw)
        _process_single(dd, args.year, not args.no_annexure)
        return
    else:
        parser.print_help()
        return

    print(f"[INFO] Found {len(pairs)} pairs to process.")
    
    total = len(pairs)
    total_generated = 0
    for i, (din, cin_data) in enumerate(pairs, 1):
        # cin_data can be a single string or a list
        cins = cin_data if isinstance(cin_data, list) else [cin_data]
        dd = fetch_full_data_from_db(din, cins, args.date)
        if dd:
            report_progress(i, total, f"Generating for {dd.name}")
            _process_single(dd, args.year, not args.no_annexure)
            total_generated += 1
        else:
            report_progress(i, total, f"Skipping DIN {din} (Not Found)")

    report_progress(total, total, "Complete")
    print(f"\n[FINISH] Total {total_generated} document(s) generated.")

def register_document_in_db(dd: DirectorData, file_path: str):
    """Registers the generated file in the document_summaries table for the repository."""
    conn = get_db_connection()
    if not conn: return
    try:
        cur = conn.cursor()
        # Relative path from Output_Disclosures
        rel_path = str(Path(file_path).relative_to(Path("Output_Disclosures")))
        
        cur.execute("""
            INSERT INTO directors_data.document_summaries 
            (director_name, din, file_path, created_at)
            VALUES (%s, %s, %s, CURRENT_TIMESTAMP)
            ON CONFLICT (file_path) DO UPDATE 
            SET created_at = EXCLUDED.created_at,
                director_name = EXCLUDED.director_name,
                din = EXCLUDED.din
        """, (dd.name, dd.din, rel_path))
        conn.commit()
        print(f"  [DB] Registered document for {dd.name}")
    except Exception as e:
        print(f"  [DB_ERROR] Failed to register document: {e}")
    finally:
        conn.close()

def _process_single(dd: DirectorData, year: str, include_annexure: bool):
    base_output = Path("Output_Disclosures") / year
    
    doc = build_document(dd, include_annexure)
    
    # Clean names for folder/file
    clean_company = "".join(c if c.isalnum() else "_" for c in dd.primary_company).strip("_")
    clean_director = "".join(c if c.isalnum() else "_" for c in dd.name).strip("_")
    
    # Directory structure: Output/Year/Company/MBP-1/
    target_dir = base_output / clean_company / "MBP-1"
    target_dir.mkdir(parents=True, exist_ok=True)
    
    file_name = f"MBP1_{clean_director}_{dd.din}.docx"
    out_path = target_dir / file_name
    
    doc.save(str(out_path))
    register_document_in_db(dd, str(out_path))
    print(f"  [OK] {clean_company} -> {clean_director}")

if __name__ == "__main__":
    main()


if __name__ == "__main__":
    main()
