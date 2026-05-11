"""
DIR-8 Form Generator — Production Grade
Python 3.12+  |  python-docx 1.2+

Generates Form DIR-8 (Intimation by Director) matching the exact formatting
of the official template: Times New Roman 11.5pt, US Letter, 1-inch margins,
tab-aligned header fields, plain 4-column table, right-indented signature block.

Usage:
    python dir8_generator.py --input company_data.json --director-din 08858955
    python dir8_generator.py --input company_data.json --director-din 08858955 --date "1st April, 2025"

The script generates one DIR-8 document per director (or for a specific director
if --director-din is supplied).

Input JSON  (same company registry schema used for DIR-12):
{
  "company_details": { "cin": "...", "company_name": "...", "auth_capital": "...",
                       "paid_capital": "...", ... },
  "contact_details": { "address": "...", "email": "..." },
  "directors": [
    { "din": "...", "director_name": "...", "designation": "...",
      "appointment_date": "YYYY-MM-DD",
      // optional:
      "father_name": "...",
      "address": "...",
      "cessation_date": "YYYY-MM-DD",
      "other_companies": [
        { "com_name": "...", "appointment_date": "YYYY-MM-DD", "cessation_date": "YYYY-MM-DD" }
      ]
    }
  ],
  "signature_date": "...",        // optional, default placeholder
  "signature_place": "Ahmedabad"  // optional, default Ahmedabad
}
"""

from __future__ import annotations

import argparse
import json
import sys
import os
import time
import psycopg2
from psycopg2.extras import RealDictCursor
from dotenv import load_dotenv
from dataclasses import dataclass, field
from datetime import date, datetime
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Inches, Twips


# ---------------------------------------------------------------------------
# XML Cleanup Utility
# ---------------------------------------------------------------------------
def clean_xml_string(s: Any) -> str:
    """Removes control characters that are illegal in XML."""
    if s is None:
        return ""
    s = str(s)
    # Remove NULL bytes and other common illegal XML control chars
    # XML 1.0 allows: #x9 | #xA | #xD | [#x20-#xD7FF] | [#xE000-#xFFFD] | [#x10000-#x10FFFF]
    return "".join(c for c in s if _is_valid_xml_char(c))

def _is_valid_xml_char(c: str) -> bool:
    codepoint = ord(c)
    return (
        codepoint == 0x9 or
        codepoint == 0xA or
        codepoint == 0xD or
        (codepoint >= 0x20 and codepoint <= 0xD7FF) or
        (codepoint >= 0xE000 and codepoint <= 0xFFFD) or
        (codepoint >= 0x10000 and codepoint <= 0x10FFFF)
    )

# ---------------------------------------------------------------------------
# Exact measurements extracted from the template XML
# ---------------------------------------------------------------------------

FONT_NAME   = "Adani"
FONT_SZ_VAL = "23"          # w:sz value = half-points → 11.5 pt
FONT_PT     = Pt(11.5)
LINE_SPACE  = 276            # w:line auto rule
PAGE_W      = Twips(12240)   # US Letter width
PAGE_H      = Twips(15840)   # US Letter height
MARGIN      = Twips(1440)    # 1 inch

# Table column widths in DXA (must sum to 9360)
COL_SR      = Twips(636)
COL_NAME    = Twips(3864)
COL_APPT    = Twips(2610)
COL_CESS    = Twips(2250)
TBL_TOTAL   = Twips(9360)

DEFAULT_PLACE = "Ahmedabad"
PROGRESS_FILE = os.path.join(os.path.dirname(__file__), "sync_progress_dir8.json")

def report_progress(current, total, status="Generating DIR-8..."):
    try:
        with open(PROGRESS_FILE, "w") as f:
            json.dump({"current": current, "total": total, "status": status, "timestamp": time.time()}, f)
    except:
        pass


# ---------------------------------------------------------------------------
# Data model
# ---------------------------------------------------------------------------

@dataclass
class OtherCompany:
    com_name:         str = "{{ROW_COMPANY_NAME}}"
    appointment_date: str = "{{ROW_APPOINTMENT_DATE}}"
    cessation_date:   str = "{{ROW_CESSATION_DATE}}"


@dataclass
class DirectorInfo:
    din:            str = "{{DIRECTOR_DIN}}"
    name:           str = "{{DIRECTOR_NAME}}"
    father_name:    str = "{{FATHER_NAME}}"
    address:        str = "{{DIRECTOR_ADDRESS}}"
    designation:    str = "{{DIRECTOR_DESIGNATION}}"
    other_companies: list[OtherCompany] = field(default_factory=list)


@dataclass
class CompanyInfo:
    cin:           str = "{{COMPANY_CIN}}"
    company_name:  str = "{{COMPANY_NAME}}"
    address:       str = "{{COMPANY_ADDRESS}}"
    auth_capital:  str = "{{NOMINAL_CAPITAL}}"
    paid_capital:  str = "{{PAID_UP_CAPITAL}}"
    signature_date:  str = "{{SIGNATURE_DATE}}"
    signature_place: str = DEFAULT_PLACE
    target_company_names: list[str] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Database Utilities
# ---------------------------------------------------------------------------

def get_db_connection():
    # Load configuration from .env
    env_path = os.path.join(os.path.dirname(__file__), '..', 'aegis_backend', '.env')
    load_dotenv(env_path)
    
    try:
        conn = psycopg2.connect(
            host=os.getenv('POSTGRES_HOST'),
            port=os.getenv('POSTGRES_PORT', '5432'),
            user=os.getenv('POSTGRES_USER'),
            password=os.getenv('POSTGRES_PASSWORD'),
            dbname=os.getenv('POSTGRES_DATABASE_DIRECTOR', 'director_disclosure_system'),
            sslmode='require'
        )
        return conn
    except Exception as e:
        print(f"[DB_ERROR] Connection failed: {e}")
        return None

def fetch_full_data_from_db(din: str, cins: list[str], sig_date: Optional[str] = None):
    conn = get_db_connection()
    if not conn: return None, None
    
    try:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        
        # 1. Fetch Director Info (Master + Profile + Family)
        cur.execute("""
            SELECT 
                d.name, d.din, 
                p.address, p.pan,
                f.father
            FROM directors_master.directors d
            LEFT JOIN directors_profile.directors_profile p ON d.din = p.din
            LEFT JOIN family_information.director_family f ON d.name = f.director_name
            WHERE d.din = %s
        """, (din,))
        d_row = cur.fetchone()
        if not d_row:
            print(f"[WARN] No director found for DIN {din}")
            return None, None
            
        # 2. Fetch Target Company Info
        target_cos = []
        if cins:
            placeholders = ', '.join(['%s'] * len(cins))
            cur.execute(f"SELECT cin, name, address, auth_capital, paid_capital FROM directors_data.companies WHERE cin IN ({placeholders})", tuple(cins))
            target_cos = cur.fetchall()
            
        if not target_cos:
            print(f"[WARN] No companies found for CINs {cins}")
            return None, None
        
        # Use first company for primary header details
        primary_co = target_cos[0]
            
        # 3. Fetch Associations (Other Companies) - Filter for Active Only
        cur.execute("""
            SELECT company_name as com_name, appointment_date, status
            FROM directors_master.external_board_members 
            WHERE din = %s
            AND (status IS NULL OR status = '' OR status = 'None' OR status ILIKE 'Active%%')
            ORDER BY appointment_date DESC
        """, (din,))
        assoc_rows = cur.fetchall()
        
        # Map to Objects
        co = CompanyInfo()
        co.cin = primary_co['cin']
        co.company_name = primary_co['name']
        co.address = primary_co['address'] or "{{COMPANY_ADDRESS}}"
        co.auth_capital = _fmt_capital(str(primary_co['auth_capital'])) if primary_co['auth_capital'] else "{{NOMINAL_CAPITAL}}"
        co.paid_capital = _fmt_capital(str(primary_co['paid_capital'])) if primary_co['paid_capital'] else "{{PAID_UP_CAPITAL}}"
        co.signature_date = sig_date if sig_date else f"{date.today().strftime('%d')}th {date.today().strftime('%B, %Y')}"
        co.target_company_names = [r['name'] for r in target_cos]
        
        di = DirectorInfo()
        di.name = d_row['name']
        di.din = d_row['din']
        di.address = d_row['address'] or "{{DIRECTOR_ADDRESS}}"
        di.father_name = d_row['father'] or "{{FATHER_NAME}}"
        
        for a in assoc_rows:
            oc = OtherCompany()
            oc.com_name = a['com_name']
            oc.appointment_date = _fmt_date(str(a['appointment_date'])) if a['appointment_date'] else ""
            oc.cessation_date = ""
            di.other_companies.append(oc)
            
        return co, [di]
        
    except Exception as e:
        print(f"[ERROR] DB Fetch failed: {e}")
        return None, None
    finally:
        conn.close()

def get_all_pairs_from_db():
    conn = get_db_connection()
    if not conn: return []
    try:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        # Fetch all (din, cin) pairs, but restricted to Group Companies (those in our master companies table)
        cur.execute("""
            SELECT ea.din, ea.cin 
            FROM directors_master.external_board_members ea
            JOIN directors_data.companies c ON ea.cin = c.cin
            WHERE ea.din IS NOT NULL AND ea.cin IS NOT NULL
        """)
        pairs = cur.fetchall()
        return [(p['din'], p['cin']) for p in pairs]
    finally:
        conn.close()


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _p(value: Any, placeholder: str) -> str:
    if value is None:
        return placeholder
    s = str(value).strip()
    return placeholder if s in ("", "---", "-", "None", "null") else s


def _fmt_date(val: Any) -> str:
    """Convert date object or ISO string → DD/MM/YYYY."""
    if isinstance(val, (date, datetime)):
        return val.strftime("%d/%m/%Y")
    if not val:
        return ""
    try:
        return date.fromisoformat(str(val)).strftime("%d/%m/%Y")
    except Exception:
        return str(val)


def _fmt_capital(raw: str) -> str:
    """
    Convert raw numeric string → Indian format with Rs. prefix and /- suffix.
    e.g. "3700000000" → "Rs. 3,70,00,00,000/-"
    """
    try:
        n = int(raw)
        s = str(n)
        if len(s) <= 3:
            return f"Rs. {s}/-"
        result = s[-3:]
        s = s[:-3]
        while len(s) > 2:
            result = s[-2:] + "," + result
            s = s[:-2]
        if s:
            result = s + "," + result
        return f"Rs. {result}/-"
    except Exception:
        return raw


def parse_input(raw: dict, director_din: str | None) -> tuple[CompanyInfo, list[DirectorInfo]]:
    cd  = raw.get("company_details", {})
    ct  = raw.get("contact_details", {})
    dirs_raw = raw.get("directors", [])

    # Company
    co = CompanyInfo()
    co.cin          = _p(cd.get("cin"),          "{{COMPANY_CIN}}")
    co.company_name = _p(cd.get("company_name"), "{{COMPANY_NAME}}")
    co.address      = _p(ct.get("address"),      "{{COMPANY_ADDRESS}}")

    raw_auth = cd.get("auth_capital", "")
    co.auth_capital = _fmt_capital(raw_auth) if raw_auth else "{{NOMINAL_CAPITAL}}"

    raw_paid = cd.get("paid_capital", "")
    co.paid_capital = _fmt_capital(raw_paid) if raw_paid else "{{PAID_UP_CAPITAL}}"

    co.signature_date  = _p(raw.get("signature_date"),  "{{SIGNATURE_DATE}}")
    co.signature_place = _p(raw.get("signature_place"), DEFAULT_PLACE)

    # Directors
    directors: list[DirectorInfo] = []
    for d in dirs_raw:
        din = _p(d.get("din"), "")
        # filter by DIN if requested
        if director_din and din and din != director_din:
            continue

        di = DirectorInfo()
        di.din         = _p(d.get("din"),           "{{DIRECTOR_DIN}}")
        di.name        = _p(d.get("director_name"), "{{DIRECTOR_NAME}}")
        di.father_name = _p(d.get("father_name"),   "{{FATHER_NAME}}")
        di.address     = _p(d.get("address"),       "{{DIRECTOR_ADDRESS}}")
        di.designation = _p(d.get("designation"),   "{{DIRECTOR_DESIGNATION}}")

        # other_companies: if not in JSON, build a single self-row from appointment
        raw_others = d.get("other_companies", [])
        if raw_others:
            for oc in raw_others:
                # Filter for active only
                status = oc.get("status", "")
                if status and not str(status).lower().startswith("active") and status not in ["", "None", "null"]:
                    continue
                    
                entry = OtherCompany()
                entry.com_name         = _p(oc.get("com_name"),         "{{ROW_COMPANY_NAME}}")
                raw_a = oc.get("appointment_date", "")
                entry.appointment_date = _fmt_date(raw_a) if raw_a else "{{ROW_APPOINTMENT_DATE}}"
                raw_c = oc.get("cessation_date", "")
                entry.cessation_date   = _fmt_date(raw_c) if raw_c else ""
                di.other_companies.append(entry)
        else:
            # Build placeholder rows so the table is never empty
            entry = OtherCompany()
            entry.com_name         = co.company_name
            raw_a = d.get("appointment_date", "")
            entry.appointment_date = _fmt_date(raw_a) if raw_a else "{{ROW_APPOINTMENT_DATE}}"
            raw_c = d.get("cessation_date", "")
            entry.cessation_date   = _fmt_date(raw_c) if raw_c else ""
            di.other_companies.append(entry)

        directors.append(di)

    return co, directors


# ---------------------------------------------------------------------------
# Low-level XML / formatting helpers
# (mirror the exact XML patterns in the template)
# ---------------------------------------------------------------------------

def _rpr_xml(bold: bool = False) -> str:
    """Return <w:rPr> XML fragment matching template runs."""
    b = "<w:b/><w:bCs/>" if bold else ""
    return (
        f'<w:rPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/>'
        f'{b}'
        f'<w:sz w:val="{FONT_SZ_VAL}"/>'
        f'<w:szCs w:val="{FONT_SZ_VAL}"/>'
        f'</w:rPr>'
    )


def _apply_rpr(run, bold: bool = False) -> None:
    """Set run font/size to match template exactly."""
    run.font.name = FONT_NAME
    run.font.size = FONT_PT
    run.font.bold = bold
    # also set cs font via XML
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"),  FONT_NAME)
    rFonts.set(qn("w:hAnsi"),  FONT_NAME)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)
    rFonts.set(qn("w:cs"),     FONT_NAME)


def _set_line_spacing(para, line: int = LINE_SPACE) -> None:
    pPr = para._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:line"),     str(line))
    spacing.set(qn("w:lineRule"), "auto")


def _set_spacing_after(para, after: int = 0) -> None:
    pPr = para._p.get_or_add_pPr()
    sp = pPr.find(qn("w:spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing")
        pPr.append(sp)
    sp.set(qn("w:after"), str(after))


def _set_jc(para, val: str = "both") -> None:
    pPr = para._p.get_or_add_pPr()
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        pPr.append(jc)
    jc.set(qn("w:val"), val)


def _set_ind(para, left: int = 0, hanging: int = 0, first_line: int = 0) -> None:
    pPr = para._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    if left:
        ind.set(qn("w:left"), str(left))
    if hanging:
        ind.set(qn("w:hanging"), str(hanging))
    if first_line:
        ind.set(qn("w:firstLine"), str(first_line))


def _add_tab(run) -> None:
    """Append a <w:tab/> element inside a run."""
    tab = OxmlElement("w:tab")
    run._r.append(tab)


def _add_text(run, text: str) -> None:
    """Append a <w:t> with xml:space=preserve."""
    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = clean_xml_string(text)
    run._r.append(t)


# ---------------------------------------------------------------------------
# Document sections  (each mirrors the template XML exactly)
# ---------------------------------------------------------------------------

def _section_title(doc: Document) -> None:
    """
    Form 'DIR-8'           ← bold, centered, line-spacing 276
    Intimation by Director ← bold, centered, line-spacing 276
    [Pursuant to …]        ← not-bold, centered, line-spacing 276
    blank line
    """
    lines = [
        ("\u2018DIR-8\u2019", True),      # smart quotes like template: 'DIR-8'
        ("Intimation by Director", True),
        ("[Pursuant to Section 164(2) and rule 14(1) of Companies "
         "(Appointment and Qualification of Directors) Rules, 2014]", False),
    ]
    for i, (text, bold) in enumerate(lines):
        para = doc.add_paragraph()
        _set_jc(para, "center")
        _set_line_spacing(para)
        if i < 2:
            _set_spacing_after(para, 0)

        run = para.add_run()
        _apply_rpr(run, bold=bold)
        run.text = ("Form " if i == 0 else "") + text

    # blank separator line (bold, left, line-spacing 276 — matches template)
    blank = doc.add_paragraph()
    _set_line_spacing(blank)
    run = blank.add_run()
    _apply_rpr(run, bold=True)


def _header_field(doc: Document, label: str, value: str,
                  extra_tabs: int = 2) -> None:
    """
    Label:\t\t\tValue
    spacing after=0, Times New Roman 11.5pt
    """
    para = doc.add_paragraph()
    _set_spacing_after(para, 0)
    run = para.add_run()
    _apply_rpr(run, bold=False)
    _add_text(run, label)
    for _ in range(extra_tabs):
        _add_tab(run)
    _add_text(run, value)


def _address_field(doc: Document, address: str) -> None:
    """
    Address of Registered Office:\tValue
    Uses header-style paragraph with left indent 3600, hanging 3600 (template pattern).
    """
    para = doc.add_paragraph()
    _set_line_spacing(para)
    _set_ind(para, left=3600, hanging=3600)
    _set_spacing_after(para, 0)
    run = para.add_run()
    _apply_rpr(run, bold=False)
    _add_text(run, "Address of Registered Office: ")
    _add_tab(run)
    _add_text(run, address)


def _to_block(doc: Document, company_names: list[str]) -> None:
    for text in ["To,", "The Board of Directors"]:
        para = doc.add_paragraph()
        _set_spacing_after(para, 0)
        run = para.add_run(text)
        _apply_rpr(run, bold=False)
        
    for name in company_names:
        para = doc.add_paragraph()
        _set_spacing_after(para, 0)
        run = para.add_run(name)
        _apply_rpr(run, bold=False)

    # (Removed extra blank line here to save vertical space)


def _opening_para(doc: Document, di: DirectorInfo) -> None:
    """
    I, <Name>, son of Mr. <Father>, resident of <address/blank>, India hereby give
    notice that I am/was a Director in the following companies during the last three years:-
    Justified text, spacing after=0.
    """
    father = di.father_name if "{{" not in di.father_name else "{{FATHER_NAME}}"
    address = di.address if "{{" not in di.address else "_" * 45

    text = (
        f"I, {di.name}, son/daughter of Mr. {father}, "
        f"resident of {address}, India hereby give notice that I am/was a Director "
        f"in the following companies during the last three years:-"
    )
    para = doc.add_paragraph()
    _set_spacing_after(para, 0)
    _set_jc(para, "both")
    run = para.add_run(text)
    _apply_rpr(run, bold=False)

    # blank line before table (matches template)
    blank = doc.add_paragraph()
    _set_spacing_after(blank, 0)
    _set_jc(blank, "both")


def _set_table_borders(tbl_element) -> None:
    """Apply single-line auto borders to table — matches template tblBorders."""
    tblPr = tbl_element.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_element.insert(0, tblPr)

    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBorders.append(el)
    tblPr.append(tblBorders)

    # fixed layout
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)


def _set_col_width(cell, width_twips: int) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcW  = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"),    str(width_twips))
    tcW.set(qn("w:type"), "dxa")


def _cell_para(cell, text: str, bold: bool = False,
               jc: str = "both", spacing_after: int = 0) -> None:
    para = cell.paragraphs[0]
    _set_spacing_after(para, spacing_after)
    _set_jc(para, jc)
    if text:
        run = para.add_run(text)
        _apply_rpr(run, bold=bold)


def _companies_table(doc: Document, di: DirectorInfo) -> None:
    """
    4-column table matching template exactly:
    Col widths: 636 | 3864 | 2610 | 2250 DXA
    Header row: bold text
    Data rows: one per company
    Minimum 3 data rows (blank rows added if fewer entries).
    """
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set table width and borders via XML
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    str(9360))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    tblInd = OxmlElement("w:tblInd")
    tblInd.set(qn("w:w"),    "-5")
    tblInd.set(qn("w:type"), "dxa")
    tblPr.append(tblInd)

    _set_table_borders(tbl)

    # Set grid columns
    tblGrid = OxmlElement("w:tblGrid")
    for w in (636, 3864, 2610, 2250):
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        tblGrid.append(gc)
    tbl.insert(list(tbl).index(tbl.find(qn("w:tr"))), tblGrid)

    # -- Header row --
    hrow = table.rows[0]
    col_widths = [636, 3864, 2610, 2250]
    headers    = [
        ("Sr. No.",                                                   "center"),
        ("Names of the Companies / bodies corporate / firms /association of individuals ", "both"),
        ("Date of appointment ",                                       "both"),
        ("Date of Cessation ",                                         "both"),
    ]
    for i, (cell, (hdr_text, jc_val)) in enumerate(zip(hrow.cells, headers)):
        _set_col_width(cell, col_widths[i])
        _cell_para(cell, hdr_text, bold=True, jc=jc_val)

    # -- Data rows — minimum 3 --
    entries = di.other_companies if di.other_companies else []
    min_rows = max(3, len(entries))

    for idx in range(min_rows):
        row = table.add_row()
        sr_cell   = row.cells[0]
        name_cell = row.cells[1]
        appt_cell = row.cells[2]
        cess_cell = row.cells[3]

        _set_col_width(sr_cell,   636)
        _set_col_width(name_cell, 3864)
        _set_col_width(appt_cell, 2610)
        _set_col_width(cess_cell, 2250)

        # Sr. No. cell — left-indented like template (ind left=360)
        sr_para = sr_cell.paragraphs[0]
        _set_spacing_after(sr_para, 0)
        _set_ind(sr_para, left=360)
        sr_run = sr_para.add_run(str(idx + 1))
        _apply_rpr(sr_run, bold=False)

        if idx < len(entries):
            entry = entries[idx]
            _cell_para(name_cell, entry.com_name,         bold=False, jc="both")
            _cell_para(appt_cell, entry.appointment_date, bold=False, jc="center")
            _cell_para(cess_cell, entry.cessation_date,   bold=False, jc="center")
        else:
            _cell_para(name_cell, "", jc="both")
            _cell_para(appt_cell, "", jc="center")
            _cell_para(cess_cell, "", jc="center")


def _confirmation_para(doc: Document) -> None:
    """
    Blank line + confirmation text (justified) + blank lines before signature.
    """
    # blank after table
    blank = doc.add_paragraph()
    _set_spacing_after(blank, 0)
    _set_jc(blank, "both")

    # confirmation
    para = doc.add_paragraph()
    _set_spacing_after(para, 0)
    _set_jc(para, "both")
    run = para.add_run(
        "I further confirm that I have not incurred disqualification under section 164(2) "
        "of the Companies Act, 2013 in any of the above companies, in the previous financial "
        "year and that I at present stand free from any disqualification from being a director. "
    )
    _apply_rpr(run, bold=False)

    # blank line before signature (reduced to 1 to keep on first page)
    b = doc.add_paragraph()
    _set_spacing_after(b, 0)
    _set_jc(b, "both")


def _signature_block(doc: Document, di: DirectorInfo, co: CompanyInfo) -> None:
    """
    Signature  :\t\t\t\t\t (right-indented, WW-PlainText style)
    Name       :\t<Name>
    DIN\t\t:\t<DIN>
    Place: <place>
    Date: ___________, 2025
    """
    # (Removed extra blank line here to save vertical space)

    # Signature line — ind left=2880, firstLine=720 like template
    sig = doc.add_paragraph()
    _set_line_spacing(sig)
    _set_ind(sig, left=2880, first_line=720)
    sig_run = sig.add_run()
    _apply_rpr(sig_run, bold=False)
    _add_text(sig_run, "Signature  ")
    _add_tab(sig_run)
    _add_text(sig_run, ":                      ")

    # Name — ind left=3600
    name_para = doc.add_paragraph()
    _set_line_spacing(name_para)
    _set_ind(name_para, left=3600)
    name_run = name_para.add_run()
    _apply_rpr(name_run, bold=False)
    _add_text(name_run, "Name           ")
    _add_tab(name_run)
    _add_text(name_run, f":  {di.name}")

    # DIN — ind left=3600, two tabs before colon (matches template)
    din_para = doc.add_paragraph()
    _set_line_spacing(din_para)
    _set_ind(din_para, left=3600)
    din_run = din_para.add_run()
    _apply_rpr(din_run, bold=False)
    _add_text(din_run, "DIN")
    _add_tab(din_run)
    _add_tab(din_run)
    _add_text(din_run, f":  {di.din}")

    # Place — no indent, spacing after=0
    place_para = doc.add_paragraph()
    _set_spacing_after(place_para, 0)
    place_run = place_para.add_run(f"Place: {co.signature_place}")
    _apply_rpr(place_run, bold=False)

    # Date
    date_para = doc.add_paragraph()
    _set_spacing_after(date_para, 0)
    _set_jc(date_para, "both")
    date_run = date_para.add_run()
    _apply_rpr(date_run, bold=False)
    date_str = co.signature_date if "{{" not in co.signature_date else "___________, 2025"
    _add_text(date_run, f"Date: {date_str}")
    _add_tab(date_run)


# ---------------------------------------------------------------------------
# Page setup
# ---------------------------------------------------------------------------

def _setup_page(doc: Document) -> None:
    """US Letter, 1-inch margins — exactly matching template sectPr."""
    sec = doc.sections[0]
    sec.page_width    = PAGE_W
    sec.page_height   = PAGE_H
    sec.top_margin    = MARGIN
    sec.bottom_margin = MARGIN
    sec.left_margin   = MARGIN
    sec.right_margin  = MARGIN


def _setup_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_PT


# ---------------------------------------------------------------------------
# Main builder — one document per director
# ---------------------------------------------------------------------------

def build_dir8(co: CompanyInfo, di: DirectorInfo) -> Document:
    doc = Document()
    _setup_page(doc)
    _setup_styles(doc)

    _section_title(doc)

    # Header fields — tab-aligned, spacing after=0
    _header_field(doc, "Registration No: ",    co.cin,          extra_tabs=3)
    _header_field(doc, "Nominal Capital: ",    co.auth_capital, extra_tabs=3)
    _header_field(doc, "Paid-up Capital: ",    co.paid_capital, extra_tabs=3)
    _header_field(doc, "Name of the Company: ", ", ".join(co.target_company_names), extra_tabs=2)
    _address_field(doc, co.address)

    _to_block(doc, co.target_company_names)
    _opening_para(doc, di)
    _companies_table(doc, di)
    _confirmation_para(doc)
    _signature_block(doc, di, co)

    return doc


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate Form DIR-8 Word documents from Azure PostgreSQL Database."
    )
    # DB Mode
    parser.add_argument("--din",    help="Director DIN")
    parser.add_argument("--cin",    help="Target Company CIN (comma separated for multiple)")
    parser.add_argument("--all",    action="store_true", help="Generate for all directors across all companies")
    parser.add_argument("--year",   default="2024-25",   help="Fiscal year for output folder")
    
    # Legacy / Overlay
    parser.add_argument("--input",  help="Path to company JSON file (legacy mode)")
    parser.add_argument("--date",   help="Signature date e.g. '1st April, 2025'")
    parser.add_argument("--output", help="Explicit output path")

    args = parser.parse_args()

    pairs = []
    if args.all:
        print("[BATCH] Fetching all Director-Company pairs from DB...")
        pairs = get_all_pairs_from_db()
    elif args.din and args.cin:
        # Split CINs if multiple provided
        cins = [c.strip() for c in args.cin.split(",")]
        # For multi-company mode, we treat them as one set
        pairs = [(args.din, cins)]
    elif args.input:
        # Legacy mode
        input_path = Path(args.input)
        if not input_path.exists():
            print(f"[ERROR] File not found: {input_path}")
            return
        raw = json.loads(input_path.read_text(encoding="utf-8"))
        if args.date: raw["signature_date"] = args.date
        co, directors = parse_input(raw, args.din)
        # Special case for legacy - just wrap in a dummy list to reuse logic
        _process_list(co, directors, args.year)
        return
    else:
        parser.print_help()
        return

    print(f"[INFO] Found {len(pairs)} pairs to process.")
    
    total = len(pairs)
    total_generated = 0
    for i, (din, cin_data) in enumerate(pairs, 1):
        # cin_data can be a single string or a list
        cins = cin_data if isinstance(cin_data, list) else [cin_data]
        co, directors = fetch_full_data_from_db(din, cins, args.date)
        if co and directors:
            report_progress(i, total, f"Generating for {directors[0].name}")
            _process_list(co, directors, args.year)
            total_generated += 1
        else:
            report_progress(i, total, f"Skipping DIN {din}")

    report_progress(total, total, "Complete")
    print(f"\n[FINISH] Total {total_generated} document(s) generated.")

def register_document_in_db(di: DirectorInfo, file_path: str):
    """Registers the generated file in the document_summaries table for the repository."""
    conn = get_db_connection()
    if not conn: return
    try:
        cur = conn.cursor()
        # Relative path from Output_Disclosures
        rel_path = str(Path(file_path).relative_to(Path("Output_Disclosures")))
        
        cur.execute("""
            INSERT INTO directors_data.document_summaries 
            (director_name, din, file_path, created_at)
            VALUES (%s, %s, %s, CURRENT_TIMESTAMP)
            ON CONFLICT (file_path) DO UPDATE 
            SET created_at = EXCLUDED.created_at,
                director_name = EXCLUDED.director_name,
                din = EXCLUDED.din
        """, (di.name, di.din, rel_path))
        conn.commit()
        print(f"  [DB] Registered document for {di.name}")
    except Exception as e:
        print(f"  [DB_ERROR] Failed to register document: {e}")
    finally:
        conn.close()

def _process_list(co: CompanyInfo, directors: list[DirectorInfo], year: str):
    base_output = Path("Output_Disclosures") / year
    
    for di in directors:
        doc = build_dir8(co, di)
        
        # Clean names for folder/file
        clean_company = "".join(c if c.isalnum() else "_" for c in co.company_name).strip("_")
        clean_director = "".join(c if c.isalnum() else "_" for c in di.name).strip("_")
        
        # Directory structure: Output/Year/Company/DIR-8/
        target_dir = base_output / clean_company / "DIR-8"
        target_dir.mkdir(parents=True, exist_ok=True)
        
        file_name = f"DIR8_{clean_director}_{di.din}.docx"
        out_path = target_dir / file_name
        
        doc.save(str(out_path))
        register_document_in_db(di, str(out_path))
        print(f"  [OK] {clean_company} -> {clean_director}")

if __name__ == "__main__":
    main()


if __name__ == "__main__":
    main()