"""
Script to analyze the structure of DOCX files and understand data extraction patterns.
"""

import docx
import re

def analyze_docx_structure(file_path):
    """Analyze the structure of a DOCX file to understand data extraction patterns."""
    doc = docx.Document(file_path)
    
    print(f"Analyzing file: {file_path}")
    print("=" * 50)
    
    # Print document properties
    print("Document Properties:")
    try:
        print(f"Title: {doc.core_properties.title}")
        print(f"Author: {doc.core_properties.author}")
        print(f"Created: {doc.core_properties.created}")
    except:
        print("No core properties found")
    
    print("\nParagraph Analysis:")
    print("-" * 30)
    
    # Analyze paragraphs
    for i, paragraph in enumerate(doc.paragraphs[:20]):  # First 20 paragraphs
        text = paragraph.text.strip()
        if text:  # Only print non-empty paragraphs
            print(f"Paragraph {i}: {text[:100]}{'...' if len(text) > 100 else ''}")
    
    print("\nTable Analysis:")
    print("-" * 30)
    
    # Analyze tables
    for i, table in enumerate(doc.tables):
        print(f"Table {i}: {len(table.rows)} rows, {len(table.columns)} columns")
        # Print first few cells of the first row
        if table.rows:
            first_row = table.rows[0]
            cell_texts = [cell.text.strip() for cell in first_row.cells[:5]]
            print(f"  First row cells: {cell_texts}")
    
    print("\nText Search for Patterns:")
    print("-" * 30)
    
    # Combine all text for pattern searching
    full_text = "\n".join([p.text for p in doc.paragraphs])
    
    # Search for DIN patterns
    din_patterns = [
        r"DIN\s*:\s*(\d+)",
        r"Director Identification Number\s*:\s*(\d+)",
        r"(\d{8})"  # 8-digit number pattern
    ]
    
    for pattern in din_patterns:
        matches = re.findall(pattern, full_text, re.IGNORECASE)
        if matches:
            print(f"DIN pattern '{pattern}': {matches[:3]}")  # Show first 3 matches
    
    # Search for company-related patterns
    company_patterns = [
        r"([\w\s&.,'-]+Limited)",
        r"([\w\s&.,'-]+Pvt\.?\s+Ltd\.?)",
        r"Company\s*:\s*([\w\s&.,'-]+)"
    ]
    
    for pattern in company_patterns:
        matches = re.findall(pattern, full_text, re.IGNORECASE)
        if matches:
            print(f"Company pattern '{pattern}': {matches[:3]}")  # Show first 3 matches
    
    # Search for position patterns
    position_patterns = [
        r"(Director|Whole-time Director|Additional Director)",
        r"Position\s*:\s*([\w\s-]+)"
    ]
    
    for pattern in position_patterns:
        matches = re.findall(pattern, full_text, re.IGNORECASE)
        if matches:
            print(f"Position pattern '{pattern}': {matches[:5]}")  # Show first 5 matches

if __name__ == "__main__":
    # Analyze the first director file
    file_path = "public/Directors Discloser Output/Abdul Ishad Khan_MBP.docx"
    analyze_docx_structure(file_path)