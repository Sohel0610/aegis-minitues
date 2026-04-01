import os
import logging
import zipfile
from xml.etree import ElementTree as ET
from datetime import datetime

# Load environment variables
from dotenv import load_dotenv
load_dotenv(dotenv_path=os.path.join(os.path.dirname(__file__), ".env"))

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Import PostgreSQL service (mandatory for production)
from utils.pgsql_service import get_pg_connection, get_pg_cursor

# Try to import docx, handle if not available
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    logger.warning("python-docx not available. Will use fallback text extraction.")
    DOCX_AVAILABLE = False

def extract_text_from_docx_fallback(file_path):
    """Extract text content from a DOCX file using fallback method (zip+xml)"""
    try:
        with zipfile.ZipFile(file_path, 'r') as docx:
            xml_content = docx.read('word/document.xml')
            tree = ET.fromstring(xml_content)
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            paragraphs = tree.findall('.//w:p', ns)
            text_parts = []
            for paragraph in paragraphs:
                texts = paragraph.findall('.//w:t', ns)
                paragraph_text = ''.join([t.text for t in texts if t.text])
                if paragraph_text.strip():
                    text_parts.append(paragraph_text)
            return '\n'.join(text_parts)
    except Exception as e:
        logger.error(f"Error extracting text from {file_path} using fallback method: {e}")
        return ""

def extract_text_from_docx(file_path):
    """Extract text content from a DOCX file"""
    if DOCX_AVAILABLE:
        try:
            doc = DocxDocument(file_path)
            content_parts = []
            for para in doc.paragraphs:
                if para.text.strip(): content_parts.append(para.text)
            if doc.tables:
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        content_parts.append(row_text)
            return "\n".join(content_parts)
        except Exception as e:
            logger.error(f"Error extracting text from {file_path} using python-docx: {e}")
            return extract_text_from_docx_fallback(file_path)
    return extract_text_from_docx_fallback(file_path)

def generate_summary_with_groq(content, max_tokens=1000):
    """Generate a summary of the content using Groq LLM"""
    try:
        from groq import Groq
        api_key = os.environ.get('GROQ_API_KEY')
        client = Groq(api_key=api_key) if api_key else Groq()
        
        prompt = f"""
        Please provide a concise summary of the following director's disclosure document. 
        Focus on: Director's name and DIN, Companies and positions, Shareholding (Active only), Disclosures, Concerns.
        Format requirements: Plain text (no markdown), use section headers with colon, bullet points with '-', concise.
        
        {content[:8000]}
        """
        
        completion = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "system", "content": "You are an expert at corporate document summarization. Use plain text only."},
                     {"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=max_tokens
        )
        return completion.choices[0].message.content.strip() if completion.choices[0].message.content else "No summary available"
    except Exception as e:
        logger.error(f"Error generating summary with Groq: {e}")
        return "Error generating summary with LLM"

def generate_summary_with_azure_openai(content, max_tokens=1000):
    """Generate a summary using Azure OpenAI."""
    # Simplified version for now - user wants Postgres focus
    # (Assuming Azure env vars are set)
    return "Azure OpenAI summary not fully implemented in this refactor. Use Groq."

def generate_summary(content, max_tokens=1000):
    """Generate a summary using available LLM."""
    if os.environ.get('USE_GROQ', 'true').lower() == 'true':
        return generate_summary_with_groq(content, max_tokens)
    return generate_summary_with_azure_openai(content, max_tokens)

def save_summary_to_db(director_name, din, file_path, full_text, summary):
    """Save full text and summary to PostgreSQL exclusively in the Director database."""
    pg_conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_DIRECTOR'))
    if pg_conn:
        try:
            cursor = get_pg_cursor(pg_conn)
            # Remove schema prefix for unified architecture
            cursor.execute("""
                INSERT INTO document_summaries (director_name, din, file_path, full_text, summary, updated_at)
                VALUES (%s, %s, %s, %s, %s, CURRENT_TIMESTAMP)
                ON CONFLICT (file_path)
                DO UPDATE SET
                    director_name = EXCLUDED.director_name,
                    din = EXCLUDED.din,
                    full_text = EXCLUDED.full_text,
                    summary = EXCLUDED.summary,
                    updated_at = CURRENT_TIMESTAMP
            """, (director_name, din, file_path, full_text, summary))
            pg_conn.commit()
            logger.info(f"Summary saved to PostgreSQL for {director_name}")
            return True
        except Exception as e:
            pg_conn.rollback()
            logger.error(f"Failed to save summary to PG: {e}")
        finally:
            pg_conn.close()
    return False

def get_summary_from_db(file_path):
    """Retrieve summary from PostgreSQL."""
    pg_conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_DIRECTOR'))
    if pg_conn:
        try:
            cursor = get_pg_cursor(pg_conn)
            cursor.execute("SELECT summary FROM document_summaries WHERE file_path = %s", (file_path,))
            res = cursor.fetchone()
            return res["summary"] if res else None
        finally:
            pg_conn.close()
    return None

def generate_and_save_summary(director_name, din, file_path):
    """Extract, generate, and save to PostgreSQL."""
    try:
        full_file_path = os.path.join(os.path.dirname(__file__), "public", "Directors Discloser Output", file_path)
        if not os.path.exists(full_file_path): return "File not found", "File not found"
        
        full_text = extract_text_from_docx(full_file_path)
        if not full_text.strip():
            summary = "Could not extract content"
        else:
            summary = generate_summary(full_text)
            
        save_summary_to_db(director_name, din, file_path, full_text, summary)
        return full_text, summary
    except Exception as e:
        logger.error(f"Error: {e}")
        return "Error processing document", "Error processing document"
