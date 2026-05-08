"""
Director Data Analysis Module
This module handles parsing of MBP-1 DOCX documents, extracting director information,
normalizing the data, and storing it in a PostgreSQL database (now mandatory).
"""

import os
import re
from datetime import datetime
from docx import Document
from typing import List, Dict, Tuple, Optional
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Import our PostgreSQL service
from utils.pgsql_service import get_pg_connection, get_pg_cursor

# Database schema in PostgreSQL (Unified Master Schema)
DB_SCHEMA = "directors_data"

def init_database():
    """Verify the PostgreSQL database schema and tables (Unified)."""
    # Use dedicated Directors database pointer
    pg_conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_DIRECTOR'))
    if pg_conn:
        try:
            cursor = get_pg_cursor(pg_conn)

            # Ensure schemas exist
            cursor.execute("CREATE SCHEMA IF NOT EXISTS directors_data")
            cursor.execute("CREATE SCHEMA IF NOT EXISTS directors_master")
            cursor.execute("CREATE SCHEMA IF NOT EXISTS directors_profile")

            # 1. Master Directors Table (Registry)
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS directors_master.directors (
                    id SERIAL PRIMARY KEY,
                    name TEXT NOT NULL,
                    din TEXT UNIQUE,
                    created_at TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP
                )
            """)

            # 2. Directors Data table (for analysis extraction)
            cursor.execute(f"""
                CREATE TABLE IF NOT EXISTS {DB_SCHEMA}.directors (
                    din TEXT PRIMARY KEY,
                    name TEXT,
                    source_file TEXT
                )
            """)

            # 3. Master Companies Table
            cursor.execute(f"""
                CREATE TABLE IF NOT EXISTS {DB_SCHEMA}.companies (
                    id SERIAL PRIMARY KEY,
                    name TEXT UNIQUE,
                    type TEXT
                )
            """)

            # 4. Directorships table
            cursor.execute(f"""
                CREATE TABLE IF NOT EXISTS {DB_SCHEMA}.directorships (
                    id SERIAL PRIMARY KEY,
                    din TEXT REFERENCES {DB_SCHEMA}.directors(din) ON DELETE CASCADE,
                    company_id INTEGER REFERENCES {DB_SCHEMA}.companies(id) ON DELETE CASCADE,
                    position TEXT,
                    appointment_date TEXT
                )
            """)

            # 5. Director Profile Table
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS directors_profile.directors_profile (
                    id SERIAL PRIMARY KEY,
                    din TEXT UNIQUE,
                    pan TEXT,
                    name_of_director TEXT,
                    address TEXT,
                    date_of_birth DATE,
                    qualification TEXT,
                    experience TEXT,
                    created_at TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP
                )
            """)

            # 6. Document Summaries Table
            cursor.execute(f"""
                CREATE TABLE IF NOT EXISTS {DB_SCHEMA}.document_summaries (
                    id SERIAL PRIMARY KEY,
                    director_name TEXT NOT NULL,
                    din TEXT,
                    file_path TEXT NOT NULL UNIQUE,
                    full_text TEXT,
                    summary TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)

            # 7. External Board Members — non-Adani directors discovered via company API sync
            #    Deliberately separate from directors_master.directors to preserve roster integrity
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS directors_master.external_board_members (
                    id SERIAL PRIMARY KEY,
                    din TEXT NOT NULL,
                    name TEXT,
                    cin TEXT NOT NULL,
                    company_name TEXT,
                    designation TEXT,
                    appointment_date TEXT,
                    status TEXT,
                    source TEXT DEFAULT 'COMPANY_API',
                    created_at TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP,
                    UNIQUE(din, cin)
                )
            """)
            # Ensure status column exists for existing tables
            cursor.execute("ALTER TABLE directors_master.external_board_members ADD COLUMN IF NOT EXISTS status TEXT")

            # Create family information schema and table
            cursor.execute("CREATE SCHEMA IF NOT EXISTS family_information")
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS family_information.director_family (
                    id SERIAL PRIMARY KEY,
                    director_name TEXT NOT NULL,
                    section_2_77_i TEXT,
                    section_2_77_ii TEXT,
                    section_2_77_iii TEXT,
                    father TEXT,
                    mother TEXT,
                    son TEXT,
                    sons_wife TEXT,
                    daughter TEXT,
                    daughters_husband TEXT,
                    brother TEXT,
                    sister TEXT,
                    father_pan TEXT,
                    mother_pan TEXT,
                    father_pan_file TEXT,
                    mother_pan_file TEXT,
                    is_submitted INTEGER DEFAULT 0,
                    created_at TIMESTAMPTZ DEFAULT NOW(),
                    updated_at TIMESTAMPTZ DEFAULT NOW()
                )
            """)

            # Create indexes
            cursor.execute(f"CREATE INDEX IF NOT EXISTS idx_doc_summ_file_path ON {DB_SCHEMA}.document_summaries (file_path)")
            cursor.execute(f"CREATE INDEX IF NOT EXISTS idx_doc_summ_dir_name ON {DB_SCHEMA}.document_summaries (director_name)")
            cursor.execute("CREATE INDEX IF NOT EXISTS idx_directors_master_din ON directors_master.directors(din)")
            cursor.execute("CREATE INDEX IF NOT EXISTS idx_directors_profile_din ON directors_profile.directors_profile(din)")

            pg_conn.commit()
            logger.info(f"PostgreSQL comprehensive schemas for Directors ({DB_SCHEMA}, master, profile) initialized successfully")
            return True

        except Exception as e:
            logger.error(f"PostgreSQL initialization failed: {e}")
            try:
                pg_conn.rollback()
            except Exception:
                pass
            raise RuntimeError(f"Database initialization failed: {e}")
        finally:
            pg_conn.close()

    else:
        logger.error("Could not get a database connection for initialization")
        raise RuntimeError("Database connection failed during initialization")

def extract_director_info(doc_path: str) -> Dict:
    """Extract director information from a DOCX file."""
    try:
        doc = Document(doc_path)
        filename = os.path.basename(doc_path)
        
        # Extract director name from filename (before _MBP.docx)
        director_name = filename.replace("_MBP.docx", "").replace("_", " ")
        
        # Extract DIN from document content
        din = None
        
        # Convert document to text
        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        
        full_text_str = "\n".join(full_text)
        
        # Look for DIN pattern (typically a number)
        din_match = re.search(r"DIN\s*:\s*(\d+)", full_text_str, re.IGNORECASE)
        if din_match:
            din = din_match.group(1)
        else:
            din_match = re.search(r"(\d{8})", full_text_str)
            if din_match:
                din = din_match.group(1)
        
        # Extract companies information
        companies = extract_companies_info(doc, full_text_str)
        
        return {
            "name": director_name,
            "din": din,
            "source_file": filename,
            "companies": companies
        }
    except Exception as e:
        logger.error(f"Error extracting info from {doc_path}: {str(e)}")
        return {"name": "Unknown", "din": None, "source_file": os.path.basename(doc_path), "companies": []}

def extract_companies_info(doc, full_text_str: str) -> List[Dict]:
    """Extract companies information from document tables and sections."""
    companies = []
    company_types = extract_company_types_from_sections(full_text_str)
    table_companies = extract_companies_from_tables(doc)
    
    for company in table_companies:
        company_name = company.get("name", "")
        if company_name:
            company_type = "Unknown"
            for section_type, section_companies in company_types.items():
                for section_company in section_companies:
                    if re.sub(r'\s+', ' ', company_name.strip()).upper() == re.sub(r'\s+', ' ', section_company.strip()).upper():
                        company_type = section_type
                        break
                if company_type != "Unknown":
                    break
            company["type"] = company_type
            companies.append(company)
    
    return companies

def extract_company_types_from_sections(full_text_str: str) -> Dict[str, List[str]]:
    """Extract company types from the section headers in the document."""
    company_types = {
        "Public": [],
        "Private - Subsidiary of Public": [],
        "Private - Not Subsidiary of Public": []
    }
    full_text_str = full_text_str.replace('\r\n', '\n').replace('\r', '\n')
    
    sections = {
        "Public": r"\(A\)\s*Public Limited Companies:\s*(.*?)(?=\n\([B-Z]\)|$)",
        "Private - Subsidiary of Public": r"\(B\)\s*Private Limited Companies which are subsidiary\(ies\) of Public Companies:\s*(.*?)(?=\n\([C-Z]\)|$)",
        "Private - Not Subsidiary of Public": r"\(C\)\s*Private Limited Companies which are not subsidiary\(ies\) of Public Companies:\s*(.*?)(?=\n\([D-Z]\)|$)"
    }
    
    for key, pattern in sections.items():
        match = re.search(pattern, full_text_str, re.DOTALL | re.IGNORECASE)
        if match:
            batch = match.group(1).strip()
            if re.search(r'\bNIL\b', batch, re.IGNORECASE):
                continue
            candidates = re.findall(r'^.*?LIMITED.*?$', batch, re.MULTILINE | re.IGNORECASE)
            others = re.findall(r'^[A-Z][A-Z\s\(\)&\-\.]*?(?:FOUNDATION|LTD|PRIVATE|PUBLIC|COMPANY|CORPORATION|INC|LLC|LLP)[A-Z\s\(\)&\-\.]*?$', batch, re.MULTILINE)
            for c in candidates + others:
                c = re.sub(r'\s+', ' ', c).strip()
                if c and not re.search(r'\bNIL\b', c, re.IGNORECASE) and c.upper() not in ["LIMITED", "LTD"]:
                    company_types[key].append(c)
    
    return company_types

def extract_companies_from_tables(doc) -> List[Dict]:
    """Extract companies information from document tables."""
    companies = []
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        if any("company" in h.lower() or "name" in h.lower() for h in headers):
            for i in range(1, len(table.rows)):
                cells = [cell.text.strip() for cell in table.rows[i].cells]
                if not any(cells): continue
                cinfo = extract_company_info_from_row(cells, headers)
                if cinfo and cinfo.get("name"):
                    companies.append(cinfo)
    return companies

def extract_company_info_from_row(cells: List[str], headers: List[str]) -> Dict:
    """Extract company information from a table row."""
    header_to_value = {h.lower(): (cells[i] if i < len(cells) else "") for i, h in enumerate(headers)}
    company_name = None
    for k, v in header_to_value.items():
        if ("company" in k or "name" in k) and v and not re.search(r'\bNIL\b', v, re.IGNORECASE):
            company_name = v
            break
    if not company_name:
        for v in cells:
            if v and re.match(r'^[A-Z][A-Z\s\(\)&\-\.]*?(?:LIMITED|FOUNDATION|LTD|PRIVATE|PUBLIC|COMPANY|CORPORATION|INC|LLC|LLP)', v, re.IGNORECASE):
                company_name = v
                break
    if not company_name: return {}
    
    position = "Director"
    for k, v in header_to_value.items():
        if ("position" in k or "nature" in k or "type" in k) and v:
            if "whole-time" in v.lower() or "whole time" in v.lower(): position = "Whole-time Director"
            elif "additional" in v.lower(): position = "Additional Director"
            elif "director" in v.lower(): position = "Director"
            else: position = v
            break
            
    appointment_date = ""
    for k, v in header_to_value.items():
        if "date" in k:
            dm = re.search(r"(\d{1,2}[/-]\d{1,2}[/-]\d{4})", v)
            if dm: appointment_date = dm.group(1); break
            
    return {"name": company_name, "position": position, "type": "Unknown", "appointment_date": appointment_date}

def normalize_company_name(name: str) -> str:
    """Normalize company name."""
    name = re.sub(r'\s+', ' ', name.strip())
    name = re.sub(r'^[^\w]+|[^\w]+$', '', name)
    return name.replace('\n', ' ')

def normalize_position(position: str) -> str:
    """Normalize position values."""
    if not position: return "Director"
    p = position.lower().strip()
    if any(x in p for x in ("whole-time", "whole time")): return "Whole-time Director"
    if "additional" in p: return "Additional Director"
    if "director" in p: return "Director"
    return position.title()

def store_director_data(director_info: Dict):
    """Store director data in PostgreSQL (mandatory)."""
    pg_conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_DIRECTOR'))
    if pg_conn:
        try:
            cursor = get_pg_cursor(pg_conn)
            if director_info.get("din"):
                cursor.execute(f"""
                    INSERT INTO {DB_SCHEMA}.directors (din, name, source_file)
                    VALUES (%s, %s, %s)
                    ON CONFLICT (din) DO UPDATE SET
                        name = EXCLUDED.name,
                        source_file = EXCLUDED.source_file
                """, (director_info["din"], director_info.get("name"), director_info.get("source_file")))

            for company in director_info.get("companies", []):
                if "name" not in company: continue
                cname = normalize_company_name(company.get("name", ""))
                if not cname: continue
                ctype = company.get("type", "Unknown")
                pos = normalize_position(company.get("position", "Director"))
                adate = company.get("appointment_date", "")

                cursor.execute(f"SELECT id, type FROM {DB_SCHEMA}.companies WHERE name = %s", (cname,))
                existing = cursor.fetchone()
                if existing:
                    cid = existing["id"]
                    if existing["type"] == "Unknown" and ctype != "Unknown":
                        cursor.execute(f"UPDATE {DB_SCHEMA}.companies SET type = %s WHERE id = %s", (ctype, cid))
                else:
                    cursor.execute(f"INSERT INTO {DB_SCHEMA}.companies (name, type) VALUES (%s, %s) RETURNING id", (cname, ctype))
                    cid = cursor.fetchone()["id"]

                if director_info.get("din"):
                    cursor.execute(f"""
                        SELECT 1 FROM {DB_SCHEMA}.directorships
                        WHERE din = %s AND company_id = %s AND position = %s
                    """, (director_info["din"], cid, pos))
                    if not cursor.fetchone():
                        cursor.execute(f"""
                            INSERT INTO {DB_SCHEMA}.directorships (din, company_id, position, appointment_date)
                            VALUES (%s, %s, %s, %s)
                        """, (director_info["din"], cid, pos, adate))
            pg_conn.commit()
            logger.info(f"Stored data for director (PostgreSQL): {director_info.get('name')}")
        except Exception as e:
            pg_conn.rollback()
            logger.error(f"PostgreSQL store failed for {director_info.get('name')}: {e}")
        finally:
            pg_conn.close()

def process_all_director_files(directory_path: str):
    """Process all director DOCX files in the given directory."""
    init_database()
    docx_files = [f for f in os.listdir(directory_path) if f.endswith('.docx')]
    for filename in docx_files:
        info = extract_director_info(os.path.join(directory_path, filename))
        store_director_data(info)

def get_all_directors():
    """Get all directors from PostgreSQL."""
    pg_conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_DIRECTOR'))
    if pg_conn:
        try:
            cursor = get_pg_cursor(pg_conn)
            cursor.execute(f"SELECT din, name, source_file FROM {DB_SCHEMA}.directors")
            rows = cursor.fetchall()
            return [{"din": r["din"], "name": r["name"], "source_file": r["source_file"]} for r in rows]
        finally:
            pg_conn.close()
    return []

def get_company_count():
    """Get company count statistics from Enriched External Registry."""
    pg_conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_DIRECTOR'))
    if pg_conn:
        try:
            cursor = get_pg_cursor(pg_conn)
            # Total unique companies in the registry
            cursor.execute("SELECT COUNT(DISTINCT cin) AS count FROM directors_master.external_board_members")
            total = cursor.fetchone()["count"] or 0
            
            # Listed companies (CIN starts with 'L')
            cursor.execute("SELECT COUNT(DISTINCT cin) AS count FROM directors_master.external_board_members WHERE cin LIKE 'L%'")
            public = cursor.fetchone()["count"] or 0
            
            # Unlisted/Private companies (CIN starts with 'U' or other)
            private = total - public
            
            return {"total": int(total), "public": int(public), "private": int(private)}
        finally:
            pg_conn.close()
    return {"total": 0, "public": 0, "private": 0}

def get_cross_directorship():
    """Get global cross-directorship information using External Registry associations."""
    pg_conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_DIRECTOR'))
    if pg_conn:
        try:
            cursor = get_pg_cursor(pg_conn)
            cursor.execute("""
                SELECT d.name, d.din, COUNT(ea.cin) AS company_count
                FROM directors_master.directors d
                JOIN directors_master.external_board_members ea ON d.din = ea.din
                GROUP BY d.name, d.din
                ORDER BY company_count DESC
                LIMIT 100
            """)
            return [{"name": r["name"], "din": r["din"], "company_count": r["company_count"]} for r in cursor.fetchall()]
        finally:
            pg_conn.close()
    return []

def get_clustering():
    """Get director clustering information (shared companies) using global associations."""
    pg_conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_DIRECTOR'))
    if pg_conn:
        try:
            cursor = get_pg_cursor(pg_conn)
            cursor.execute("""
                SELECT d1.name AS director1, d2.name AS director2, COUNT(ea1.cin) AS shared_companies
                FROM directors_master.external_board_members ea1
                JOIN directors_master.external_board_members ea2 ON ea1.cin = ea2.cin AND ea1.din < ea2.din
                JOIN directors_master.directors d1 ON ea1.din = d1.din
                JOIN directors_master.directors d2 ON ea2.din = d2.din
                GROUP BY d1.name, d2.name
                ORDER BY shared_companies DESC
                LIMIT 30
            """)
            return [{"director1": r["director1"], "director2": r["director2"], "sharedCompanies": r["shared_companies"]} for r in cursor.fetchall()]
        finally:
            pg_conn.close()
    return []

def get_network():
    """Get global network data for visualization."""
    pg_conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_DIRECTOR'))
    if pg_conn:
        try:
            cursor = get_pg_cursor(pg_conn)
            # Fetch most connected directors
            cursor.execute("SELECT din, name FROM directors_master.directors LIMIT 150")
            directors = cursor.fetchall()
            
            # Fetch companies with most directors
            cursor.execute("""
                SELECT cin as id, company_name as name 
                FROM directors_master.external_board_members 
                GROUP BY cin, company_name 
                ORDER BY COUNT(din) DESC 
                LIMIT 80
            """)
            companies = cursor.fetchall()
            
            # Fetch links
            cursor.execute("SELECT din, cin FROM directors_master.external_board_members LIMIT 800")
            links = cursor.fetchall()
            
            nodes = []
            for d in directors: nodes.append({"id": d["din"], "type": "director", "label": d["name"]})
            for c in companies: nodes.append({"id": c["id"], "type": "company", "label": c["name"]})
            
            link_data = []
            for l in links:
                link_data.append({"source": l["din"], "target": l["cin"]})
                
            return {"nodes": nodes, "links": link_data}
        finally:
            pg_conn.close()
    return {"nodes": [], "links": []}

def get_wtd_count():
    """Get global positions count from External Registry."""
    pg_conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_DIRECTOR'))
    if pg_conn:
        try:
            cursor = get_pg_cursor(pg_conn)
            cursor.execute("""
                SELECT d.name, COUNT(ea.id) AS positions
                FROM directors_master.directors d
                JOIN directors_master.external_board_members ea ON d.din = ea.din
                GROUP BY d.name
                ORDER BY positions DESC
                LIMIT 50
            """)
            return [dict(r) for r in cursor.fetchall()]
        finally:
            pg_conn.close()
    return []

def get_all_companies_with_director_count():
    """Get all companies with their director counts from External Registry (Largest Boards)."""
    pg_conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_DIRECTOR'))
    if pg_conn:
        try:
            cursor = get_pg_cursor(pg_conn)
            cursor.execute("""
                SELECT ea.company_name as name, 
                       ea.cin as cin,
                       CASE WHEN ea.cin LIKE 'L%' THEN 'Public' ELSE 'Private' END as type, 
                       COUNT(ea.din) AS director_count,
                       EXISTS (
                           SELECT 1 FROM directors_data.companies c WHERE c.cin = ea.cin
                       ) as is_group
                FROM directors_master.external_board_members ea
                GROUP BY ea.cin, ea.company_name
                ORDER BY is_group DESC, director_count DESC, ea.company_name ASC
            """)
            return [dict(r) for r in cursor.fetchall()]
        finally:
            pg_conn.close()
    return []
