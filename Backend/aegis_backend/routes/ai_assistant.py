# AI Assistant Route Module
# This module handles AI assistant functionality for meeting transcript processing and Meeting Minutes generation
from fastapi import APIRouter, HTTPException, UploadFile, File
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
import os
import json
import logging
import asyncio
import concurrent.futures
import uuid
from datetime import datetime
import subprocess
import platform
from docx import Document as DocxDocument
import shlex
import re

logger = logging.getLogger(__name__)

# Custom thread pool for handling blocking operations
thread_pool = concurrent.futures.ThreadPoolExecutor(max_workers=4)

# Create a router instance for AI assistant endpoints
router = APIRouter()

# Response model for transcript upload
class TranscriptUploadResponse(BaseModel):
    task_id: str
    message: str

# Response model for Meeting Minutes generation
class MoMGenerationResponse(BaseModel):
    task_id: str
    status: str
    message: str

# Model for Meeting Minutes content structure
class MoMContent(BaseModel):
    title: str
    date: str
    attendees: List[Dict[str, str]]
    agenda: List[str]
    decisions: List[str]
    action_items: List[Dict[str, str]]
    next_meeting: str

# Enhanced configuration for LLM processing
LLM_CONFIG = {
    "model": "llama-3.3-70b-versatile",
    "max_tokens": 2000,
    "temperature": 0.3,
    "max_input_chars": 15000,  # Maximum characters before chunking
    "chunk_size": 10000,       # Size of chunks for processing
    "summary_chunk_size": 8000 # Size of chunks for summarization
}

# Enhanced helper functions for Teams transcript processing
def detect_transcript_format(text_content, file_path):
    """
    Detect the format of the transcript based on structural patterns
    Returns: 'structured' or 'teams'
    """
    # Check for Teams transcript patterns
    lines = text_content.split('\n')
    
    # Teams transcripts have speaker lines with timestamp format: "Speaker Name   0:22"
    teams_pattern_count = sum(1 for line in lines if re.match(r'^[^:]+?\s+\d+:\d+', line.strip()))
    
    # Structured transcripts have explicit metadata sections
    structured_indicators = ['[Meeting Title]:', '[Date]:', '[Attendees]:']
    structured_pattern_count = sum(1 for indicator in structured_indicators if indicator in text_content)
    
    # Also check filename for Teams patterns
    filename = os.path.basename(file_path)
    teams_filename_indicators = ['-2025', '-Meeting Recording', 'Recording']
    teams_filename_matches = sum(1 for indicator in teams_filename_indicators if indicator in filename)
    
    if teams_pattern_count > 5 or teams_filename_matches > 0:
        return 'teams'
    elif structured_pattern_count > 2:
        return 'structured'
    else:
        return 'unknown'

def extract_metadata_from_filename(filename):
    """
    Extract metadata from Teams transcript filename
    Extract meeting title and date from uploaded document filename. 
    For example, filename '@AGE23L - Board Meeting - 3.00 p.m. IST, Monday, 27th October, 2025.docx' 
    should yield title 'AGE23L - Board Meeting' and date '27th October, 2025'.
    """
    metadata = {
        'title': 'Untitled Meeting',
        'date': 'Unknown Date',
        'speakers': []
    }
    
    # Clean filename by removing extension and recording identifiers
    clean_filename = filename
    if '.' in clean_filename:
        clean_filename = '.'.join(clean_filename.split('.')[:-1])
    
    # Remove recording identifiers
    recording_indicators = ['-2025', '-Meeting Recording', 'Recording']
    for indicator in recording_indicators:
        if indicator in clean_filename:
            parts = clean_filename.split(indicator)
            clean_filename = parts[0] if parts else clean_filename
    
    # Extract title and date from filename pattern
    # Pattern: "TITLE - TIME DETAILS, DAY, DATE"
    if ' - ' in clean_filename:
        parts = clean_filename.split(' - ')
        if len(parts) >= 2:
            # First part is the title
            metadata['title'] = parts[0].strip()
            
            # Look for date in all parts, not just the last one
            for part in parts:
                # Extract date pattern like "27th October, 2025"
                date_match = re.search(r'(\d{1,2}[a-z]{2}\s+[A-Za-z]+\s*,\s*\d{4})', part)
                if date_match:
                    metadata['date'] = date_match.group(1)
                    break
                # Try alternative date patterns
                alt_date_match = re.search(r'(\d{1,2}[^\d]+\d{1,2}[^\d]+\d{4})', part)
                if alt_date_match:
                    metadata['date'] = alt_date_match.group(1)
                    break
    
    return metadata

def is_system_speaker(speaker):
    """
    Check if a speaker is a system/location identifier rather than a real participant
    """
    system_patterns = [
        r'CT\d+F.*',  # Location identifiers like "CT2F - Boardroom"
        r'CT\s+\d+F.*',  # Location identifiers with spaces
        r'.*Boardroom.*',  # Generic boardroom labels
        r'.*Room.*',  # Generic room labels
        r'.*started transcription.*',  # System messages
        r'.*joined.*',  # Join notifications
        r'.*left.*',  # Leave notifications
        r'.*recording.*',  # Recording notifications
    ]
    
    for pattern in system_patterns:
        if re.match(pattern, speaker, re.IGNORECASE):
            return True
    return False

def extract_speakers_from_content(lines):
    """
    Extract real speakers from transcript content, filtering out system identifiers
    """
    speakers = set()
    
    for line in lines:
        # Extract speakers from lines with timestamp format: "Speaker Name   0:22"
        speaker_match = re.match(r'^([^:]+?)\s+(\d+:\d+)', line.strip())
        if speaker_match:
            speaker = speaker_match.group(1).strip()
            # Filter out system speakers
            if not is_system_speaker(speaker):
                speakers.add(speaker)
    
    return list(speakers)

def preprocess_teams_transcript(text_content, file_path):
    """
    Preprocess Teams transcript to extract metadata and clean content
    """
    lines = text_content.split('\n')
    filename = os.path.basename(file_path)
    
    # Extract metadata from filename
    metadata = extract_metadata_from_filename(filename)
    
    # Extract real speakers from content
    real_speakers = extract_speakers_from_content(lines)
    metadata['speakers'] = real_speakers
    
    # Clean content by removing transcription artifacts and system messages
    cleaned_lines = []
    
    for i, line in enumerate(lines):
        # Skip transcription start indicator
        if 'started transcription' in line.lower():
            continue
            
        # Skip redundant timestamp lines
        if re.match(r'^\d+:\d+:\d+\s*-->\s*\d+:\d+:\d+$', line.strip()):
            continue
            
        # Skip system speaker lines
        speaker_match = re.match(r'^([^:]+?)\s+(\d+:\d+)', line.strip())
        if speaker_match:
            speaker = speaker_match.group(1).strip()
            if is_system_speaker(speaker):
                continue
            # Keep the line but clean it
            cleaned_lines.append(f"{speaker}: ")
        else:
            cleaned_lines.append(line)
    
    # Reconstruct content
    cleaned_content = '\n'.join(cleaned_lines)
    
    # Add inferred metadata to content
    header = f"[Meeting Title]: {metadata['title']}\n"
    header += f"[Date]: {metadata['date']}\n"
    header += "[Attendees]:\n"
    for speaker in metadata['speakers']:
        header += f"- {speaker}\n"
    header += "\n" + "="*60 + "\n"
    
    return header + cleaned_content, metadata

def create_fallback_mom_data(transcript_format, metadata=None):
    """
    Create fallback Meeting Minutes structure with format-specific notes
    """
    if transcript_format == 'teams':
        notes = "Note: This Meeting Minutes was generated from a Microsoft Teams transcript. Some information may have been inferred due to the unstructured nature of the source format."
        title = metadata.get('title', 'Untitled Meeting') if metadata else 'Untitled Meeting'
        date = metadata.get('date', 'Unknown Date') if metadata else 'Unknown Date'
        attendees = [{"name": speaker, "role": "Attendee"} for speaker in metadata.get('speakers', [])] if metadata else []
    else:
        notes = "Note: Some information could not be extracted from the transcript."
        title = "Meeting Minutes"
        date = "N/A"
        attendees = []
        
    return {
        "title": title,
        "date": date,
        "attendees": attendees,
        "agenda": [notes],
        "decisions": [],
        "action_items": [],
        "next_meeting": "Not specified"
    }

def chunk_text(text, chunk_size):
    """
    Split text into chunks of specified size, trying to break at natural points
    """
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = min(start + chunk_size, len(text))
        
        # Try to find a natural break point (double newline or single newline near the end)
        if end < len(text):
            # Look for double newline first
            double_newline = text.rfind('\n\n', start, end)
            if double_newline > start + chunk_size * 0.8:  # If found near the end
                end = double_newline + 2
            else:
                # Look for single newline
                single_newline = text.rfind('\n', start, end)
                if single_newline > start + chunk_size * 0.9:  # If found near the end
                    end = single_newline + 1
        
        chunks.append(text[start:end])
        start = end
    
    return chunks

def extract_key_points_from_chunk(client, chunk, chunk_index, total_chunks):
    """
    Extract key points from a chunk of transcript
    """
    try:
        system_prompt = """You are an AI assistant that extracts key points from meeting transcript chunks. 
Focus on formal agenda items, decisions, action items, and important discussion points. 
Distinguish between formal agenda items (usually numbered) and discussion topics.
Only extract information that is explicitly present in the transcript.
Ignore system messages and location labels as speakers."""

        user_prompt = f"""Extract ONLY the key points from this chunk of meeting transcript (chunk {chunk_index+1} of {total_chunks}). 
Focus ONLY on what is explicitly stated in the transcript:

1. Formal agenda items (look for numbered items like "Item 1", "Item 2", etc.)
2. Decisions made (look for "approved", "noted", "committee decided", "accorded")
3. Action items assigned (only those with clear ownership/assignment)
4. Important discussion points within formal agenda items

DO NOT generate or infer information that is not explicitly present in the transcript.
DO NOT treat every topic mentioned as a separate agenda item.
DO NOT include system messages or location labels as speakers or agenda items.

Chunk content:
{chunk}

Respond with a concise summary of ONLY the key points that are explicitly mentioned."""

        response = client.chat.completions.create(
            model=LLM_CONFIG["model"],
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=LLM_CONFIG["temperature"],
            max_tokens=500,
            top_p=1,
            stream=False
        )
        
        return response.choices[0].message.content
    except Exception as e:
        logger.error(f"Error extracting key points from chunk {chunk_index}: {str(e)}")
        return f"Error processing chunk {chunk_index}: {str(e)}"

def synthesize_mom_from_key_points(client, key_points_list, metadata):
    """
    Synthesize final Meeting Minutes from key points extracted from chunks
    """
    try:
        system_prompt = """You are an AI assistant that synthesizes Meeting Minutes from key points extracted from transcript chunks. 
Create well-structured Meeting Minutes with ONLY information that is explicitly present in the key points.
Do not add information not present in the key points.
Do not generate or infer information.
Distinguish between formal agenda items (usually numbered) and discussion topics.
For action items, only include those with clear ownership/assignment.
Look for explicit agenda structure markers (Item 1, Item 2, etc.).
Limit agenda items to top-level formal items (max 15-20)."""

        key_points_combined = "\n\n".join([f"Chunk {i+1} Key Points:\n{points}" for i, points in enumerate(key_points_list)])
        
        user_prompt = f"""Synthesize Meeting Minutes from the following key points extracted from transcript chunks:

Meeting Metadata:
Title: {metadata.get('title', 'Unknown')}
Date: {metadata.get('date', 'Unknown')}
Attendees: {', '.join(metadata.get('speakers', [])) if metadata.get('speakers') else 'Unknown'}

Key Points:
{key_points_combined}

Respond ONLY with valid JSON in the following format and include ONLY information explicitly present in the key points: 
{{"title": "Meeting Title", "date": "Meeting Date", "attendees":[{{"name": "Attendee Name", "role": "Attendee Role"}}], "agenda":["Formal Agenda Item 1", "Formal Agenda Item 2"], "decisions":["Decision 1", "Decision 2"], "action_items":[{{"task": "Task Description", "assignee": "Assignee Name"}}], "next_meeting": "Next meeting details"}}

IMPORTANT GUIDELINES:
1. Distinguish between formal agenda items (usually numbered) and discussion topics
2. Extract meeting title from document header or filename if not explicitly in transcript
3. For action items, only include those with clear ownership/assignment
4. Limit agenda items to top-level formal items (max 15-20)
5. Only include information explicitly present in the key points
6. If title is generic ("Untitled Meeting"), use filename-derived title
7. If date is "Unknown", use filename-derived date
8. If attendees are missing, use filename-derived speakers"""

        response = client.chat.completions.create(
            model=LLM_CONFIG["model"],
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=LLM_CONFIG["temperature"],
            max_tokens=LLM_CONFIG["max_tokens"],
            top_p=1,
            stream=False
        )
        
        return response.choices[0].message.content
    except Exception as e:
        logger.error(f"Error synthesizing MoM from key points: {str(e)}")
        raise Exception(f"Error synthesizing Meeting Minutes: {str(e)}")

def process_large_transcript(client, content, metadata):
    """
    Process large transcripts by chunking and synthesizing
    """
    try:
        logger.info(f"Processing large transcript of {len(content)} characters")
        
        # Chunk the content
        chunks = chunk_text(content, LLM_CONFIG["chunk_size"])
        logger.info(f"Split transcript into {len(chunks)} chunks")
        
        # Extract key points from each chunk
        key_points_list = []
        for i, chunk in enumerate(chunks):
            logger.info(f"Processing chunk {i+1}/{len(chunks)}")
            key_points = extract_key_points_from_chunk(client, chunk, i, len(chunks))
            key_points_list.append(key_points)
        
        # Synthesize final Meeting Minutes
        logger.info("Synthesizing final Meeting Minutes")
        final_mom_content = synthesize_mom_from_key_points(client, key_points_list, metadata)
        
        return final_mom_content
    except Exception as e:
        logger.error(f"Error processing large transcript: {str(e)}")
        raise Exception(f"Error processing large transcript: {str(e)}")

# Endpoint to upload a transcript file (DOCX or TXT) for processing
@router.post("/ai-assistant/upload", response_model=TranscriptUploadResponse)
async def upload_transcript(file: UploadFile = File(...)):
    """Upload a transcript file (DOCX or TXT) for processing"""
    try:
        # Generate a unique task ID
        task_id = str(uuid.uuid4())
        
        # Create directory for this task if it doesn't exist
        task_dir = os.path.join(os.path.dirname(__file__), "..", "public", "ai_assistant_mom", task_id)
        os.makedirs(task_dir, exist_ok=True)
        
        # Save the uploaded file
        file_extension = os.path.splitext(file.filename)[1].lower()
        if file_extension not in ['.docx', '.txt', '.pdf']:
            raise HTTPException(status_code=400, detail="Only DOCX, TXT, and PDF files are supported")
        
        file_path = os.path.join(task_dir, f"transcript{file_extension}")
        
        # Save file content
        with open(file_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        # Create initial status file
        status_file = os.path.join(task_dir, "status.json")
        status_data = {
            "task_id": task_id,
            "status": "uploaded",
            "message": "File uploaded successfully",
            "created_at": datetime.now().isoformat()
        }
        
        with open(status_file, "w") as f:
            json.dump(status_data, f)
        
        return TranscriptUploadResponse(
            task_id=task_id,
            message="File uploaded successfully"
        )
    
    except Exception as e:
        logger.error(f"Error uploading transcript: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to upload transcript: {str(e)}")

# Endpoint to generate Meeting Minutes from uploaded transcript using LLM
@router.post("/ai-assistant/generate-mom", response_model=MoMGenerationResponse)
async def generate_mom(task_id: str):
    """Generate Meeting Minutes from uploaded transcript using LLM"""
    try:
        # Check if task directory exists
        task_dir = os.path.join(os.path.dirname(__file__), "..", "public", "ai_assistant_mom", task_id)
        if not os.path.exists(task_dir):
            raise HTTPException(status_code=404, detail="Task not found")
        
        # Check if transcript exists
        transcript_path = None
        for ext in ['.docx', '.txt', '.pdf']:
            path = os.path.join(task_dir, f"transcript{ext}")
            if os.path.exists(path):
                transcript_path = path
                break
        
        if not transcript_path:
            raise HTTPException(status_code=404, detail="Transcript file not found")
        
        # Update status to processing
        status_file = os.path.join(task_dir, "status.json")
        with open(status_file, "r") as f:
            status_data = json.load(f)
        
        status_data["status"] = "processing"
        status_data["message"] = "Generating Meeting Minutes"
        status_data["started_at"] = datetime.now().isoformat()
        
        with open(status_file, "w") as f:
            json.dump(status_data, f)
        
        def process_transcript():
            # Extract text from transcript
            if transcript_path.endswith('.docx'):
                doc = DocxDocument(transcript_path)
                text_content = "\n".join([para.text for para in doc.paragraphs])
            elif transcript_path.endswith('.pdf'):
                try:
                    import PyPDF2
                    with open(transcript_path, "rb") as fh:
                        reader = PyPDF2.PdfReader(fh)
                        text_content = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
                except ImportError:
                    import pdfplumber
                    with pdfplumber.open(transcript_path) as pdf:
                        text_content = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
            else:  # .txt file
                with open(transcript_path, "r", encoding="utf-8") as f:
                    text_content = f.read()
            
            # Detect transcript format
            transcript_format = detect_transcript_format(text_content, transcript_path)
            
            # Process based on format
            if transcript_format == 'teams':
                # Preprocess Teams transcript
                processed_content, metadata = preprocess_teams_transcript(text_content, transcript_path)
                format_instructions = """
IMPORTANT: This is a Microsoft Teams meeting transcript with the following characteristics:
1. Speaker names are followed by timestamps in format "Speaker Name   0:22"
2. Meeting metadata has been inferred from filename and content
3. Some speaker attributions may be incomplete or missing
4. Agenda items are implicit and must be inferred from content flow
5. System speakers like "CT2F - Boardroom" should be filtered out

Please extract information carefully, noting when metadata was inferred rather than explicitly stated. 
ONLY extract information that is explicitly present in the transcript.
DO NOT generate or infer information that is not present in the transcript.
Distinguish between formal agenda items (usually numbered) and discussion topics.
For action items, only include those with clear ownership/assignment.
"""
            else:
                processed_content = text_content
                metadata = {"title": "Unknown", "date": "Unknown", "speakers": []}
                format_instructions = """
This is a structured meeting transcript with explicit metadata sections.
Extract information as usual from the clearly marked sections.
ONLY extract information that is explicitly present in the transcript.
DO NOT generate or infer information that is not present in the transcript.
"""

            # Determine which script to use based on OS
            if platform.system() == "Windows":
                script_path = os.path.join(os.path.dirname(__file__), "..", "..", "llm_client.bat")
            else:
                script_path = os.path.join(os.path.dirname(__file__), "..", "..", "llm_client.sh")
            
            # Run LLM client script or use Python Groq library
            try:
                # Check if we should use Groq with Python library
                # Force Groq to be true by default, only use Azure if explicitly set to false
                use_groq_env = os.environ.get('USE_GROQ', 'false').lower()
                use_groq = use_groq_env != 'false'
                
                logger.info(f"USE_GROQ environment variable: {use_groq_env}")
                logger.info(f"Using Groq: {use_groq}")
                
                if use_groq:
                    # Use Python Groq library
                    from groq import Groq
                    
                    # Initialize Groq client with explicit API key from environment
                    groq_api_key = os.getenv("GROQ_API_KEY")
                    client = Groq(api_key=groq_api_key) if groq_api_key else Groq()
                    
                    # Check if content is too large and handle appropriately
                    if len(processed_content) > LLM_CONFIG["max_input_chars"]:
                        logger.info(f"Large transcript detected ({len(processed_content)} chars), using chunking approach")
                        # Process large transcript using chunking approach
                        final_content = process_large_transcript(client, processed_content, metadata)
                        
                        # Try to parse the final content as JSON
                        try:
                            mom_data = json.loads(final_content)
                        except json.JSONDecodeError:
                            # If not JSON, try to extract JSON from markdown code blocks
                            json_match = re.search(r"```(?:json)?\s*({.*?})\s*```", final_content, re.DOTALL)
                            if json_match:
                                mom_data = json.loads(json_match.group(1))
                            else:
                                # Create fallback structure with metadata
                                mom_data = create_fallback_mom_data(transcript_format, metadata)
                    else:
                        # Process normally for smaller content
                        logger.info(f"Normal transcript size ({len(processed_content)} chars), using standard approach")
                        
                        # Create prompt for LLM
                        prompt_data = {
                            "messages": [
                                {
                                    "role": "system",
                                    "content": f"You are an AI assistant that converts meeting transcripts into structured Meeting Minutes. Extract key information including title, date, attendees, agenda items, decisions made, action items with assignees, and next meeting details. {format_instructions} Respond ONLY with valid JSON in the following format: {{\"title\": \"Meeting Title\", \"date\": \"Meeting Date\", \"attendees\":[{{\"name\": \"Attendee Name\", \"role\": \"Attendee Role\"}}], \"agenda\":[\"Formal Agenda Item 1\", \"Formal Agenda Item 2\"], \"decisions\":[\"Decision 1\", \"Decision 2\"], \"action_items\":[{{\"task\": \"Task Description\", \"assignee\": \"Assignee Name\"}}], \"next_meeting\": \"Next meeting details\"}}. ONLY include information that is explicitly present in the transcript. Distinguish between formal agenda items (usually numbered) and discussion topics. For action items, only include those with clear ownership/assignment."
                                },
                                {
                                    "role": "user",
                                    "content": f"Please convert the following meeting transcript into structured Meeting Minutes in JSON format:\n\n{processed_content}\n\nIMPORTANT: Respond ONLY with valid JSON. Do not include any other text, markdown, or explanations. ONLY include information that is explicitly present in the transcript. Distinguish between formal agenda items (usually numbered) and discussion topics. For action items, only include those with clear ownership/assignment."
                                }
                            ],
                            "temperature": LLM_CONFIG["temperature"],
                            "max_tokens": LLM_CONFIG["max_tokens"]
                        }
                        
                        # Make API call to Groq
                        completion = client.chat.completions.create(
                            model=LLM_CONFIG["model"],
                            messages=prompt_data["messages"],
                            temperature=prompt_data["temperature"],
                            max_tokens=prompt_data["max_tokens"],
                            top_p=1,
                            stream=False
                        )
                        
                        # Extract content from response
                        content = completion.choices[0].message.content
                        
                        # Try to parse as JSON
                        try:
                            mom_data = json.loads(content)
                        except json.JSONDecodeError:
                            # If not JSON, try to extract JSON from markdown code blocks
                            json_match = re.search(r"```(?:json)?\s*({.*?})\s*```", content, re.DOTALL)
                            if json_match:
                                mom_data = json.loads(json_match.group(1))
                            else:
                                # Create fallback structure with metadata
                                mom_data = create_fallback_mom_data(transcript_format, metadata)
                else:
                    # Use shell script for Azure OpenAI (simplified approach for now)
                    logger.info("Using Azure OpenAI API via shell script")
                    
                    # For Azure, we'll use a simplified approach - truncate if too large
                    if len(processed_content) > LLM_CONFIG["max_input_chars"]:
                        processed_content = processed_content[:LLM_CONFIG["max_input_chars"]]
                        logger.info("Truncated content for Azure OpenAI API")
                    
                    # Create prompt for LLM
                    prompt_data = {
                        "messages": [
                            {
                                "role": "system",
                                "content": f"You are an AI assistant that converts meeting transcripts into structured Meeting Minutes. Extract key information including title, date, attendees, agenda items, decisions made, action items with assignees, and next meeting details. {format_instructions} Respond ONLY with valid JSON in the following format: {{\"title\": \"Meeting Title\", \"date\": \"Meeting Date\", \"attendees\":[{{\"name\": \"Attendee Name\", \"role\": \"Attendee Role\"}}], \"agenda\":[\"Formal Agenda Item 1\", \"Formal Agenda Item 2\"], \"decisions\":[\"Decision 1\", \"Decision 2\"], \"action_items\":[{{\"task\": \"Task Description\", \"assignee\": \"Assignee Name\"}}], \"next_meeting\": \"Next meeting details\"}}. ONLY include information that is explicitly present in the transcript. Distinguish between formal agenda items (usually numbered) and discussion topics. For action items, only include those with clear ownership/assignment."
                            },
                            {
                                "role": "user",
                                "content": f"Please convert the following meeting transcript into structured Meeting Minutes in JSON format:\n\n{processed_content}\n\nIMPORTANT: Respond ONLY with valid JSON. Do not include any other text, markdown, or explanations. ONLY include information that is explicitly present in the transcript. Distinguish between formal agenda items (usually numbered) and discussion topics. For action items, only include those with clear ownership/assignment."
                            }
                        ],
                        "temperature": LLM_CONFIG["temperature"],
                        "max_tokens": LLM_CONFIG["max_tokens"]
                    }
                    
                    # Save prompt to file
                    prompt_file = os.path.join(task_dir, "prompt.json")
                    with open(prompt_file, "w") as f:
                        json.dump(prompt_data, f, indent=2)
                    
                    # Properly handle file paths with spaces on Windows
                    if platform.system() == "Windows":
                        # For Windows batch files, we need to properly quote the file path
                        prompt_file_quoted = f'"{prompt_file}"' if ' ' in prompt_file else prompt_file
                        result = subprocess.run([script_path, prompt_file_quoted], capture_output=True, text=True, cwd=os.path.dirname(script_path), shell=True)
                    else:
                        # For Unix-like systems, shlex.quote handles spaces properly
                        result = subprocess.run([script_path, shlex.quote(prompt_file)], capture_output=True, text=True, cwd=os.path.dirname(script_path))
                    
                    if result.returncode != 0:
                        raise Exception(f"LLM client script failed: {result.stderr}")
                    
                    # Parse LLM response
                    llm_response = json.loads(result.stdout)
                    
                    # Extract content from response
                    if "choices" in llm_response and len(llm_response["choices"]) > 0:
                        content = llm_response["choices"][0]["message"]["content"]
                        
                        # Try to parse as JSON
                        try:
                            mom_data = json.loads(content)
                        except json.JSONDecodeError:
                            # If not JSON, try to extract JSON from markdown code blocks
                            json_match = re.search(r"```(?:json)?\s*({.*?})\s*```", content, re.DOTALL)
                            if json_match:
                                mom_data = json.loads(json_match.group(1))
                            else:
                                # Create fallback structure with metadata
                                mom_data = create_fallback_mom_data(transcript_format, metadata)
                    else:
                        # Create fallback structure with metadata
                        mom_data = create_fallback_mom_data(transcript_format, metadata)
                
                # Ensure metadata from filename is preserved if not extracted properly
                if transcript_format == 'teams' and metadata:
                    if mom_data.get('title') in ['Meeting Minutes', 'Untitled Meeting', 'Unknown'] or not mom_data.get('title'):
                        mom_data['title'] = metadata.get('title', 'Untitled Meeting')
                    if mom_data.get('date') in ['N/A', 'Unknown Date', 'Unknown'] or not mom_data.get('date'):
                        mom_data['date'] = metadata.get('date', 'Unknown Date')
                    if not mom_data.get('attendees') and metadata.get('speakers'):
                        mom_data['attendees'] = [{"name": speaker, "role": "Attendee"} for speaker in metadata.get('speakers', [])]
                
                # Save structured MoM data
                mom_json_path = os.path.join(task_dir, "mom.json")
                with open(mom_json_path, "w") as f:
                    json.dump(mom_data, f, indent=2)
                
                # Create DOCX document from MoM data
                from docx import Document
                doc = Document()
                
                # Add title
                doc.add_heading(mom_data.get('title', 'Meeting Minutes'), 0)
                
                # Add metadata
                doc.add_paragraph(f"Date: {mom_data.get('date', 'N/A')}")
                
                # Add attendees
                doc.add_heading('Attendees', level=1)
                if 'attendees' in mom_data and mom_data['attendees']:
                    for attendee in mom_data['attendees']:
                        doc.add_paragraph(f"{attendee.get('name', 'N/A')} ({attendee.get('role', 'N/A')})")
                else:
                    doc.add_paragraph('Not specified')
                
                # Add agenda
                doc.add_heading('Agenda', level=1)
                if 'agenda' in mom_data and mom_data['agenda']:
                    # Limit to reasonable number of agenda items
                    agenda_items = mom_data['agenda'][:20]  # Cap at 20 items
                    for item in agenda_items:
                        doc.add_paragraph(item, style='List Bullet')
                else:
                    doc.add_paragraph('Not specified')
                
                # Add decisions
                doc.add_heading('Decisions Made', level=1)
                if 'decisions' in mom_data and mom_data['decisions']:
                    for decision in mom_data['decisions']:
                        doc.add_paragraph(decision, style='List Bullet')
                else:
                    doc.add_paragraph('Not specified')
                
                # Add action items
                doc.add_heading('Action Items', level=1)
                if 'action_items' in mom_data and mom_data['action_items']:
                    for item in mom_data['action_items']:
                        assignee = item.get('assignee', 'Unassigned')
                        task = item.get('task', 'N/A')
                        if assignee and assignee != 'Unassigned':
                            doc.add_paragraph(f"{task} - {assignee}", style='List Bullet')
                        else:
                            # Only include action items with clear assignees
                            pass
                else:
                    doc.add_paragraph('Not specified')
                
                # Add next meeting
                doc.add_heading('Next Meeting', level=1)
                doc.add_paragraph(mom_data.get('next_meeting', 'Not specified'))
                
                # Save DOCX file
                mom_docx_path = os.path.join(task_dir, "mom.docx")
                doc.save(mom_docx_path)
                
                return True
                
            except Exception as e:
                logger.error(f"Error in process_transcript: {str(e)}")
                raise Exception(f"Error processing with LLM: {str(e)}")
        
        # Run processing in thread pool
        loop = asyncio.get_event_loop()
        success = await loop.run_in_executor(thread_pool, process_transcript)
        
        # Update status
        with open(status_file, "r") as f:
            status_data = json.load(f)
        
        if success:
            status_data["status"] = "completed"
            status_data["message"] = "Meeting Minutes generated successfully"
            status_data["completed_at"] = datetime.now().isoformat()
        else:
            status_data["status"] = "failed"
            status_data["message"] = "Failed to generate Meeting Minutes"
        
        with open(status_file, "w") as f:
            json.dump(status_data, f)
        
        return MoMGenerationResponse(
            task_id=task_id,
            status=status_data["status"],
            message=status_data["message"]
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generating MoM: {str(e)}")
        
        # Update status to failed
        try:
            task_dir = os.path.join(os.path.dirname(__file__), "..", "public", "ai_assistant_mom", task_id)
            status_file = os.path.join(task_dir, "status.json")
            if os.path.exists(status_file):
                with open(status_file, "r") as f:
                    status_data = json.load(f)
                status_data["status"] = "failed"
                status_data["message"] = f"Failed to generate Meeting Minutes: {str(e)}"
                status_data["completed_at"] = datetime.now().isoformat()
                with open(status_file, "w") as f:
                    json.dump(status_data, f)
        except:
            pass
        
        raise HTTPException(status_code=500, detail=f"Failed to generate Meeting Minutes: {str(e)}")

# Endpoint to get the status of Meeting Minutes generation task
@router.get("/ai-assistant/status/{task_id}")
async def get_mom_status(task_id: str):
    """Get the status of Meeting Minutes generation task"""
    try:
        task_dir = os.path.join(os.path.dirname(__file__), "..", "public", "ai_assistant_mom", task_id)
        if not os.path.exists(task_dir):
            raise HTTPException(status_code=404, detail="Task not found")
        
        status_file = os.path.join(task_dir, "status.json")
        if not os.path.exists(status_file):
            raise HTTPException(status_code=404, detail="Status file not found")
        
        with open(status_file, "r") as f:
            status_data = json.load(f)
        
        return status_data
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting MoM status: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to get status: {str(e)}")

# Endpoint to get the structured Meeting Minutes content
@router.get("/ai-assistant/mom/{task_id}")
async def get_mom_content(task_id: str):
    """Get the structured Meeting Minutes content"""
    try:
        task_dir = os.path.join(os.path.dirname(__file__), "..", "public", "ai_assistant_mom", task_id)
        if not os.path.exists(task_dir):
            raise HTTPException(status_code=404, detail="Task not found")
        
        mom_json_path = os.path.join(task_dir, "mom.json")
        if not os.path.exists(mom_json_path):
            raise HTTPException(status_code=404, detail="Meeting Minutes not found")
        
        with open(mom_json_path, "r") as f:
            mom_data = json.load(f)
        
        return mom_data
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting MoM content: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to get Meeting Minutes content: {str(e)}")

# Endpoint to download the generated Meeting Minutes DOCX file
@router.get("/ai-assistant/download/{task_id}")
async def download_mom(task_id: str):
    """Download the generated Meeting Minutes DOCX file"""
    try:
        task_dir = os.path.join(os.path.dirname(__file__), "..", "public", "ai_assistant_mom", task_id)
        if not os.path.exists(task_dir):
            raise HTTPException(status_code=404, detail="Task not found")
        
        mom_docx_path = os.path.join(task_dir, "mom.docx")
        if not os.path.exists(mom_docx_path):
            raise HTTPException(status_code=404, detail="Meeting Minutes DOCX file not found")
        
        return FileResponse(
            path=mom_docx_path,
            filename=f"Meeting_Minutes_{task_id}.docx",
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error downloading MoM: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to download Meeting Minutes: {str(e)}")