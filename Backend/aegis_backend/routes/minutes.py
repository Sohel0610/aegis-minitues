# Minutes Generation Route Module
# This module handles minutes preparation functionality including place management and document generation using PostgreSQL
from fastapi import APIRouter, HTTPException, UploadFile, File, Form, Request
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
import os
import re
import io
import json
import sqlite3
import logging
import asyncio
import concurrent.futures
from datetime import datetime
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
try:
    import pdfplumber
    PDF_PLUMBER_AVAILABLE = True
except ImportError:
    PDF_PLUMBER_AVAILABLE = False
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
try:
    import pytesseract
    from pdf2image import convert_from_bytes
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
import shutil
from utils.pgsql_service import get_pg_connection, get_pg_cursor
from utils.auth_dep import require_session
from utils.audit_logger import AuditLogger, get_client_ip, get_user_agent
from fastapi import Depends


# --- Document Content Extraction Helpers ---

def extract_text_from_docx(file_bytes: bytes) -> Dict[str, Any]:
    """Extract text paragraphs and tables from a .docx file."""
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables = []
    for tbl_idx, table in enumerate(doc.tables):
        rows_data = []
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells]
            rows_data.append(row_cells)
        if rows_data:
            headers = rows_data[0] if rows_data else []
            tables.append({
                "table_index": tbl_idx,
                "headers": headers,
                "rows": rows_data[1:] if len(rows_data) > 1 else [],
                "total_rows": len(rows_data) - 1
            })
    return {
        "text": "\n".join(paragraphs),
        "paragraph_count": len(paragraphs),
        "tables": tables,
        "table_count": len(tables)
    }


def render_resolutions_into_doc(doc, resolutions_text: str, anchor_para=None):
    """Insert resolution content preserving structure (MOM #8): plain lines become
    paragraphs; pipe/tab-delimited blocks become real DOCX tables. If anchor_para
    is given, content is relocated directly after it."""
    def is_table_line(ln: str) -> bool:
        return ('|' in ln and ln.count('|') >= 2) or ('\t' in ln)

    lines = resolutions_text.replace('\r\n', '\n').split('\n')
    created_elements = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if is_table_line(line):
            block = []
            while i < len(lines) and is_table_line(lines[i]):
                raw = lines[i].strip().strip('|')
                cells = [c.strip() for c in (raw.split('|') if '|' in lines[i] else raw.split('\t'))]
                # skip markdown separator rows such as ---|---
                if not all(set(c) <= set('-: ') for c in cells):
                    block.append(cells)
                i += 1
            if block:
                cols = max(len(r) for r in block)
                table = doc.add_table(rows=len(block), cols=cols)
                try:
                    table.style = 'Table Grid'
                except Exception:
                    pass
                for r_idx, row_cells in enumerate(block):
                    for c_idx in range(cols):
                        table.rows[r_idx].cells[c_idx].text = row_cells[c_idx] if c_idx < len(row_cells) else ''
                created_elements.append(table._tbl)
        else:
            if line.strip():
                p = doc.add_paragraph(line)
                created_elements.append(p._p)
            i += 1

    if anchor_para is not None:
        for el in reversed(created_elements):
            anchor_para._p.addnext(el)


def extract_text_from_pdf(file_bytes: bytes) -> Dict[str, Any]:
    """Extract text and tables from a .pdf file using pdfplumber."""
    all_text = []
    tables = []
    tbl_idx = 0

    if PDF_PLUMBER_AVAILABLE:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page_num, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    all_text.append(page_text)
                page_tables = page.extract_tables()
                if page_tables:
                    for pt in page_tables:
                        if pt and len(pt) > 0:
                            headers = [str(c or '') for c in pt[0]]
                            rows = [[str(c or '') for c in row] for row in pt[1:]]
                            tables.append({
                                "table_index": tbl_idx,
                                "page": page_num + 1,
                                "headers": headers,
                                "rows": rows,
                                "total_rows": len(rows)
                            })
                            tbl_idx += 1
    elif PYPDF2_AVAILABLE:
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            text = page.extract_text()
            if text:
                all_text.append(text)

    # OCR fallback for scanned/image-only PDFs (BRD #4 / #17)
    ocr_used = False
    if len("".join(all_text).strip()) < 50 and OCR_AVAILABLE:
        try:
            images = convert_from_bytes(file_bytes, dpi=200)
            ocr_text = []
            for img in images:
                page_text = pytesseract.image_to_string(img)
                if page_text and page_text.strip():
                    ocr_text.append(page_text.strip())
            if ocr_text:
                all_text = ocr_text
                ocr_used = True
        except Exception as ocr_err:
            logging.getLogger(__name__).warning(f"OCR fallback failed: {ocr_err}")

    return {
        "text": "\n".join(all_text),
        "paragraph_count": len(all_text),
        "tables": tables,
        "table_count": len(tables),
        "ocr_used": ocr_used
    }

logger = logging.getLogger(__name__)

# Custom thread pool for handling blocking operations
thread_pool = concurrent.futures.ThreadPoolExecutor(max_workers=4)

# Create a router instance for minutes endpoints
router = APIRouter()

# --- Models ---

class GeneratedMinuteResponse(BaseModel):
    id: int
    company_name: str
    meeting_type: str
    meeting_date: str
    file_path: str
    created_at: str
    download_url: Optional[str] = None
    meeting_number: Optional[str] = None
    meeting_year: Optional[str] = None
    status: Optional[str] = "draft"  # draft | finalized
    finalized_at: Optional[str] = None
    finalized_by: Optional[str] = None
    is_signed: Optional[bool] = False
    unsigned_file_path: Optional[str] = None

class MinutesHistoryResponse(BaseModel):
    data: List[GeneratedMinuteResponse]
    count: int

class CompanyMeetingsListResponse(BaseModel):
    data: List[GeneratedMinuteResponse]
    count: int
    meeting_type: Optional[str] = None
    next_meeting_number: Optional[str] = None


def _row_to_generated_minute(r) -> GeneratedMinuteResponse:
    """Map a generated_minutes DB row to API response (status-aware)."""
    fp = r.get('file_path') or ''
    is_signed = r.get('is_signed')
    if isinstance(is_signed, str):
        is_signed = is_signed.lower() in ('1', 'true', 'yes')
    return GeneratedMinuteResponse(
        id=r['id'],
        company_name=r.get('company_name') or '',
        meeting_type=r.get('meeting_type') or '',
        meeting_date=str(r.get('meeting_date') or ''),
        file_path=fp,
        created_at=str(r.get('created_at') or ''),
        download_url=f"/api/generated-minutes/download/{fp}" if fp else None,
        meeting_number=r.get('meeting_number') or '',
        meeting_year=r.get('meeting_year') or '',
        status=(r.get('status') or 'draft'),
        finalized_at=str(r['finalized_at']) if r.get('finalized_at') else None,
        finalized_by=r.get('finalized_by'),
        is_signed=bool(is_signed),
        unsigned_file_path=r.get('unsigned_file_path') or None,
    )

class PlaceResponse(BaseModel):
    id: int
    name: str
    address: str
    is_default: bool
    created_at: str

class PlacesListResponse(BaseModel):
    data: List[PlaceResponse]
    count: int

class PlaceCreateRequest(BaseModel):
    name: str
    address: str
    is_default: bool = False
    
class ResolutionTemplateResponse(BaseModel):
    id: int
    template_name: str
    resolution_text: str
    created_at: str

class ResolutionTemplatesList(BaseModel):
    data: List[ResolutionTemplateResponse]
    count: int

class ResolutionCreate(BaseModel):
    template_name: str
    resolution_text: str

class ComplianceResponse(BaseModel):
    id: int
    form: str
    description: str
    due_date: str
    status: str
    priority: str
    company_name: Optional[str] = None
    vertical_name: Optional[str] = None
    created_at: str

class ComplianceCreate(BaseModel):
    form: str
    description: str
    due_date: str
    status: str
    priority: str
    company_name: Optional[str] = None
    vertical_name: Optional[str] = None

class CompliancesList(BaseModel):
    data: List[ComplianceResponse]
    count: int

class MinutesGenerationRequest(BaseModel):
    template: str
    companyName: str
    meetingNumber: str
    meetingType: str
    meetingDay: str
    meetingDate: str
    meetingStartTime: str
    meetingEndTime: str
    meetingPlace: str
    chairmanName: str
    presentDirectors: List[Dict[str, str]]
    inAttendance: List[Dict[str, str]]
    companySecretary: str
    previousMeetingDate: str
    authorisedOfficer: str
    quorum: str
    concerns: str
    declarations: str
    auditorPaymentAmount: str
    auditorPaymentWords: str
    financialYear: str
    agmNumber: str
    agmDay: str
    agmMonthName: str
    agmDate: str
    agmTime: str
    agmPlace: str
    recordingDate: str
    signingDate: str
    signingPlace: str

# --- Database Init ---

def _normalize_company_key(name: Optional[str]) -> str:
    """Normalize company legal names so Ltd. / Limited / spacing variants match."""
    s = (name or "").lower()
    s = re.sub(r"\bprivate\s+limited\b", "pvt ltd", s)
    s = re.sub(r"\bpvt\.?\s*ltd\.?\b", "pvt ltd", s)
    s = re.sub(r"\blimited\b", "ltd", s)
    s = re.sub(r"\bltd\.?\b", "ltd", s)
    return re.sub(r"[^a-z0-9]", "", s)


def _company_names_match(a: Optional[str], b: Optional[str]) -> bool:
    if not a or not b:
        return False
    if a.strip().upper() == b.strip().upper():
        return True
    return _normalize_company_key(a) == _normalize_company_key(b)


def _seed_minutes_directors_from_json(cursor, seed_company: bool = True, seed_external: bool = True) -> int:
    """Load committed JSON seeds into minutes director tables (for machines without local_fallback.db)."""
    seed_dir = os.path.join(os.path.dirname(__file__), "..", "public", "seeds")
    inserted = 0

    if seed_external:
        ebm_path = os.path.join(seed_dir, "minutes_external_board_members.json")
        if os.path.exists(ebm_path):
            with open(ebm_path, encoding="utf-8") as f:
                payload = json.load(f)
            for r in payload.get("members") or []:
                company = (r.get("company_name") or "").strip()
                name = (r.get("name") or "").strip()
                if not company or not name:
                    continue
                din = (r.get("din") or "").strip() or f"SEED-{company[:12]}-{name[:12]}"
                try:
                    cursor.execute(
                        """
                        INSERT INTO external_board_members
                            (din, name, cin, company_name, designation, appointment_date, status, source)
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                        ON CONFLICT DO NOTHING
                        """,
                        (
                            din,
                            name,
                            (r.get("cin") or "").strip(),
                            company,
                            (r.get("designation") or "Director"),
                            r.get("appointment_date"),
                            r.get("status") or "Active",
                            r.get("source") or "SEED",
                        ),
                    )
                    inserted += 1
                except Exception:
                    try:
                        cursor.execute(
                            """
                            INSERT OR IGNORE INTO external_board_members
                                (din, name, cin, company_name, designation, appointment_date, status, source)
                            VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                            """,
                            (
                                din,
                                name,
                                (r.get("cin") or "").strip(),
                                company,
                                (r.get("designation") or "Director"),
                                r.get("appointment_date"),
                                r.get("status") or "Active",
                                r.get("source") or "SEED",
                            ),
                        )
                        inserted += 1
                    except Exception:
                        pass

    if seed_company:
        cd_path = os.path.join(seed_dir, "minutes_company_directors.json")
        if os.path.exists(cd_path):
            with open(cd_path, encoding="utf-8") as f:
                payload = json.load(f)
            for r in payload.get("directors") or []:
                company = (r.get("company_name") or "").strip()
                name = (r.get("name") or "").strip()
                if not company or not name:
                    continue
                try:
                    cursor.execute(
                        """
                        INSERT INTO company_directors (company_name, name, din, designation)
                        VALUES (%s, %s, %s, %s)
                        """,
                        (
                            company,
                            name,
                            (r.get("din") or "").strip(),
                            (r.get("designation") or "Director"),
                        ),
                    )
                    inserted += 1
                except Exception:
                    pass

    return inserted


def init_minutes_pg():
    """Initialize minutes tables in PostgreSQL public schema."""
    target_db = os.getenv('POSTGRES_DATABASE_MINUTES')
    conn = get_pg_connection(target_db)
    if conn:
        try:
            cursor = get_pg_cursor(conn)
            
            # Generated Minutes
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS generated_minutes (
                    id SERIAL PRIMARY KEY,
                    company_name TEXT,
                    meeting_type TEXT,
                    meeting_date TEXT,
                    file_path TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)
            
            # Verticals Table
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS verticals (
                    id SERIAL PRIMARY KEY,
                    name TEXT UNIQUE,
                    code TEXT UNIQUE
                )
            """)

            # Companies Table (Enhanced with audit fields)
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS companies (
                    id SERIAL PRIMARY KEY,
                    name TEXT UNIQUE,
                    code TEXT,
                    cin TEXT,
                    type TEXT,
                    vertical_id INTEGER REFERENCES verticals(id),
                    status TEXT DEFAULT 'Active',
                    secretary_name TEXT,
                    created_by TEXT,
                    updated_by TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)

            # Ensure companies table column extensions exist (safety for SQLite fallback schema updates).
            # NOTE: SQLite forbids non-constant defaults (e.g. CURRENT_TIMESTAMP) on ALTER TABLE ADD COLUMN.
            for col, col_type in [
                ("cin", "TEXT"),
                ("type", "TEXT"),
                ("vertical_id", "INTEGER"),
                ("status", "TEXT DEFAULT 'Active'"),
                ("code", "TEXT"),
                ("secretary_name", "TEXT"),
                ("created_by", "TEXT"),
                ("updated_by", "TEXT"),
                ("created_at", "TIMESTAMP"),
                ("updated_at", "TIMESTAMP"),
            ]:
                try:
                    cursor.execute(f"ALTER TABLE companies ADD COLUMN {col} {col_type}")
                except Exception:
                    pass

            # Ensure generated_minutes schema extensions exist
            for col, col_type in [
                ("vertical_name", "TEXT"),
                ("meeting_number", "TEXT"),
                ("meeting_year", "TEXT"),
                ("status", "TEXT"),
                ("finalized_at", "TIMESTAMP"),
                ("finalized_by", "TEXT"),
                ("is_signed", "INTEGER"),
                ("unsigned_file_path", "TEXT"),
            ]:
                try:
                    cursor.execute(f"ALTER TABLE generated_minutes ADD COLUMN {col} {col_type}")
                except Exception:
                    pass
            # Default existing rows without status → draft (user must finalize intentionally)
            try:
                cursor.execute(
                    "UPDATE generated_minutes SET status = 'draft' WHERE status IS NULL OR status = ''"
                )
            except Exception:
                pass

            # Seed default Verticals
            cursor.execute("SELECT COUNT(*) as count FROM verticals")
            if cursor.fetchone()['count'] == 0:
                logger.info("Seeding default verticals...")
                verticals = [
                    ('Renewables', 'REN'),
                    ('Ports & Logistics', 'PRT'),
                    ('Infrastructure', 'INF'),
                    ('Transmission & Distribution', 'TRD'),
                    ('Realty', 'RLT'),
                    ('Cement', 'CEM'),
                    ('Airport', 'AIR'),
                    ('Natural Resources', 'NAT'),
                    ('Thermal Power', 'THP'),
                    ('Promoter', 'PRO'),
                    ('PLL', 'PLL'),
                    ('Solar Manufacturing', 'SOL'),
                    ('Gas Distribution', 'GAS'),
                    ('Data Centre', 'DAT'),
                    ('Others', 'OTH')
                ]
                for name, code in verticals:
                    try:
                        cursor.execute("INSERT INTO verticals (name, code) VALUES (%s, %s)", (name, code))
                    except Exception:
                        pass

            # Auto-run Excel import if available and database is empty/small
            try:
                cursor.execute("SELECT COUNT(*) as count FROM companies")
                c_cnt = cursor.fetchone()['count']
                if c_cnt < 50:
                    excel_file = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "Vertical and Entity name.xlsx"))
                    if os.path.exists(excel_file):
                        logger.info(f"Triggering automatic Excel migration from {excel_file}...")
                        import subprocess
                        mig_script = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "Backend", "migrate_excel_to_db.py"))
                        if os.path.exists(mig_script):
                            subprocess.Popen(["python", mig_script])
            except Exception as _mig_err:
                logger.warning(f"Auto Excel migration check failed: {_mig_err}")
            
            # Resolution Templates
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS resolution_templates (
                    id SERIAL PRIMARY KEY,
                    template_name TEXT UNIQUE,
                    resolution_text TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)

            # Templates Repository Metadata
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS templates (
                    id SERIAL PRIMARY KEY,
                    template_name TEXT UNIQUE,
                    category TEXT,
                    company_name TEXT,
                    quarter TEXT,
                    file_path TEXT,
                    file_size BIGINT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)
            
            # Compliance Tracking
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS compliances (
                    id SERIAL PRIMARY KEY,
                    form TEXT,
                    description TEXT,
                    due_date TEXT,
                    status TEXT,
                    priority TEXT,
                    company_name TEXT,
                    vertical_name TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)

            # Ensure columns exist
            for col, col_type in [("company_name", "TEXT"), ("vertical_name", "TEXT")]:
                try:
                    cursor.execute(f"ALTER TABLE compliances ADD COLUMN {col} {col_type}")
                except Exception:
                    pass

            # Seed default Compliances
            cursor.execute("SELECT COUNT(*) as count FROM compliances")
            if cursor.fetchone()['count'] == 0:
                logger.info("Seeding default compliances...")
                seeded_compliances = [
                    ('Form MGT-7', 'Annual Return Filing', '2026-11-29', 'Pending', 'High', 'Adani Green Energy Ltd.', 'Energy'),
                    ('Form AOC-4', 'Filing of Audited Financial Statements', '2026-10-30', 'Pending', 'Critical', 'Adani Ports and Special Economic Zone Ltd.', 'Ports & Logistics'),
                    ('Form DIR-12', 'Filing for Appointment of statutory director', '2026-08-15', 'Completed', 'Medium', 'Adani Total Gas Ltd.', 'Gas & Utilities'),
                    ('Form MBP-1', 'First Board Meeting Interest Disclosure', '2026-04-30', 'Completed', 'High', 'Adani Power Ltd.', 'Energy'),
                    ('Form ADT-1', 'Statutory Auditor Appointment Filing', '2026-10-15', 'Pending', 'High', 'Adani Enterprises Ltd.', 'Infrastructure'),
                    ('Form DIR-8', 'Director Disqualification Declaration', '2026-04-30', 'Completed', 'Medium', 'Adani Digital Labs Private Ltd.', 'Digital & Media')
                ]
                for form, desc, due, status, priority, comp, vert in seeded_compliances:
                    cursor.execute(
                        "INSERT INTO compliances (form, description, due_date, status, priority, company_name, vertical_name) VALUES (%s, %s, %s, %s, %s, %s, %s)",
                        (form, desc, due, status, priority, comp, vert)
                    )

            # Places Table (Local to minutes if needed, or shared)
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS places (
                    id SERIAL PRIMARY KEY,
                    name TEXT,
                    address TEXT,
                    is_default BOOLEAN DEFAULT FALSE,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)

            # Local manual-director overlay: writes here never touch the Director Disclosure DB
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS company_directors (
                    id SERIAL PRIMARY KEY,
                    company_name TEXT NOT NULL,
                    name TEXT NOT NULL,
                    din TEXT,
                    designation TEXT DEFAULT 'Director',
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)
            try:
                cursor.execute("ALTER TABLE company_directors ADD COLUMN designation TEXT DEFAULT 'Director'")
            except Exception:
                pass

            cursor.execute("""
                CREATE TABLE IF NOT EXISTS external_board_members (
                    id SERIAL PRIMARY KEY,
                    din TEXT NOT NULL,
                    name TEXT,
                    cin TEXT NOT NULL DEFAULT '',
                    company_name TEXT,
                    designation TEXT DEFAULT 'Director',
                    appointment_date TEXT,
                    status TEXT DEFAULT 'Active',
                    source TEXT DEFAULT 'DISCLOSURE_DOCS',
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    UNIQUE(din, company_name)
                )
            """)

            # Auto-seed directors for local/SQLite if empty (other machines won't have *.db from git)
            try:
                cursor.execute("SELECT COUNT(*) as count FROM company_directors")
                cd_count = cursor.fetchone()["count"] or 0
                cursor.execute("SELECT COUNT(*) as count FROM external_board_members")
                ebm_count = cursor.fetchone()["count"] or 0
                if cd_count == 0 or ebm_count == 0:
                    seeded = _seed_minutes_directors_from_json(
                        cursor,
                        seed_company=(cd_count == 0),
                        seed_external=(ebm_count == 0),
                    )
                    if seeded:
                        logger.info(f"Auto-seeded minutes director data ({seeded} row(s))")
            except Exception as seed_err:
                logger.warning(f"Minutes director auto-seed skipped: {seed_err}")

            # Per-meeting attendance records for real meeting/person-wise reporting
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS meeting_attendance (
                    id SERIAL PRIMARY KEY,
                    minutes_id INTEGER,
                    company_name TEXT,
                    meeting_type TEXT,
                    meeting_date TEXT,
                    director_name TEXT NOT NULL,
                    din TEXT,
                    status TEXT DEFAULT 'Present',
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)

            # Audit Logs Table - Track all critical operations
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS audit_logs (
                    id SERIAL PRIMARY KEY,
                    entity_type TEXT NOT NULL,
                    entity_id INTEGER,
                    entity_name TEXT,
                    action TEXT NOT NULL,
                    performed_by TEXT NOT NULL,
                    performed_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    old_data JSONB,
                    new_data JSONB,
                    ip_address TEXT,
                    user_agent TEXT,
                    remarks TEXT,
                    vertical_id INTEGER,
                    company_name TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)

            # Create indexes for audit_logs for fast querying
            try:
                cursor.execute("CREATE INDEX IF NOT EXISTS idx_audit_logs_entity ON audit_logs(entity_type, entity_id)")
                cursor.execute("CREATE INDEX IF NOT EXISTS idx_audit_logs_user ON audit_logs(performed_by)")
                cursor.execute("CREATE INDEX IF NOT EXISTS idx_audit_logs_time ON audit_logs(performed_at DESC)")
                cursor.execute("CREATE INDEX IF NOT EXISTS idx_audit_logs_company ON audit_logs(company_name)")
            except Exception:
                pass  # Indexes might already exist

            # Draft minutes forms (Save / resume across steps, refresh, Back)
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS minutes_drafts (
                    id SERIAL PRIMARY KEY,
                    draft_key TEXT UNIQUE,
                    company_name TEXT,
                    meeting_type TEXT,
                    meeting_date TEXT,
                    committee_name TEXT,
                    current_step INTEGER DEFAULT 0,
                    form_json TEXT,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)
            try:
                cursor.execute("CREATE INDEX IF NOT EXISTS idx_minutes_drafts_company ON minutes_drafts(company_name)")
            except Exception:
                pass

            # Index public/templates DOCX files into generated_minutes for company meeting stats
            try:
                synced = _sync_template_meetings(cursor)
                if synced:
                    logger.info(f"Synced {synced} template meeting file(s) into generated_minutes")
            except Exception as sync_err:
                logger.warning(f"Template meetings sync skipped: {sync_err}")

            conn.commit()
            logger.info(f"Minutes tables initialized successfully in {target_db or 'default'}")
        except Exception as e:
            conn.rollback()
            logger.error(f"Minutes init failed: {e}")
        finally:
            conn.close()


# Abbreviation → canonical company name (must match companies.name)
TEMPLATE_COMPANY_MAP = {
    "AGEL": "Adani Green Energy Limited",
    "AGE(UP)L": "Adani Green Energy (UP) Limited",
    "AGEUPL": "Adani Green Energy (UP) Limited",
    "AGE25BL": "Adani Green Energy Twenty Five B Limited",
}

TEMPLATE_TYPE_MAP = {
    "BM": "Board Meeting",
    "AC": "Audit Committee",
    "NRC": "Nomination and Remuneration Committee",
    "SRC": "Stakeholders Relationship Committee",
    "CSR": "CSR Committee",
    "RMC": "Risk Management Committee",
    "AGM": "AGM",
    "EGM": "EGM",
}


def _ordinal_meeting_number(num: int) -> str:
    if num <= 0:
        return ""
    suffix = "TH"
    if num % 10 == 1 and num % 100 != 11:
        suffix = "ST"
    elif num % 10 == 2 and num % 100 != 12:
        suffix = "ND"
    elif num % 10 == 3 and num % 100 != 13:
        suffix = "RD"
    return f"{num}{suffix}"


def _parse_template_filename(filename: str) -> Optional[Dict[str, Any]]:
    """
    Parse meeting metadata from template filenames, e.g.:
      90. AGEL - BM - 23.01.2026.docx
      70. AGEL - AC - Minutes - 22.01.2026.docx
      3. AGE(UP)L - AC - 17.10.2025.docx
      59. 04.12.2025 - AGE25BL.docx
      79. 17.10.2025 - AGE(UP)L.docx
    """
    import re
    name = filename
    if name.lower().endswith(".docx"):
        name = name[:-5]

    # Pattern A: {num}. {CODE} - {TYPE} [- Minutes] - {DD.MM.YYYY}
    m = re.match(
        r"^\s*(\d+)\.\s*(.+?)\s*-\s*(BM|AC|NRC|SRC|CSR|RMC|AGM|EGM)\s*(?:-\s*Minutes)?\s*-\s*(\d{2}\.\d{2}\.\d{4})\s*$",
        name,
        flags=re.IGNORECASE,
    )
    if m:
        num, code, mtype, date_str = m.groups()
        code_key = code.strip().upper().replace(" ", "")
        company = TEMPLATE_COMPANY_MAP.get(code_key) or TEMPLATE_COMPANY_MAP.get(code.strip().upper())
        if not company:
            return None
        dd, mm, yyyy = date_str.split(".")
        return {
            "meeting_number": _ordinal_meeting_number(int(num)),
            "meeting_number_int": int(num),
            "company_name": company,
            "company_code": code.strip(),
            "meeting_type": TEMPLATE_TYPE_MAP.get(mtype.upper(), mtype),
            "meeting_date": f"{yyyy}-{mm}-{dd}",
            "meeting_year": yyyy,
            "file_path": filename,
        }

    # Pattern B: {num}. {DD.MM.YYYY} - {CODE}
    m = re.match(
        r"^\s*(\d+)\.\s*(\d{2}\.\d{2}\.\d{4})\s*-\s*(.+?)\s*$",
        name,
        flags=re.IGNORECASE,
    )
    if m:
        num, date_str, code = m.groups()
        code_key = code.strip().upper().replace(" ", "")
        company = TEMPLATE_COMPANY_MAP.get(code_key) or TEMPLATE_COMPANY_MAP.get(code.strip().upper())
        if not company:
            return None
        dd, mm, yyyy = date_str.split(".")
        return {
            "meeting_number": _ordinal_meeting_number(int(num)),
            "meeting_number_int": int(num),
            "company_name": company,
            "company_code": code.strip(),
            "meeting_type": "Board Meeting",  # type omitted in this naming style
            "meeting_date": f"{yyyy}-{mm}-{dd}",
            "meeting_year": yyyy,
            "file_path": filename,
        }

    return None


def _sync_template_meetings(cursor) -> int:
    """
    Scan public/templates for real meeting DOCX files and upsert into generated_minutes.
    Also stamps company codes (AGEL, etc.) onto matching companies rows.
    Returns number of inserted/updated rows.
    """
    templates_dir = os.path.join(os.path.dirname(__file__), "..", "public", "templates")
    if not os.path.isdir(templates_dir):
        return 0

    # Ensure company codes for known abbreviations (helps UI abbreviations)
    for code, company_name in TEMPLATE_COMPANY_MAP.items():
        try:
            cursor.execute(
                """
                UPDATE companies
                SET code = %s
                WHERE UPPER(name) = UPPER(%s)
                  AND (code IS NULL OR code = '')
                """,
                (code if code != "AGEUPL" else "AGE(UP)L", company_name),
            )
        except Exception:
            pass

    synced = 0
    for fname in os.listdir(templates_dir):
        if not fname.lower().endswith(".docx"):
            continue
        # Skip custom uploads / generated copies parked in templates
        if fname.lower().startswith("custom_") or fname.lower().startswith("meeting_minutes_"):
            continue

        meta = _parse_template_filename(fname)
        if not meta:
            continue

        cursor.execute(
            "SELECT id FROM generated_minutes WHERE file_path = %s",
            (meta["file_path"],),
        )
        existing = cursor.fetchone()
        if existing:
            cursor.execute(
                """
                UPDATE generated_minutes
                SET company_name = %s,
                    meeting_type = %s,
                    meeting_date = %s,
                    meeting_number = %s,
                    meeting_year = %s,
                    status = CASE
                        WHEN COALESCE(is_signed, 0) = 1 THEN status
                        WHEN LOWER(COALESCE(status, '')) = 'finalized' AND finalized_at IS NOT NULL THEN status
                        ELSE COALESCE(NULLIF(status, ''), 'draft')
                    END
                WHERE id = %s
                """,
                (
                    meta["company_name"],
                    meta["meeting_type"],
                    meta["meeting_date"],
                    meta["meeting_number"],
                    meta["meeting_year"],
                    existing["id"],
                ),
            )
        else:
            cursor.execute(
                """
                INSERT INTO generated_minutes
                    (company_name, meeting_type, meeting_date, file_path, meeting_number, meeting_year, status)
                VALUES (%s, %s, %s, %s, %s, %s, %s)
                """,
                (
                    meta["company_name"],
                    meta["meeting_type"],
                    meta["meeting_date"],
                    meta["file_path"],
                    meta["meeting_number"],
                    meta["meeting_year"],
                    "draft",
                ),
            )
        synced += 1

    return synced


# Ensure tables exist on module load (idempotent: CREATE TABLE IF NOT EXISTS)
try:
    init_minutes_pg()
except Exception as _init_err:
    logger.error(f"init_minutes_pg on import failed: {_init_err}")

# --- API Endpoints ---

@router.post("/templates/sync-meetings")
async def sync_template_meetings_endpoint():
    """Re-scan public/templates and sync real meeting files into generated_minutes."""
    try:
        def run():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn:
                raise RuntimeError("Database connection unavailable")
            cursor = get_pg_cursor(conn)
            try:
                count = _sync_template_meetings(cursor)
                conn.commit()
                return count
            except Exception:
                conn.rollback()
                raise
            finally:
                conn.close()

        count = await asyncio.get_running_loop().run_in_executor(thread_pool, run)
        return {"success": True, "synced": count, "message": f"Synced {count} template meeting file(s)"}
    except Exception as e:
        logger.error(f"Template sync failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


class MinutesDraftUpsertRequest(BaseModel):
    draft_key: str
    company_name: str = ""
    meeting_type: str = ""
    meeting_date: str = ""
    committee_name: str = ""
    current_step: int = 0
    form_data: Dict[str, Any] = {}


@router.put("/minutes-drafts")
async def upsert_minutes_draft(request: MinutesDraftUpsertRequest):
    """Save / update a multi-step minutes form draft (survives refresh, Back, reconnect)."""
    try:
        def save():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn:
                raise RuntimeError("Database connection unavailable")
            cursor = get_pg_cursor(conn)
            try:
                # Ensure table exists (for servers that started before this migration)
                cursor.execute("""
                    CREATE TABLE IF NOT EXISTS minutes_drafts (
                        id SERIAL PRIMARY KEY,
                        draft_key TEXT UNIQUE,
                        company_name TEXT,
                        meeting_type TEXT,
                        meeting_date TEXT,
                        committee_name TEXT,
                        current_step INTEGER DEFAULT 0,
                        form_json TEXT,
                        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                    )
                """)
                form_json = json.dumps(request.form_data or {})
                cursor.execute("SELECT id FROM minutes_drafts WHERE draft_key = %s", (request.draft_key,))
                row = cursor.fetchone()
                if row:
                    cursor.execute(
                        """
                        UPDATE minutes_drafts
                        SET company_name = %s, meeting_type = %s, meeting_date = %s,
                            committee_name = %s, current_step = %s, form_json = %s,
                            updated_at = CURRENT_TIMESTAMP
                        WHERE draft_key = %s
                        RETURNING id, draft_key, current_step, updated_at
                        """,
                        (
                            request.company_name, request.meeting_type, request.meeting_date,
                            request.committee_name, request.current_step, form_json, request.draft_key,
                        ),
                    )
                else:
                    cursor.execute(
                        """
                        INSERT INTO minutes_drafts
                            (draft_key, company_name, meeting_type, meeting_date, committee_name, current_step, form_json)
                        VALUES (%s, %s, %s, %s, %s, %s, %s)
                        RETURNING id, draft_key, current_step, updated_at
                        """,
                        (
                            request.draft_key, request.company_name, request.meeting_type,
                            request.meeting_date, request.committee_name, request.current_step, form_json,
                        ),
                    )
                saved = cursor.fetchone()
                conn.commit()
                return {
                    "id": saved["id"],
                    "draft_key": saved["draft_key"],
                    "current_step": saved["current_step"],
                    "updated_at": str(saved["updated_at"]),
                }
            except Exception:
                conn.rollback()
                raise
            finally:
                conn.close()

        return await asyncio.get_running_loop().run_in_executor(thread_pool, save)
    except Exception as e:
        logger.error(f"Failed to save minutes draft: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/minutes-drafts")
async def get_minutes_draft(draft_key: str):
    """Load a saved minutes form draft by key."""
    try:
        def fetch():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn:
                raise RuntimeError("Database connection unavailable")
            cursor = get_pg_cursor(conn)
            try:
                cursor.execute(
                    """
                    SELECT id, draft_key, company_name, meeting_type, meeting_date, committee_name,
                           current_step, form_json, updated_at
                    FROM minutes_drafts WHERE draft_key = %s
                    """,
                    (draft_key,),
                )
                row = cursor.fetchone()
                if not row:
                    return None
                form_data = {}
                try:
                    form_data = json.loads(row["form_json"] or "{}")
                except Exception:
                    form_data = {}
                return {
                    "id": row["id"],
                    "draft_key": row["draft_key"],
                    "company_name": row["company_name"],
                    "meeting_type": row["meeting_type"],
                    "meeting_date": row["meeting_date"],
                    "committee_name": row.get("committee_name") or "",
                    "current_step": row["current_step"] or 0,
                    "form_data": form_data,
                    "updated_at": str(row["updated_at"]),
                }
            finally:
                conn.close()

        data = await asyncio.get_running_loop().run_in_executor(thread_pool, fetch)
        if not data:
            raise HTTPException(status_code=404, detail="Draft not found")
        return data
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to load minutes draft: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/minutes-drafts")
async def delete_minutes_draft(draft_key: str):
    """Delete a minutes form draft (after finalize or Reset)."""
    try:
        def delete():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn:
                raise RuntimeError("Database connection unavailable")
            cursor = get_pg_cursor(conn)
            try:
                cursor.execute("DELETE FROM minutes_drafts WHERE draft_key = %s", (draft_key,))
                conn.commit()
                return cursor.rowcount or 0
            except Exception:
                conn.rollback()
                raise
            finally:
                conn.close()

        deleted = await asyncio.get_running_loop().run_in_executor(thread_pool, delete)
        return {"success": True, "deleted": deleted}
    except Exception as e:
        logger.error(f"Failed to delete minutes draft: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/generated-minutes", response_model=MinutesHistoryResponse)
async def get_history():
    """Get history of generated minutes from PostgreSQL."""
    try:
        def fetch():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn: raise RuntimeError("Database connection unavailable")
            cursor = get_pg_cursor(conn)
            try:
                cursor.execute(
                    """
                    SELECT id, company_name, meeting_type, meeting_date, file_path, created_at,
                           meeting_number, meeting_year, status, finalized_at, finalized_by,
                           is_signed, unsigned_file_path
                    FROM generated_minutes
                    ORDER BY id DESC
                    """
                )
                rows = cursor.fetchall()
                data = [_row_to_generated_minute(r) for r in rows]
                return data, len(data)
            finally:
                conn.close()
        
        loop = asyncio.get_running_loop()
        data, count = await loop.run_in_executor(thread_pool, fetch)
        return MinutesHistoryResponse(data=data, count=count)
    except Exception as e:
        logger.error(f"Error fetching history: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/minutes/history", response_model=MinutesHistoryResponse)
async def get_minutes_history_post():
    """Fallback for POST history request."""
    return await get_history()

@router.post("/generated-minutes/{id}/finalize")
async def finalize_minute(id: int, request: Request):
    """
    Lock an approved minutes record (Draft → Finalized).
    Finalized records cannot be deleted casually.
    """
    try:
        body = {}
        try:
            body = await request.json()
        except Exception:
            body = {}
        finalized_by = (body.get("finalized_by") or body.get("user_email") or "system").strip()

        def run():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn:
                raise RuntimeError("Database connection unavailable")
            cursor = get_pg_cursor(conn)
            try:
                for col, col_type in [("status", "TEXT"), ("finalized_at", "TIMESTAMP"), ("finalized_by", "TEXT")]:
                    try:
                        cursor.execute(f"ALTER TABLE generated_minutes ADD COLUMN {col} {col_type}")
                    except Exception:
                        pass

                cursor.execute(
                    """
                    SELECT id, company_name, meeting_type, meeting_date, meeting_number, status
                    FROM generated_minutes WHERE id = %s
                    """,
                    (id,),
                )
                row = cursor.fetchone()
                if not row:
                    return None, "not_found"
                if (row.get("status") or "").lower() == "finalized":
                    return dict(row), "already_finalized"

                cursor.execute(
                    """
                    UPDATE generated_minutes
                    SET status = 'finalized',
                        finalized_at = CURRENT_TIMESTAMP,
                        finalized_by = %s
                    WHERE id = %s
                    RETURNING id, company_name, meeting_type, meeting_date, meeting_number,
                              status, finalized_at, finalized_by
                    """,
                    (finalized_by, id),
                )
                updated = cursor.fetchone()
                try:
                    from utils.audit_logger import AuditLogger
                    AuditLogger.log_meeting_finalized(
                        conn=conn,
                        meeting_id=id,
                        company_name=updated.get("company_name") or "",
                        meeting_type=updated.get("meeting_type") or "",
                        meeting_number=updated.get("meeting_number") or "",
                        user_email=finalized_by,
                    )
                except Exception as audit_err:
                    logger.warning(f"Finalize audit log skipped: {audit_err}")
                conn.commit()
                return dict(updated), "ok"
            except Exception:
                conn.rollback()
                raise
            finally:
                conn.close()

        result, code = await asyncio.get_running_loop().run_in_executor(thread_pool, run)
        if code == "not_found":
            raise HTTPException(status_code=404, detail="Minutes record not found")
        return {
            "success": True,
            "already_finalized": code == "already_finalized",
            "id": result["id"],
            "status": "finalized",
            "company_name": result.get("company_name"),
            "meeting_type": result.get("meeting_type"),
            "finalized_by": result.get("finalized_by") or finalized_by,
            "finalized_at": str(result.get("finalized_at") or ""),
            "message": "Minutes already finalized" if code == "already_finalized" else "Minutes finalized and locked",
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to finalize minutes {id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))


def _resolve_minutes_file_path(filename: str) -> Optional[str]:
    """Find a minutes file under public/generated or public/templates."""
    if not filename or os.path.basename(filename) != filename:
        return None
    base = os.path.join(os.path.dirname(__file__), "..", "public")
    for sub in ("generated", "templates"):
        fp = os.path.join(base, sub, filename)
        if os.path.exists(fp):
            return fp
    return None


def _extract_chairman_from_docx(file_path: str) -> Optional[str]:
    """
    Read a minutes DOCX and find the meeting chairman dynamically.
    Prefers attendance lines like 'Mr. X - Chairman', then 'X occupied the Chair'.
    """
    if not file_path or not os.path.exists(file_path) or not DOCX_AVAILABLE:
        return None
    try:
        doc = Document(file_path)
        lines = [p.text.strip() for p in doc.paragraphs if (p.text or "").strip()]
        # Also scan tables lightly
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        t = (p.text or "").strip()
                        if t:
                            lines.append(t)

        # 1) "Name - Chairman" / "Name – Chairman" in attendance
        for line in lines:
            m = re.search(
                r"(?:Mr\.|Mrs\.|Ms\.|Dr\.)?\s*([A-Za-z][A-Za-z\.\s]+?)\s*[-–—]\s*Chairman\b",
                line,
                re.IGNORECASE,
            )
            if m:
                name = re.sub(r"\s+", " ", m.group(1)).strip(" .")
                if name and len(name) > 2:
                    return name

        # 2) "Name occupied the Chair"
        for line in lines:
            m = re.search(
                r"(?:Mr\.|Mrs\.|Ms\.|Dr\.)?\s*([A-Za-z][A-Za-z\.\s]+?)(?:,\s*Chairman)?\s+occupied the [Cc]hair",
                line,
            )
            if m:
                name = re.sub(r"\s+", " ", m.group(1)).strip(" .")
                if name and len(name) > 2 and "thereafter" not in name.lower():
                    return name
    except Exception as ex:
        logger.warning(f"Chairman extract failed for {file_path}: {ex}")
    return None


def _meeting_type_matches(stored: Optional[str], wanted: Optional[str]) -> bool:
    if not wanted or wanted.lower() in ("all", ""):
        return True
    if not stored:
        return False
    a = stored.lower().strip()
    b = wanted.lower().strip()
    if a == b:
        return True
    # Board vs Board Meeting; Audit Committee vs Committee Meeting + committee
    if "board" in a and "board" in b:
        return True
    if "audit" in a and "audit" in b:
        return True
    if b in a or a in b:
        return True
    return False


def _resolve_default_chairman(company_name: str, meeting_type: str = "") -> Dict[str, Any]:
    """
    Resolve meeting chairman dynamically for a company + meeting type:
      1) Most recent generated_minutes DOCX for same company (+ type)
      2) Matching template DOCX for same company (+ type)
    """
    result = {
        "chairman_name": "",
        "source": None,
        "file_path": None,
        "company_name": company_name,
        "meeting_type": meeting_type or "",
    }
    if not (company_name or "").strip():
        return result

    # Scan generated_minutes (newest first)
    try:
        conn = get_pg_connection(os.getenv("POSTGRES_DATABASE_MINUTES"))
        if conn:
            cursor = get_pg_cursor(conn)
            try:
                cursor.execute(
                    """
                    SELECT id, company_name, meeting_type, file_path, meeting_date
                    FROM generated_minutes
                    WHERE company_name IS NOT NULL AND file_path IS NOT NULL
                    ORDER BY id DESC
                    """
                )
                rows = cursor.fetchall()
            finally:
                conn.close()

            for r in rows:
                if not _company_names_match(company_name, r.get("company_name")):
                    continue
                if meeting_type and not _meeting_type_matches(r.get("meeting_type"), meeting_type):
                    continue
                fp = _resolve_minutes_file_path(r.get("file_path") or "")
                if not fp:
                    continue
                chair = _extract_chairman_from_docx(fp)
                if chair:
                    result.update({
                        "chairman_name": chair,
                        "source": "previous_minutes",
                        "file_path": r.get("file_path"),
                        "meeting_type": r.get("meeting_type") or meeting_type,
                    })
                    return result
    except Exception as ex:
        logger.warning(f"default chairman DB scan failed: {ex}")

    # Fall back to templates on disk for this company
    templates_dir = os.path.join(os.path.dirname(__file__), "..", "public", "templates")
    if os.path.isdir(templates_dir):
        candidates = []
        for fname in os.listdir(templates_dir):
            if not fname.lower().endswith(".docx"):
                continue
            if fname.lower().startswith("custom_") or fname.lower().startswith("meeting_minutes_"):
                continue
            meta = _parse_template_filename(fname)
            if not meta:
                continue
            if not _company_names_match(company_name, meta.get("company_name")):
                continue
            if meeting_type and not _meeting_type_matches(meta.get("meeting_type"), meeting_type):
                continue
            candidates.append((meta.get("meeting_date") or "", fname, meta))
        candidates.sort(reverse=True)  # newest date first
        for _, fname, meta in candidates:
            fp = os.path.join(templates_dir, fname)
            chair = _extract_chairman_from_docx(fp)
            if chair:
                result.update({
                    "chairman_name": chair,
                    "source": "template",
                    "file_path": fname,
                    "meeting_type": meta.get("meeting_type") or meeting_type,
                })
                return result

    # Seed JSON extracted from templates (portable across machines without scanning docs)
    seed_path = os.path.join(os.path.dirname(__file__), "..", "public", "seeds", "minutes_default_chairmen.json")
    if os.path.exists(seed_path):
        try:
            with open(seed_path, encoding="utf-8") as f:
                payload = json.load(f)
            for row in payload.get("chairmen") or []:
                if not _company_names_match(company_name, row.get("company_name")):
                    continue
                if meeting_type and not _meeting_type_matches(row.get("meeting_type"), meeting_type):
                    continue
                chair = (row.get("chairman_name") or "").strip()
                if chair:
                    result.update({
                        "chairman_name": chair,
                        "source": "template_seed",
                        "file_path": row.get("file_path"),
                        "meeting_type": row.get("meeting_type") or meeting_type,
                    })
                    return result
        except Exception as ex:
            logger.warning(f"default chairman seed read failed: {ex}")

    return result


@router.get("/templates/{filename}/chairman")
async def get_template_chairman(filename: str):
    """Extract Meeting Chairman from a specific official template DOCX."""
    safe = os.path.basename(filename or "")
    if not safe or safe != filename:
        raise HTTPException(status_code=400, detail="Invalid filename")
    fp = os.path.join(os.path.dirname(__file__), "..", "public", "templates", safe)
    if not os.path.exists(fp):
        raise HTTPException(status_code=404, detail="Template not found")
    try:
        loop = asyncio.get_running_loop()
        chair = await loop.run_in_executor(thread_pool, lambda: _extract_chairman_from_docx(fp))
        meta = _parse_template_filename(safe) or {}
        return {
            "success": True,
            "filename": safe,
            "chairman_name": chair or "",
            "company_name": meta.get("company_name"),
            "meeting_type": meta.get("meeting_type"),
            "meeting_date": meta.get("meeting_date"),
        }
    except Exception as e:
        logger.error(f"template chairman extract failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/companies/{company_name}/default-chairman")
async def get_default_chairman(company_name: str, meeting_type: Optional[str] = None):
    """
    Auto-resolve Meeting Chairman for a company (and optional meeting type)
    from previous minutes / templates — not hardcoded.
    """
    try:
        loop = asyncio.get_running_loop()
        data = await loop.run_in_executor(
            thread_pool,
            lambda: _resolve_default_chairman(company_name, meeting_type or ""),
        )
        return {"success": True, **data}
    except Exception as e:
        logger.error(f"default-chairman failed: {e}")
        return {
            "success": False,
            "chairman_name": "",
            "source": None,
            "company_name": company_name,
            "meeting_type": meeting_type or "",
            "detail": str(e),
        }


@router.post("/generated-minutes/{id}/replace-file")
async def replace_minutes_file(
    id: int,
    file: UploadFile = File(...),
):
    """
    Replace the working document for a DRAFT minutes record (fix mistakes before finalize).
    Finalized unsigned docs can also be replaced only via upload-signed.
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    lower = file.filename.lower()
    if not (lower.endswith('.docx') or lower.endswith('.pdf')):
        raise HTTPException(status_code=400, detail="Only .docx or .pdf files are supported")

    try:
        content = await file.read()

        def run():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn:
                raise RuntimeError("Database connection unavailable")
            cursor = get_pg_cursor(conn)
            try:
                for col, col_type in [("is_signed", "INTEGER"), ("unsigned_file_path", "TEXT"), ("status", "TEXT")]:
                    try:
                        cursor.execute(f"ALTER TABLE generated_minutes ADD COLUMN {col} {col_type}")
                    except Exception:
                        pass

                cursor.execute(
                    "SELECT id, file_path, status, is_signed FROM generated_minutes WHERE id = %s",
                    (id,),
                )
                row = cursor.fetchone()
                if not row:
                    return None, "not_found"
                if (row.get("status") or "").lower() == "finalized" and bool(row.get("is_signed")):
                    return None, "signed_locked"

                old_name = row.get("file_path") or ""
                ext = os.path.splitext(file.filename)[1].lower() or ".docx"
                new_name = f"revised_{id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}{ext}"
                generated_dir = os.path.join(os.path.dirname(__file__), "..", "public", "generated")
                os.makedirs(generated_dir, exist_ok=True)
                new_path = os.path.join(generated_dir, new_name)
                with open(new_path, "wb") as f:
                    f.write(content)

                # Keep old file as backup name if present
                old_fp = _resolve_minutes_file_path(old_name)
                if old_fp and os.path.exists(old_fp):
                    backup = os.path.join(
                        os.path.dirname(old_fp),
                        f"backup_before_revise_{id}_{os.path.basename(old_name)}",
                    )
                    try:
                        if not os.path.exists(backup):
                            shutil.copy2(old_fp, backup)
                    except Exception:
                        pass

                cursor.execute(
                    """
                    UPDATE generated_minutes
                    SET file_path = %s, status = 'draft', is_signed = 0
                    WHERE id = %s
                    RETURNING id, file_path, status
                    """,
                    (new_name, id),
                )
                updated = cursor.fetchone()
                conn.commit()
                return dict(updated), "ok"
            except Exception:
                conn.rollback()
                raise
            finally:
                conn.close()

        result, code = await asyncio.get_running_loop().run_in_executor(thread_pool, run)
        if code == "not_found":
            raise HTTPException(status_code=404, detail="Minutes record not found")
        if code == "signed_locked":
            raise HTTPException(status_code=409, detail="Signed final document is locked. Ask Master Admin to unlock.")
        return {
            "success": True,
            "id": result["id"],
            "file_path": result["file_path"],
            "status": result.get("status") or "draft",
            "download_url": f"/api/generated-minutes/download/{result['file_path']}",
            "message": "Document replaced. Record is Draft — review and Finalize when ready.",
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to replace minutes file {id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/generated-minutes/{id}/upload-signed")
async def upload_signed_minutes(
    id: int,
    file: UploadFile = File(...),
    confirm_final: str = Form("false"),
    uploaded_by: str = Form("user"),
):
    """
    Upload signed PDF/DOCX after approval.
    Requires confirm_final=true. Overrides the unsigned file (kept as unsigned_file_path backup)
    and marks the record Final + signed.
    """
    if str(confirm_final).lower() not in ("1", "true", "yes"):
        raise HTTPException(
            status_code=400,
            detail="Confirmation required. Set confirm_final=true after user confirms this is the final signed copy.",
        )
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    lower = file.filename.lower()
    if not (lower.endswith('.pdf') or lower.endswith('.docx')):
        raise HTTPException(status_code=400, detail="Only signed .pdf or .docx files are supported")

    try:
        content = await file.read()

        def run():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn:
                raise RuntimeError("Database connection unavailable")
            cursor = get_pg_cursor(conn)
            try:
                for col, col_type in [
                    ("status", "TEXT"),
                    ("finalized_at", "TIMESTAMP"),
                    ("finalized_by", "TEXT"),
                    ("is_signed", "INTEGER"),
                    ("unsigned_file_path", "TEXT"),
                ]:
                    try:
                        cursor.execute(f"ALTER TABLE generated_minutes ADD COLUMN {col} {col_type}")
                    except Exception:
                        pass

                cursor.execute(
                    "SELECT id, file_path, status, is_signed, unsigned_file_path FROM generated_minutes WHERE id = %s",
                    (id,),
                )
                row = cursor.fetchone()
                if not row:
                    return None, "not_found"

                old_name = row.get("file_path") or ""
                unsigned_keep = row.get("unsigned_file_path") or old_name

                ext = os.path.splitext(file.filename)[1].lower() or ".pdf"
                new_name = f"signed_{id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}{ext}"
                generated_dir = os.path.join(os.path.dirname(__file__), "..", "public", "generated")
                os.makedirs(generated_dir, exist_ok=True)
                new_path = os.path.join(generated_dir, new_name)
                with open(new_path, "wb") as f:
                    f.write(content)

                # Backup unsigned file path reference (do not delete unsigned immediately)
                old_fp = _resolve_minutes_file_path(old_name)
                if old_fp and os.path.exists(old_fp) and not (row.get("is_signed")):
                    backup_name = f"unsigned_backup_{id}_{os.path.basename(old_name)}"
                    backup_path = os.path.join(generated_dir, backup_name)
                    try:
                        if not os.path.exists(backup_path):
                            shutil.copy2(old_fp, backup_path)
                        unsigned_keep = backup_name
                    except Exception as copy_err:
                        logger.warning(f"Could not backup unsigned file: {copy_err}")
                        unsigned_keep = old_name

                cursor.execute(
                    """
                    UPDATE generated_minutes
                    SET file_path = %s,
                        unsigned_file_path = %s,
                        is_signed = 1,
                        status = 'finalized',
                        finalized_at = CURRENT_TIMESTAMP,
                        finalized_by = %s
                    WHERE id = %s
                    RETURNING id, file_path, status, is_signed, unsigned_file_path, finalized_at, finalized_by
                    """,
                    (new_name, unsigned_keep, uploaded_by, id),
                )
                updated = cursor.fetchone()
                try:
                    from utils.audit_logger import AuditLogger
                    AuditLogger.log_meeting_finalized(
                        conn=conn,
                        meeting_id=id,
                        company_name="",
                        meeting_type="",
                        meeting_number="",
                        user_email=uploaded_by,
                    )
                except Exception:
                    pass
                conn.commit()
                return dict(updated), "ok"
            except Exception:
                conn.rollback()
                raise
            finally:
                conn.close()

        result, code = await asyncio.get_running_loop().run_in_executor(thread_pool, run)
        if code == "not_found":
            raise HTTPException(status_code=404, detail="Minutes record not found")
        return {
            "success": True,
            "id": result["id"],
            "file_path": result["file_path"],
            "status": "finalized",
            "is_signed": True,
            "unsigned_file_path": result.get("unsigned_file_path"),
            "download_url": f"/api/generated-minutes/download/{result['file_path']}",
            "message": "Signed document uploaded. It replaced the unsigned file and is now Final.",
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to upload signed minutes {id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/generated-minutes/{id}")
async def delete_minute(id: int):
    """Delete a minutes record. Signed finals are blocked; draft/final unsigned can be removed."""
    try:
        def delete():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn: raise RuntimeError("Database connection unavailable")
            cursor = get_pg_cursor(conn)
            try:
                try:
                    cursor.execute(f"ALTER TABLE generated_minutes ADD COLUMN is_signed INTEGER")
                except Exception:
                    pass

                cursor.execute(
                    "SELECT file_path, status, is_signed FROM generated_minutes WHERE id = %s",
                    (id,),
                )
                row = cursor.fetchone()
                if not row:
                    return {"success": False, "reason": "not_found"}
                if bool(row.get("is_signed")):
                    return {"success": False, "reason": "signed"}

                fp_name = row.get("file_path") or ""
                # Only remove generated copies — never delete shared templates on disk
                if fp_name and os.path.basename(fp_name) == fp_name:
                    generated_fp = os.path.join(
                        os.path.dirname(__file__), "..", "public", "generated", fp_name
                    )
                    if os.path.exists(generated_fp):
                        try:
                            os.remove(generated_fp)
                        except Exception:
                            pass

                cursor.execute("DELETE FROM generated_minutes WHERE id = %s", (id,))
                conn.commit()
                return {"success": True, "reason": "deleted"}
            except Exception:
                conn.rollback()
                raise
            finally:
                conn.close()
        result = await asyncio.get_running_loop().run_in_executor(thread_pool, delete)
        if result.get("reason") == "signed":
            raise HTTPException(
                status_code=409,
                detail="This signed final document is locked and cannot be deleted.",
            )
        if result.get("reason") == "not_found":
            raise HTTPException(status_code=404, detail="Minutes record not found")
        return result
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

def _minutes_media_type(filename: str) -> str:
    lower = (filename or "").lower()
    if lower.endswith(".pdf"):
        return "application/pdf"
    return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


@router.get("/generated-minutes/download/{filename}")
@router.get("/templates/download/{filename}")
async def download_file(filename: str):
    if os.path.basename(filename) != filename:
        raise HTTPException(status_code=400, detail="Invalid filename")
    fp = _resolve_minutes_file_path(filename)
    if not fp:
        raise HTTPException(status_code=404, detail="File not found")
    from fastapi.responses import FileResponse
    return FileResponse(
        path=fp,
        filename=filename,
        media_type=_minutes_media_type(filename),
        content_disposition_type="attachment",
    )


@router.get("/generated-minutes/view/{filename}")
async def view_file_inline(filename: str):
    """Open document inline in the browser (View) — does not force download."""
    if os.path.basename(filename) != filename:
        raise HTTPException(status_code=400, detail="Invalid filename")
    fp = _resolve_minutes_file_path(filename)
    if not fp:
        raise HTTPException(status_code=404, detail="File not found")
    from fastapi.responses import FileResponse
    return FileResponse(
        path=fp,
        filename=filename,
        media_type=_minutes_media_type(filename),
        content_disposition_type="inline",
    )


class MinutesContentUpdate(BaseModel):
    extracted_text: str
    edited_by: Optional[str] = "user"


@router.put("/generated-minutes/{id}/content")
async def update_minutes_content(id: int, body: MinutesContentUpdate):
    """
    Save in-app text edits back to the working DOCX (draft / unsigned only).
    Rebuilds paragraph text from the edited content and keeps the record as draft.
    """
    if not DOCX_AVAILABLE:
        raise HTTPException(status_code=500, detail="python-docx is not installed on the server")
    try:
        def run():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn:
                raise RuntimeError("Database connection unavailable")
            cursor = get_pg_cursor(conn)
            try:
                for col, col_type in [("is_signed", "INTEGER"), ("status", "TEXT")]:
                    try:
                        cursor.execute(f"ALTER TABLE generated_minutes ADD COLUMN {col} {col_type}")
                    except Exception:
                        pass

                cursor.execute(
                    "SELECT id, file_path, status, is_signed FROM generated_minutes WHERE id = %s",
                    (id,),
                )
                row = cursor.fetchone()
                if not row:
                    raise HTTPException(status_code=404, detail="Minutes record not found")
                if bool(row.get("is_signed")):
                    raise HTTPException(status_code=403, detail="Signed final documents cannot be edited")

                old_name = row.get("file_path") or ""
                old_fp = _resolve_minutes_file_path(old_name)

                doc = Document()
                text = (body.extracted_text or "").replace("\r\n", "\n")
                for line in text.split("\n"):
                    doc.add_paragraph(line)

                generated_dir = os.path.join(os.path.dirname(__file__), "..", "public", "generated")
                os.makedirs(generated_dir, exist_ok=True)
                stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                base_name = os.path.splitext(os.path.basename(old_name) or "minutes")[0]
                # Strip previous edit stamps
                base_name = re.sub(r"_edited_\d{8}_\d{6}$", "", base_name)
                new_name = f"{base_name}_edited_{stamp}.docx"
                new_path = os.path.join(generated_dir, new_name)
                doc.save(new_path)

                cursor.execute(
                    """
                    UPDATE generated_minutes
                    SET file_path = %s, status = 'draft', is_signed = 0
                    WHERE id = %s
                    RETURNING id, file_path, status, company_name, meeting_type, meeting_date
                    """,
                    (new_name, id),
                )
                updated = cursor.fetchone()
                conn.commit()
                return dict(updated) if updated else {"id": id, "file_path": new_name, "status": "draft"}
            finally:
                conn.close()

        loop = asyncio.get_running_loop()
        result = await loop.run_in_executor(thread_pool, run)
        return {
            "success": True,
            "message": "Document updated and saved as draft",
            "id": result.get("id"),
            "file_path": result.get("file_path"),
            "status": result.get("status", "draft"),
            "download_url": f"/api/generated-minutes/download/{result.get('file_path')}",
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to update minutes content {id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/templates")
async def list_templates():
    """List templates from database repository with metadata."""
    td = os.path.join(os.path.dirname(__file__), "..", "public", "templates")
    fs = []
    
    # Query database table first
    db_path = os.path.join(os.path.dirname(__file__), "..", "public", "local_fallback.db")
    if os.path.exists(db_path):
        import sqlite3
        try:
            conn = sqlite3.connect(db_path)
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            cursor.execute("SELECT id, template_name, category, company_name, quarter, file_path, file_size, created_at FROM templates ORDER BY template_name")
            rows = cursor.fetchall()
            for r in rows:
                fs.append({
                    "id": r["id"],
                    "name": r["template_name"],
                    "category": r["category"],
                    "companyName": r["company_name"],
                    "quarter": r["quarter"],
                    "size": r["file_size"],
                    "lastModified": str(r["created_at"]),
                    "path": r["file_path"]
                })
            conn.close()
            if fs:
                return {"data": fs, "count": len(fs)}
        except Exception as err:
            logger.warning(f"Error querying templates table: {err}")

    # Fallback to filesystem scan if DB query is empty
    if os.path.exists(td):
        for f in os.listdir(td):
            if f.endswith('.docx') and not f.startswith('~'):
                stats = os.stat(os.path.join(td, f))
                fs.append({"name": f, "size": stats.st_size, "lastModified": datetime.fromtimestamp(stats.st_mtime).strftime('%Y-%m-%d %H:%M:%S'), "path": f})
    return {"data": fs, "count": len(fs)}

# --- Place Endpoints ---

@router.get("/places", response_model=PlacesListResponse)
async def get_places():
    try:
        def fetch():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn: raise RuntimeError("Database connection unavailable")
            cursor = get_pg_cursor(conn)
            try:
                cursor.execute("SELECT id, name, address, is_default, created_at FROM places ORDER BY name")
                rows = cursor.fetchall()
                return [PlaceResponse(id=r['id'], name=r['name'], address=r['address'], is_default=r['is_default'], created_at=str(r['created_at'])) for r in rows]
            finally:
                conn.close()
        data = await asyncio.get_running_loop().run_in_executor(thread_pool, fetch)
        return PlacesListResponse(data=data, count=len(data))
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# --- Compliance Endpoints ---

@router.get("/compliances", response_model=CompliancesList)
async def get_compliances():
    try:
        def fetch():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn: raise RuntimeError("Database connection unavailable")
            cursor = get_pg_cursor(conn)
            try:
                cursor.execute("SELECT id, form, description, due_date, status, priority, company_name, vertical_name, created_at FROM compliances ORDER BY created_at DESC")
                rows = cursor.fetchall()
                return [ComplianceResponse(id=r['id'], form=r['form'], description=r['description'],
                                          due_date=r['due_date'], status=r['status'], priority=r['priority'],
                                          company_name=r['company_name'] if 'company_name' in r.keys() else None,
                                          vertical_name=r['vertical_name'] if 'vertical_name' in r.keys() else None,
                                          created_at=str(r['created_at'])) for r in rows]
            finally:
                conn.close()
        data = await asyncio.get_running_loop().run_in_executor(thread_pool, fetch)
        return CompliancesList(data=data, count=len(data))
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/compliances", response_model=ComplianceResponse)
async def create_compliance(request: ComplianceCreate, user: dict = Depends(require_session)):
    try:
        def insert():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn: raise RuntimeError("DB connection failed")
            cursor = get_pg_cursor(conn)
            try:
                cursor.execute(
                    "INSERT INTO compliances (form, description, due_date, status, priority, company_name, vertical_name) VALUES (%s, %s, %s, %s, %s, %s, %s) RETURNING id, form, description, due_date, status, priority, company_name, vertical_name, created_at",
                    (request.form, request.description, request.due_date, request.status, request.priority, request.company_name, request.vertical_name))
                row = cursor.fetchone()
                conn.commit()
                return ComplianceResponse(id=row['id'], form=row['form'], description=row['description'],
                                         due_date=row['due_date'], status=row['status'], priority=row['priority'],
                                         company_name=row['company_name'] if 'company_name' in row.keys() else None,
                                         vertical_name=row['vertical_name'] if 'vertical_name' in row.keys() else None,
                                         created_at=str(row['created_at']))
            finally:
                conn.close()
        return await asyncio.get_running_loop().run_in_executor(thread_pool, insert)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# --- Place Create Endpoint ---

@router.post("/places", response_model=PlaceResponse)
async def create_place(request: PlaceCreateRequest):
    """Create a new meeting place in PostgreSQL."""
    try:
        def insert():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn:
                raise RuntimeError("DB connection failed")
            cursor = get_pg_cursor(conn)
            try:
                # If this is set as default, unset other defaults
                if request.is_default:
                    cursor.execute("UPDATE places SET is_default = FALSE WHERE is_default = TRUE")

                cursor.execute(
                    "INSERT INTO places (name, address, is_default) VALUES (%s, %s, %s) RETURNING id, name, address, is_default, created_at",
                    (request.name, request.address, request.is_default)
                )
                row = cursor.fetchone()
                conn.commit()
                return PlaceResponse(
                    id=row['id'], name=row['name'], address=row['address'],
                    is_default=row['is_default'], created_at=str(row['created_at'])
                )
            finally:
                conn.close()
        return await asyncio.get_running_loop().run_in_executor(thread_pool, insert)
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e))
    except Exception as e:
        logger.error(f"Error creating place: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# --- Upload Custom Template Endpoint ---

@router.post("/upload-template")
async def upload_template(file: UploadFile = File(...)):
    """Upload a custom DOCX template for minutes generation."""
    if not file.filename or not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Only .docx files are supported")

    try:
        templates_dir = os.path.join(os.path.dirname(__file__), "..", "public", "templates")
        os.makedirs(templates_dir, exist_ok=True)

        # Generate a unique filename to avoid collisions
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_name = file.filename.replace(' ', '_')
        filename = f"custom_{timestamp}_{safe_name}"
        file_path = os.path.join(templates_dir, filename)

        content = await file.read()
        with open(file_path, "wb") as f:
            f.write(content)

        logger.info(f"Custom template uploaded: {filename}")
        return {"filename": filename, "message": "Template uploaded successfully"}
    except Exception as e:
        logger.error(f"Error uploading template: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to upload template: {str(e)}")


# --- Minutes Generation Endpoint ---

class MinutesGenerationRequest(BaseModel):
    template: str
    companyName: str
    meetingNumber: str = ""
    meetingType: str = ""
    meetingDay: str = ""
    meetingDate: str = ""
    meetingStartTime: str = ""
    meetingEndTime: str = ""
    meetingPlace: str = ""
    chairmanName: str = ""
    presentDirectors: List[Dict[str, Any]] = []
    inAttendance: List[Dict[str, Any]] = []
    companySecretary: str = ""
    previousMeetingDate: str = ""
    authorisedOfficer: str = ""
    quorum: str = ""
    concerns: str = ""
    declarations: str = ""
    auditorPaymentAmount: str = ""
    auditorPaymentWords: str = ""
    financialYear: str = ""
    agmNumber: str = ""
    agmDay: str = ""
    agmMonthName: str = ""
    agmDate: str = ""
    agmTime: str = ""
    agmPlace: str = ""
    recordingDate: str = ""
    signingDate: str = ""
    signingPlace: str = ""
    # Extended fields from FormBasedGenerator
    hasSection184Disclosure: bool = False
    section184Subject: str = ""
    section184Text: str = ""
    resolutions: str = ""
    customTemplateFilename: Optional[str] = None
    vertical_name: Optional[str] = "Energy"
    # TemplateRenderer sends directors as a list too
    directors: List[Dict[str, Any]] = []


@router.post("/generate-minutes")
async def generate_minutes(request: MinutesGenerationRequest):
    """Generate meeting minutes document from a DOCX template with placeholder replacement."""
    if not DOCX_AVAILABLE:
        raise HTTPException(status_code=500, detail="python-docx is not installed on the server")

    try:
        templates_dir = os.path.join(os.path.dirname(__file__), "..", "public", "templates")

        # Validate template name is provided
        if not request.template or not request.template.strip():
            raise HTTPException(status_code=400, detail="Template name is required. Please select a template before generating.")

        # Resolve the template path
        template_path = None
        if request.template == 'custom' and request.customTemplateFilename:
            template_path = os.path.join(templates_dir, request.customTemplateFilename)
            if not os.path.exists(template_path):
                raise HTTPException(status_code=400, detail=f"Custom template file '{request.customTemplateFilename}' not found on server. Please re-upload the template.")
        elif os.path.exists(os.path.join(templates_dir, request.template)):
            template_path = os.path.join(templates_dir, request.template)
        else:
            official_q_map = {
                'q1': '87. AGEL - BM - 28.04.2025.docx',
                'q2': '88. AGEL - BM - 28.07.2025.docx',
                'q3': '89. AGEL - BM - 28.10.2025.docx',
                'q4': '90. AGEL - BM - 23.01.2026.docx',
            }
            template_filename = official_q_map.get(request.template.lower(), f"{request.template.lower()}_meeting_template.docx")
            template_path = os.path.join(templates_dir, template_filename)

        # Final check: ensure the resolved template file exists
        if not os.path.exists(template_path):
            raise HTTPException(
                status_code=404,
                detail=f"Template '{request.template}' could not be resolved to an existing file. Available templates can be viewed on the Templates page."
            )

        def generate_document():
            doc = Document(template_path)

            # Merge presentDirectors and directors, then filter by status
            raw_directors = request.presentDirectors or request.directors or []
            seen_d_names = set()
            all_directors = []
            for d in raw_directors:
                d_n = (d.get('name') or '').strip()
                if d_n and d_n.lower() not in seen_d_names:
                    seen_d_names.add(d_n.lower())
                    all_directors.append(d)
            if not all_directors and raw_directors:
                all_directors = raw_directors

            present_directors = [d for d in all_directors if d.get('status', 'Present') != 'Leave of Absence']
            absent_directors = [d for d in all_directors if d.get('status') == 'Leave of Absence']
            directors = present_directors if present_directors else all_directors

            # Get non-chairman director for signature tables
            non_chairman = [d for d in all_directors if d.get('name') != request.chairmanName]
            sig_director = non_chairman[0] if non_chairman else (all_directors[0] if all_directors else {'name': '', 'din': ''})

            # Auto-pull real Director Disclosure DTOs (MBP-1 & DIR-8) via Service Repository
            mbp1_disclosures_list = []
            dir8_declarations_list = []

            for d in all_directors:
                d_din = d.get('din', '').strip()
                if d_din:
                    try:
                        from routes.directors_disclosure import get_director_disclosure_dto
                        dto = get_director_disclosure_dto(d_din)
                        if dto.dir8_confirmation:
                            dir8_declarations_list.append(dto.dir8_confirmation)
                        if dto.mbp1_disclosure_text:
                            mbp1_disclosures_list.append(dto.mbp1_disclosure_text)
                    except Exception as ex:
                        logger.warning(f"Failed to query director disclosure DTO for DIN {d_din}: {ex}")

            sec_184_text = request.section184Text or "\n".join(mbp1_disclosures_list) or "Notices of Interest in Form MBP-1 pursuant to Section 184(1) of the Companies Act, 2013 received from attending directors were placed before the Board and noted."
            sec_164_text = "\n".join(dir8_declarations_list) or "Declarations in Form DIR-8 in terms of Section 164(2) of the Companies Act, 2013 received from attending directors confirming no disqualification were taken on record."

            # Dynamic placeholders dictionary
            placeholders = {
                '[No. of Meeting]': request.meetingNumber,
                '[Type of Meeting]': request.meetingType,
                '[Name of Company]': request.companyName,
                '[Day of Meeting]': request.meetingDay,
                '[Date of Meeting]': request.meetingDate,
                '[Time: COMMENCED AT]': request.meetingStartTime,
                '[Time: CONCLUDED AT]': request.meetingEndTime,
                '[Place of Meeting]': request.meetingPlace,
                '[Chairman]': request.chairmanName,
                '[Manual]': request.chairmanName,
                '[Director]': sig_director.get('name', ''),
                '[Date-previous-meeting]': request.previousMeetingDate,
                '[amount]': request.auditorPaymentAmount,
                '[Amount-in-words]': request.auditorPaymentWords,
                '[amount-in-words]': request.auditorPaymentWords,
                '[Year]': request.financialYear,
                '[year]': request.financialYear,
                '[YEAR]': request.financialYear,
                '[start-year]': request.financialYear.split('-')[0] if '-' in request.financialYear else request.financialYear,
                '[end-year]': request.financialYear.split('-')[1] if '-' in request.financialYear else (str(int(request.financialYear) + 1) if request.financialYear.isdigit() else ''),
                '[Day-of-meeting]': request.meetingDay,
                '[Month-of-meeting]': request.agmMonthName,
                '[TIME]': request.agmTime,
                '[Office-address]': request.agmPlace,
                '[Date-of-Recording]': request.recordingDate,
                '[Date-of-signing]': request.signingDate,
                '[Place of signing]': request.signingPlace,
                '[Officer]': request.companySecretary or request.authorisedOfficer,
                '[Resolutions]': request.resolutions,
                '[resolutions]': request.resolutions,
                '[RESOLUTIONS]': request.resolutions,
                '[Section-184-Disclosures]': sec_184_text,
                '[Section-164-Declarations]': sec_164_text,
            }

            # Explicit Sample Place & Address Replacement Rules
            if request.meetingPlace and request.meetingPlace.strip():
                placeholders['Plot No. 83, Sector 32, Institutional Area, Gurgaon, Haryana 122001'] = request.meetingPlace
                placeholders['Plot No. 83, Sector 32, Institutional Area, Gurgaon, Haryana-122001'] = request.meetingPlace

            if request.signingPlace and request.signingPlace.strip():
                placeholders['hyderabad'] = request.signingPlace
                placeholders['Hyderabad'] = request.signingPlace
                placeholders['HYDERABAD'] = request.signingPlace.upper()

            # ════════════════════════════════════════════════════════════
            # SMART AUTO-DETECTION ENGINE
            # Instead of manually listing every sample string from every
            # template, we scan the document text at generation time
            # and dynamically discover company names, director names,
            # dates, and meeting numbers — then replace them with the
            # user's actual form data.
            # ════════════════════════════════════════════════════════════

            import re as _re

            # Gather the full text of the template for pattern scanning
            _full_text = "\n".join(p.text for p in doc.paragraphs)
            for _tbl in doc.tables:
                for _row in _tbl.rows:
                    for _cell in _row.cells:
                        _full_text += "\n" + "\n".join(p.text for p in _cell.paragraphs)

            # ── 1. AUTO-DETECT & REPLACE COMPANY NAMES ──────────────
            # Find all "ALL CAPS ... LIMITED/LTD" strings in the template
            if request.companyName and request.companyName.strip():
                upper_comp = request.companyName.upper()
                # Match company name patterns (10+ chars ending with LIMITED or LTD)
                _found_companies = _re.findall(
                    r'[A-Z][A-Z\s\(\)\-\']{8,}(?:LIMITED|LTD\.?)',
                    _full_text
                )
                _found_companies = list(set(c.strip() for c in _found_companies if len(c.strip()) > 10))

                # Filter: only replace companies that look like Adani entities or
                # match the template's BU company. Skip generic bank names, legal refs, etc.
                _skip_keywords = ['BANK ', 'STOCK EXCHANGE', 'SECURITIES', 'RESERVE BANK',
                                  'REGISTRAR', 'INSURANCE', 'SEBI', 'NSDL', 'CDSL']
                for _comp in _found_companies:
                    # Skip bank names and regulatory bodies
                    if any(sk in _comp for sk in _skip_keywords):
                        continue
                    # Skip if it's already the user's company name
                    if _comp == upper_comp:
                        continue
                    # Replace the sample company name with user's company
                    placeholders[_comp] = upper_comp

            # ── 2. AUTO-DETECT & REPLACE TITLE LINE (MINUTES OF THE...) ──
            # The title line follows the pattern:
            # "MINUTES OF THE <ORDINAL> MEETING OF ... <COMPANY NAME>"
            _title_pattern = _re.compile(
                r'(MINUTES\s+OF\s+THE\s+)([A-Z\s\-]+?)'
                r'(\s+MEETING\s+OF\s+(?:THE\s+)?(?:BOARD\s+OF\s+DIRECTORS|AUDIT\s+COMMITTEE|NOMINATION\s+AND\s+REMUNERATION\s+COMMITTEE|STAKEHOLDERS\s+RELATIONSHIP\s+COMMITTEE|CORPORATE\s+SOCIAL\s+RESPONSIBILITY\s+COMMITTEE|RISK\s+MANAGEMENT\s+COMMITTEE)\s+OF\s+)'
                r'(.+?)(?=\s*$|\s*\n)',
                _re.MULTILINE
            )
            _title_matches = _title_pattern.findall(_full_text)
            for _match in _title_matches:
                _old_ordinal = _match[1].strip()
                _old_company = _match[3].strip()
                # Replace the ordinal meeting number in the title
                if request.meetingNumber and request.meetingNumber.strip():
                    if _old_ordinal and _old_ordinal not in placeholders:
                        placeholders[_old_ordinal] = request.meetingNumber.upper()
                # Replace the company name in the title
                if request.companyName and request.companyName.strip():
                    if _old_company and _old_company not in placeholders:
                        placeholders[_old_company] = request.companyName.upper()

            # ── 3. AUTO-DETECT & REPLACE MEETING NUMBER ORDINALS ─────
            if request.meetingNumber and request.meetingNumber.strip():
                m_num = request.meetingNumber
                # Find word-ordinals like "EIGHTY SEVENTH", "FIFTY NINTH" etc.
                _ordinal_pattern = _re.compile(
                    r'\b('
                    r'(?:TWENTY|THIRTY|FORTY|FIFTY|SIXTY|SEVENTY|EIGHTY|NINETY)'
                    r'[\s\-]?'
                    r'(?:FIRST|SECOND|THIRD|FOURTH|FIFTH|SIXTH|SEVENTH|EIGHTH|NINTH)'
                    r'|FIRST|SECOND|THIRD|FOURTH|FIFTH|SIXTH|SEVENTH|EIGHTH|NINTH|TENTH'
                    r'|ELEVENTH|TWELFTH|THIRTEENTH|FOURTEENTH|FIFTEENTH|SIXTEENTH'
                    r'|SEVENTEENTH|EIGHTEENTH|NINETEENTH|TWENTIETH|THIRTIETH|FORTIETH'
                    r'|FIFTIETH|SIXTIETH|SEVENTIETH|EIGHTIETH|NINETIETH|HUNDREDTH'
                    r')\b'
                )
                _found_ordinals = set(_ordinal_pattern.findall(_full_text))
                for _ord in _found_ordinals:
                    # Only replace ordinals that appear in the TITLE context
                    # (i.e. near "MEETING OF"), not random uses like "FOURTH quarter"
                    _context_check = _re.search(
                        _re.escape(_ord) + r'\s+MEETING',
                        _full_text
                    )
                    if _context_check and _ord not in placeholders:
                        placeholders[_ord] = m_num

                # Also find numeric ordinals that reference previous meetings
                # e.g. "86TH", "87TH" (previous meeting number references)
                _num_ord_pattern = _re.compile(r'\b(\d+)(?:ST|ND|RD|TH)\s+MEETING\b', _re.IGNORECASE)
                _found_num_ords = _num_ord_pattern.findall(_full_text)
                for _n in set(_found_num_ords):
                    _full_match = _re.search(r'\b(' + _n + r'(?:ST|ND|RD|TH))\s+MEETING\b', _full_text, _re.IGNORECASE)
                    if _full_match:
                        _old_num_ord = _full_match.group(1)
                        if _old_num_ord not in placeholders:
                            placeholders[_old_num_ord] = m_num

            # ── 4. AUTO-DETECT & REPLACE DIRECTOR/PERSON NAMES ───────
            # Find all "Mr./Mrs./Ms./Shri <Name>" patterns in the template
            _person_pattern = _re.compile(
                r'(?:Mr\.|Mrs\.|Ms\.|Shri|Smt\.)\s+[A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+){1,3}'
            )
            _found_persons = list(set(_person_pattern.findall(_full_text)))

            # Build a set of the user's entered director names for matching
            _user_director_names = set()
            for _d in all_directors:
                _dn = _d.get('name', '').strip()
                if _dn:
                    _user_director_names.add(_dn)
            if request.chairmanName and request.chairmanName.strip():
                _user_director_names.add(request.chairmanName)

            # Only replace sample person names that are NOT in the user's director list
            # (If the user actually has a director with the same name, don't replace it)
            _sample_persons = [p for p in _found_persons if p not in _user_director_names]

            # Don't auto-replace person names with a single value — that would make
            # every name in the doc the same. Instead, we only replace the CHAIRMAN
            # and the SIGNING DIRECTOR names. The rest of the attendance list is
            # handled by the [Dir-name]/[Din-num] placeholder system.
            # But we DO need to replace the specific chairman/signing names embedded
            # in the template's body text (resolutions, authorizations, etc.)
            if request.chairmanName and request.chairmanName.strip():
                # Find the first person mentioned in the template — that's typically
                # the chairman in the original sample
                # Look for the chairman pattern: appears right after "IN THE CHAIR" or
                # near "Chairman" context
                _chair_context = _re.search(
                    r'(?:Mr\.|Mrs\.|Ms\.)\s+[A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+){1,3}'
                    r'(?=.*?(?:Chair|chair|CHAIR|presided))',
                    _full_text, _re.DOTALL
                )
                if _chair_context:
                    _old_chair = _chair_context.group(0).strip()
                    if _old_chair not in _user_director_names and _old_chair not in placeholders:
                        placeholders[_old_chair] = request.chairmanName
                        # Also add without honorific
                        _bare_old = _re.sub(r'^(?:Mr\.|Mrs\.|Ms\.)\s+', '', _old_chair)
                        if _bare_old and _bare_old not in placeholders:
                            _bare_new = _re.sub(r'^(?:Mr\.|Mrs\.|Ms\.)\s+', '', request.chairmanName)
                            placeholders[_bare_old] = _bare_new if _bare_new else request.chairmanName

            # ── 5. AUTO-DETECT & REPLACE KEY DATES ────────────────────
            # We only replace the PRIMARY meeting date (the one in the title/header)
            # and the previous meeting date. We do NOT replace random dates in
            # resolution bodies (like agreement dates, appointment dates, etc.)
            # because those are part of the legal content.
            if request.meetingDate and request.meetingDate.strip():
                m_date = request.meetingDate
                try:
                    dt_obj = datetime.strptime(m_date, '%Y-%m-%d')
                    day_num = dt_obj.day
                    ord_suffix = 'th' if 11 <= day_num <= 13 else {1: 'st', 2: 'nd', 3: 'rd'}.get(day_num % 10, 'th')
                    formatted_date_upper = f"{dt_obj.day}{ord_suffix.upper()} {dt_obj.strftime('%B %Y').upper()}"
                    formatted_date_title = f"{dt_obj.day}{ord_suffix} {dt_obj.strftime('%B, %Y')}"
                except Exception:
                    formatted_date_upper = m_date.upper()
                    formatted_date_title = m_date

                # Find the primary meeting date in the template
                # It's typically the date that appears in the title or right after
                # "held on" / "HELD ON" or matches the filename date
                # Strategy: extract the date from the template filename
                _template_basename = os.path.basename(template_path)
                _fname_date_match = _re.search(r'(\d{2})\.(\d{2})\.(\d{4})', _template_basename)
                if _fname_date_match:
                    _fd, _fm, _fy = _fname_date_match.groups()
                    _fname_date_str = f"{_fd}.{_fm}.{_fy}"
                    # Replace the DD.MM.YYYY format date
                    placeholders[_fname_date_str] = formatted_date_upper

                    # Build the spelled-out version of this date for replacement
                    try:
                        _tpl_dt = datetime(int(_fy), int(_fm), int(_fd))
                        _tpl_day = _tpl_dt.day
                        _tpl_suffix = 'th' if 11 <= _tpl_day <= 13 else {1: 'st', 2: 'nd', 3: 'rd'}.get(_tpl_day % 10, 'th')
                        _month_name = _tpl_dt.strftime('%B')
                        _year = _tpl_dt.strftime('%Y')

                        # All common date format variations of the template's primary date
                        _date_variants = [
                            f"{_tpl_day}{_tpl_suffix.upper()} {_month_name.upper()}, {_year}",
                            f"{_tpl_day}{_tpl_suffix.upper()} {_month_name.upper()} {_year}",
                            f"{_tpl_day}{_tpl_suffix} {_month_name}, {_year}",
                            f"{_tpl_day}{_tpl_suffix} {_month_name} {_year}",
                            f"{_fd}{_tpl_suffix.upper()} {_month_name.upper()}, {_year}",
                            f"{_fd}{_tpl_suffix.upper()} {_month_name.upper()} {_year}",
                            f"{_fd}{_tpl_suffix} {_month_name}, {_year}",
                            f"{_fd}{_tpl_suffix} {_month_name} {_year}",
                        ]
                        for _dv in _date_variants:
                            if _dv in _full_text and _dv not in placeholders:
                                placeholders[_dv] = formatted_date_upper
                    except Exception:
                        pass

                # Also find the signing/recording date pattern (DD.MM.YYYY after the main date)
                _all_dotdates = _re.findall(r'\b(\d{2}\.\d{2}\.\d{4})\b', _full_text)
                for _dd in set(_all_dotdates):
                    # Only replace if it's the signing/conclusion date, not random dates
                    # The signing date typically appears near the end of the document
                    _dd_pos = _full_text.rfind(_dd)
                    _doc_len = len(_full_text)
                    if _dd_pos > _doc_len * 0.8:  # Last 20% of the document
                        if request.signingDate and request.signingDate.strip():
                            try:
                                _sd_obj = datetime.strptime(request.signingDate, '%Y-%m-%d')
                                _sd_day = _sd_obj.day
                                _sd_suf = 'th' if 11 <= _sd_day <= 13 else {1: 'st', 2: 'nd', 3: 'rd'}.get(_sd_day % 10, 'th')
                                _sd_formatted = f"{_sd_day}{_sd_suf.upper()} {_sd_obj.strftime('%B %Y').upper()}"
                                placeholders[_dd] = _sd_formatted
                            except Exception:
                                placeholders[_dd] = request.signingDate

            # ── Style-preserving run-level replacement ──────────────────
            # Instead of setting para.text (which destroys all formatting),
            # we replace text inside individual runs so that each run's
            # font size, bold, italic, underline, color, etc. are kept.

            def _replace_in_runs(runs, placeholder, value):
                """Replace *placeholder* with *value* across a list of runs,
                preserving all run-level formatting.

                Case 1 – placeholder sits inside a single run:
                    Simply do run.text = run.text.replace(...)

                Case 2 – placeholder is split across consecutive runs:
                    Concatenate adjacent run texts, find the placeholder span,
                    put the replacement in the first participating run (keeping
                    its formatting), and clear the remaining runs.
                """
                val = str(value)

                # --- Single-run replacement (most common) ---
                for run in runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, val)

                # --- Cross-run replacement (placeholder split by Word) ---
                # Word sometimes splits a single placeholder across 2-3 runs
                # e.g. run1="[Name of ", run2="Company]"
                full = "".join(r.text for r in runs)
                if placeholder not in full:
                    return  # nothing left to do

                while placeholder in "".join(r.text for r in runs):
                    combined = ""
                    start_idx = None
                    end_idx = None
                    for i, run in enumerate(runs):
                        combined += run.text
                        if start_idx is None and placeholder[:1] in run.text:
                            # potential start — re-check from this run
                            partial = "".join(r.text for r in runs[i:])
                            if partial.startswith(placeholder) or placeholder in "".join(r.text for r in runs[:i+1]):
                                pass  # continue scanning
                        if placeholder in combined:
                            end_idx = i
                            # walk backwards to find where placeholder text begins
                            back = ""
                            for j in range(i, -1, -1):
                                back = runs[j].text + back
                                if placeholder in back:
                                    start_idx = j
                                    break
                            break

                    if start_idx is None or end_idx is None:
                        break  # safety – avoid infinite loop

                    # Merge text of participating runs
                    merged = "".join(runs[k].text for k in range(start_idx, end_idx + 1))
                    replaced = merged.replace(placeholder, val, 1)

                    # Put replaced text in the FIRST run (keeps its formatting)
                    runs[start_idx].text = replaced
                    # Clear the rest
                    for k in range(start_idx + 1, end_idx + 1):
                        runs[k].text = ""

            def replace_in_paragraph(para):
                """Replace all placeholders in a paragraph while preserving
                each run's original font properties."""
                for placeholder, value in placeholders.items():
                    if not value:
                        continue
                    if placeholder in para.text:
                        _replace_in_runs(list(para.runs), placeholder, value)

            def replace_in_cell_preserving_style(cell):
                """Replace placeholders inside every paragraph of a table cell,
                preserving run-level formatting."""
                for para in cell.paragraphs:
                    replace_in_paragraph(para)

            # Replace in document body paragraphs
            for para in doc.paragraphs:
                replace_in_paragraph(para)

            # Replace in document body tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        replace_in_cell_preserving_style(cell)

            # Replace in headers and footers
            for section in doc.sections:
                for para in section.header.paragraphs:
                    replace_in_paragraph(para)
                for table in section.header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            replace_in_cell_preserving_style(cell)
                for para in section.footer.paragraphs:
                    replace_in_paragraph(para)
                for table in section.footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            replace_in_cell_preserving_style(cell)

            # Smart [Dir-name] / [Din-num] & Attendance List replacement — each occurrence gets filled
            # Uses run-level replacement to preserve formatting
            if all_directors:
                director_index = 0
                total_directors = len(all_directors)
                for para in doc.paragraphs:
                    while '[Dir-name]' in para.text or '[Din-num]' in para.text:
                        if director_index < total_directors:
                            d = all_directors[director_index]
                            if '[Dir-name]' in para.text:
                                _replace_in_runs(list(para.runs), '[Dir-name]', d.get('name', ''))
                            if '[Din-num]' in para.text:
                                _replace_in_runs(list(para.runs), '[Din-num]', d.get('din', ''))
                            director_index += 1
                        else:
                            _replace_in_runs(list(para.runs), '[Dir-name]', '')
                            _replace_in_runs(list(para.runs), '[Din-num]', '')
                            break

            # ── 6. SMART TRANSFORM FOR HARDCODED TEXT (DATES, TIMES, ATTENDANCE, RESOLUTIONS) ──
            comp_name_upper = request.companyName.upper() if request.companyName else ""
            day_upper = request.meetingDay.upper() if request.meetingDay else ""
            
            # Format start time e.g. "10:11" -> "10:11 AM" & "10.11 A.M."
            start_time_str = request.meetingStartTime or ""
            start_time_dot_str = ""
            if start_time_str:
                try:
                    clean_time = start_time_str.strip()
                    if ":" in clean_time or "." in clean_time:
                        delim = ":" if ":" in clean_time else "."
                        parts = clean_time.split(delim)
                        hh = int(parts[0])
                        mm = parts[1][:2]
                        ampm = "AM" if hh < 12 else "PM"
                        ampm_dot = "A.M." if hh < 12 else "P.M."
                        hh12 = hh if hh <= 12 else hh - 12
                        if hh12 == 0: hh12 = 12
                        start_time_str = f"{hh12:02d}:{mm} {ampm}"
                        start_time_dot_str = f"{hh12:02d}.{mm} {ampm_dot}"
                except Exception:
                    pass

            # Format end time e.g. "11:11" -> "11:11 AM" & "11.11 A.M."
            end_time_str = request.meetingEndTime or ""
            end_time_dot_str = ""
            if end_time_str:
                try:
                    clean_etime = end_time_str.strip()
                    if ":" in clean_etime or "." in clean_etime:
                        delim = ":" if ":" in clean_etime else "."
                        parts = clean_etime.split(delim)
                        hh = int(parts[0])
                        mm = parts[1][:2]
                        ampm = "AM" if hh < 12 else "PM"
                        ampm_dot = "A.M." if hh < 12 else "P.M."
                        hh12 = hh if hh <= 12 else hh - 12
                        if hh12 == 0: hh12 = 12
                        end_time_str = f"{hh12:02d}:{mm} {ampm}"
                        end_time_dot_str = f"{hh12:02d}.{mm} {ampm_dot}"
                except Exception:
                    pass

            def format_date_ordinal(date_str):
                if not date_str:
                    return ""
                try:
                    dt = datetime.strptime(date_str[:10], "%Y-%m-%d")
                    day = dt.day
                    suffix = "TH" if 11 <= day <= 13 else {1: "ST", 2: "ND", 3: "RD"}.get(day % 10, "TH")
                    month_name = dt.strftime("%B").upper()
                    return f"{day}{suffix} {month_name} {dt.year}"
                except Exception:
                    return date_str

            m_date_formatted = format_date_ordinal(request.meetingDate)
            rec_date_formatted = format_date_ordinal(request.recordingDate)
            sign_date_formatted = format_date_ordinal(request.signingDate)
            prev_date_formatted = format_date_ordinal(request.previousMeetingDate) or m_date_formatted

            # Build Global Director List from Website Input
            directors = all_directors

            def transform_text(text):
                if not text or not text.strip():
                    return text
                
                # 0. Replace bracket placeholders
                for pk, pv in placeholders.items():
                    if pk in text and pv:
                        text = text.replace(pk, str(pv))

                # 1. Company Name Replacements (both exact template names & generic pattern)
                if comp_name_upper:
                    for old_c in [
                        "ADANI GREEN ENERGY LIMITED", "Adani Green Energy Limited",
                        "Adani Green Energy (UP) Limited", "ADANI GREEN ENERGY (UP) LIMITED",
                        "ADANI RENEWABLE ENERGY HOLDING FOUR LIMITED", "Adani Renewable Energy Holding Four Limited",
                        "AGEL", "AGE(UP)L", "AGE25BL"
                    ]:
                        if old_c in text:
                            text = text.replace(old_c, comp_name_upper if old_c.isupper() else request.companyName)
                    text = re.sub(r'\b[A-Z0-9\s\(\)\&\.\-]{3,60}\s+(?:LIMITED|PVT\s+LTD|PRIVATE\s+LIMITED|LTD)\b', comp_name_upper, text, flags=re.IGNORECASE)
                
                # 2. Meeting Place & Address Replacement
                if request.meetingPlace and request.meetingPlace.strip():
                    sample_places = [
                        "Adani Corporate House, Shantigram, Near Vaishno Devi Circle, S. G. Highway, Khodiyar, Ahmedabad – 382 421, Gujarat, India",
                        "Adani Corporate House, Shantigram, Near Vaishno Devi Circle, S. G. Highway, Khodiyar, Ahmedabad – 382 421, Gujarat",
                        "Adani Corporate House, Shantigram, Near Vaishno Devi Circle, S. G. Highway, Khodiyar, Ahmedabad - 382 421",
                        "Plot No. 83, Sector 32, Institutional Area, Gurgaon, Haryana 122001",
                        "Plot No. 83, Sector 32, Institutional Area, Gurgaon, Haryana-122001"
                    ]
                    for sp in sample_places:
                        if sp in text:
                            text = text.replace(sp, request.meetingPlace)

                # 3. Company Secretary Name Replacement
                if request.companySecretary and request.companySecretary.strip():
                    for sample_cs in ["Kuntal Pandya", "Kuntal Chandya", "Chandan Lakhwani"]:
                        if sample_cs in text:
                            text = text.replace(sample_cs, request.companySecretary)

                # 4. Replace day names in headings (MONDAY..SUNDAY)
                if day_upper:
                    text = re.sub(r'\b(MONDAY|TUESDAY|WEDNESDAY|THURSDAY|FRIDAY|SATURDAY|SUNDAY)\b', day_upper, text, flags=re.IGNORECASE)
                
                # 5. Replace ALL internal meeting & financial dates (e.g. 22ND JULY 2025, 30TH SEPTEMBER 2025)
                if m_date_formatted:
                    text = re.sub(r'\b\d{1,2}(?:ST|ND|RD|TH)\s+(?:JANUARY|FEBRUARY|MARCH|APRIL|MAY|JUNE|JULY|AUGUST|SEPTEMBER|OCTOBER|NOVEMBER|DECEMBER)(?:,\s*|\s+)\d{4}\b', m_date_formatted, text, flags=re.IGNORECASE)

                # 6. Replace Commencement Time
                if start_time_str:
                    target_start = start_time_dot_str if ("." in text and ("P.M." in text or "A.M." in text or "p.m." in text)) else start_time_str
                    text = re.sub(r'(?:commenced|held)\s+at\s+\d{1,2}[\.:]\d{2}\s*(?:AM|PM|a\.m\.|p\.m\.|P\.M\.|A\.M\.)', f'commenced at {target_start}', text, flags=re.IGNORECASE)
                    text = re.sub(r'AT\s+\d{1,2}[\.:]\d{2}\s*(?:AM|PM|a\.m\.|p\.m\.|P\.M\.|A\.M\.)\s+AT', f'AT {target_start} AT', text, flags=re.IGNORECASE)

                # 7. Replace Conclusion Time (Vote of thanks)
                if end_time_str:
                    target_end = end_time_dot_str if ("." in text and ("P.M." in text or "A.M." in text or "p.m." in text)) else end_time_str
                    text = re.sub(r'(?:concluded|thanks\s+to\s+the\s+chair)\s+at\s+\d{1,2}[\.:]\d{2}\s*(?:AM|PM|a\.m\.|p\.m\.|P\.M\.|A\.M\.)', f'concluded with a vote of thanks to the chair at {target_end}', text, flags=re.IGNORECASE)

                # 8. Replace sample signing places
                if request.signingPlace:
                    text = re.sub(r'\b[A-Z][a-z]{2,20}\b(?=\s+CHAIRMAN|\s*Date)', request.signingPlace, text)

                # 9. Handle Date of Entry & Date of Signing specifically
                if "Date of entry" in text or "Date of Recording" in text or "Date of Entry" in text:
                    if rec_date_formatted:
                        text = re.sub(r'(Date of entry|Date of Recording|Date of Entry):?\s*.*', f'Date of entry:\t{rec_date_formatted}', text, flags=re.IGNORECASE)

                if "Date of signing" in text or "Date of Signing" in text:
                    if sign_date_formatted:
                        text = re.sub(r'(Date of signing|Date of Signing):?\s*.*', f'Date of signing:\t{sign_date_formatted}', text, flags=re.IGNORECASE)
                    
                return text

            for para in doc.paragraphs:
                if para.text:
                    new_text = transform_text(para.text)
                    if new_text != para.text:
                        para.text = new_text

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if para.text:
                                new_text = transform_text(para.text)
                                if new_text != para.text:
                                    para.text = new_text

            for section in doc.sections:
                for hp in section.header.paragraphs:
                    if hp.text:
                        hp.text = transform_text(hp.text)
                for ht in section.header.tables:
                    for r in ht.rows:
                        for c in r.cells:
                            for p in c.paragraphs:
                                if p.text:
                                    p.text = transform_text(p.text)
                for fp in section.footer.paragraphs:
                    if fp.text:
                        fp.text = transform_text(fp.text)

            # Ensure Chairman occupied the chair text matches request.chairmanName
            if request.chairmanName:
                for para in doc.paragraphs:
                    if "occupied the Chair" in para.text or "occupied the chair" in para.text:
                        para.text = re.sub(r'^.*?\boccupied the [Cc]hair', f"{request.chairmanName} occupied the Chair", para.text)

            # Rebuild Attendance Tables & Director lists dynamically
            if directors:
                # 1. Update Word Tables containing director rows
                for table in doc.tables:
                    headers = [cell.text.lower() for cell in table.rows[0].cells] if table.rows else []
                    if any("director" in h or "din" in h or "name" in h for h in headers):
                        row_idx = 1
                        for d_idx, d in enumerate(directors):
                            d_name = d.get('name', '')
                            d_din = d.get('din', '')
                            d_role = "Chairman" if d_name == request.chairmanName else (d.get('designation') or d.get('role') or "Director")

                            if row_idx < len(table.rows):
                                row = table.rows[row_idx]
                                if len(row.cells) >= 3:
                                    row.cells[0].text = str(d_idx + 1)
                                    row.cells[1].text = d_name
                                    row.cells[2].text = d_din or d_role
                                    if len(row.cells) >= 4:
                                        row.cells[3].text = d_role
                                elif len(row.cells) >= 2:
                                    row.cells[0].text = d_name
                                    row.cells[1].text = d_din or d_role
                                row_idx += 1
                            else:
                                new_row = table.add_row()
                                if len(new_row.cells) >= 3:
                                    new_row.cells[0].text = str(d_idx + 1)
                                    new_row.cells[1].text = d_name
                                    new_row.cells[2].text = d_din or d_role
                                    if len(new_row.cells) >= 4:
                                        new_row.cells[3].text = d_role
                                elif len(new_row.cells) >= 2:
                                    new_row.cells[0].text = d_name
                                    new_row.cells[1].text = d_din or d_role
                                row_idx += 1

                        while row_idx < len(table.rows):
                            for cell in table.rows[row_idx].cells:
                                cell.text = ""
                            row_idx += 1

                # 2. Update Paragraph-based attendance lists & remove extra sample paragraphs
                in_present_section = False
                dir_paras = []
                next_para = None

                for para in doc.paragraphs:
                    t = para.text.strip()
                    if "where the following directors were present" in t.lower() or "present physically" in t.lower():
                        in_present_section = True
                        continue
                    if in_present_section:
                        if "invitee" in t.lower() or "chairman" in t.lower() or "leave of absence" in t.lower():
                            in_present_section = False
                            next_para = para
                            break
                        elif re.match(r'^\d+\.\s*', t):
                            dir_paras.append(para)

                existing_count = len(dir_paras)
                for idx in range(max(len(present_directors), existing_count)):
                    if idx < len(present_directors):
                        d = present_directors[idx]
                        d_name = d.get('name', '')
                        d_role = "Chairman" if d_name == request.chairmanName else (d.get('designation') or d.get('role') or "Director")
                        formatted_line = f"{idx + 1}.\t{d_name}\t\t-\t{d_role}"
                        if idx < existing_count:
                            target_p = dir_paras[idx]
                            if target_p.runs:
                                first_run = target_p.runs[0]
                                f_name = first_run.font.name
                                f_size = first_run.font.size
                                is_b = first_run.bold
                                is_i = first_run.italic
                                for r in target_p.runs[1:]:
                                    try:
                                        r._element.getparent().remove(r._element)
                                    except Exception:
                                        pass
                                first_run.text = formatted_line
                                if f_name: first_run.font.name = f_name
                                if f_size: first_run.font.size = f_size
                                if is_b is not None: first_run.bold = is_b
                                if is_i is not None: first_run.italic = is_i
                            else:
                                target_p.text = formatted_line
                        else:
                            if next_para:
                                next_para.insert_paragraph_before(formatted_line)
                            else:
                                doc.add_paragraph(formatted_line)
                    else:
                        if idx < existing_count:
                            try:
                                p_elem = dir_paras[idx]._element
                                p_elem.getparent().remove(p_elem)
                            except Exception:
                                dir_paras[idx].text = ""

                # 3. Dynamic Leave of Absence Section Update
                if absent_directors:
                    absent_names = []
                    for d in absent_directors:
                        n = d.get('name', '').strip()
                        if n:
                            if not n.startswith(('Mr.', 'Mrs.', 'Ms.', 'Dr.')):
                                n = f"Mr. {n}"
                            absent_names.append(n)
                    if absent_names:
                        if len(absent_names) == 1:
                            absent_str = absent_names[0]
                        else:
                            absent_str = ", ".join(absent_names[:-1]) + " and " + absent_names[-1]
                        loa_text = f"Leaves of absence were granted to {absent_str}, who couldn't make it convenient to attend the meeting."

                        replaced_loa = False
                        for para in doc.paragraphs:
                            p_lower = para.text.lower()
                            if "leave" in p_lower and "absence" in p_lower:
                                if re.search(r'\b[Ll]eaves?\s+of\s+absence\s+(?:was|were)\s+granted\b.*', para.text):
                                    para.text = loa_text
                                    replaced_loa = True
                                elif "leaves of absence" in p_lower and not replaced_loa:
                                    para.text = f"LEAVES OF ABSENCE\n\n{loa_text}"
                                    replaced_loa = True
                else:
                    for para in doc.paragraphs:
                        if re.search(r'\b[Ll]eaves?\s+of\s+absence\s+(?:was|were)\s+granted\b.*', para.text):
                            para.text = "Leaves of absence: Nil."

            # 3. Dynamic Resolutions Replacement (structure-preserving: text + tables)
            if request.resolutions and request.resolutions.strip():
                anchor = None
                for para in doc.paragraphs:
                    if "[Resolutions]" in para.text or "[resolutions]" in para.text or "[RESOLUTIONS]" in para.text:
                        para.text = para.text.replace("[Resolutions]", "").replace("[resolutions]", "").replace("[RESOLUTIONS]", "")
                        anchor = para
                        break

                if anchor is None:
                    doc.add_paragraph("\nMEETING RESOLUTIONS & BUSINESS TRANSACTED:\n")
                    render_resolutions_into_doc(doc, request.resolutions)
                else:
                    render_resolutions_into_doc(doc, request.resolutions, anchor_para=anchor)

            # Save generated document (kept separate from templates for production)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"meeting_minutes_{request.template}_{timestamp}.docx"
            generated_dir = os.path.join(os.path.dirname(__file__), "..", "public", "generated")
            os.makedirs(generated_dir, exist_ok=True)
            output_path = os.path.join(generated_dir, filename)
            doc.save(output_path)

            # BRD Req #7: Save to Structured Path Repository
            # Hierarchy: Vertical (BU) -> Company Name -> Meeting -> Type of Meeting -> Date(year)
            def _clean_str(s: str) -> str:
                return "".join(c for c in s if c.isalnum() or c in (' ', '_', '-')).strip().replace(' ', '_')

            v_name = _clean_str(request.vertical_name or "Energy")
            c_name = _clean_str(request.companyName or "General_Company")
            m_name = _clean_str(f"Meeting_{request.meetingNumber}") if request.meetingNumber else "Meeting"
            t_name = _clean_str(request.meetingType or "Board_Meeting")
            y_name = request.meetingDate.split('-')[0] if (request.meetingDate and '-' in request.meetingDate) else str(datetime.now().year)

            structured_dir = os.path.join(
                os.path.dirname(__file__), "..", "public", "repository",
                v_name, c_name, m_name, t_name, y_name
            )
            os.makedirs(structured_dir, exist_ok=True)
            doc.save(os.path.join(structured_dir, filename))
            logger.info(f"Saved copy to structured directory: {structured_dir}")

            return filename, output_path

        loop = asyncio.get_running_loop()
        filename, output_path = await loop.run_in_executor(thread_pool, generate_document)

        # Record in PostgreSQL history
        try:
            def record_history():
                conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
                if not conn:
                    return
                cursor = get_pg_cursor(conn)
                try:
                    cursor.execute(
                        """
                        INSERT INTO generated_minutes
                            (company_name, meeting_type, meeting_date, file_path, meeting_number, status)
                        VALUES (%s, %s, %s, %s, %s, %s)
                        RETURNING id
                        """,
                        (
                            request.companyName,
                            request.meetingType,
                            request.meetingDate,
                            filename,
                            request.meetingNumber or "",
                            "draft",
                        ),
                    )
                    row = cursor.fetchone()
                    minutes_id = row["id"] if row else None

                    # Persist per-director attendance for meeting/person-wise reports
                    attendees = request.presentDirectors or request.directors or []
                    for d in attendees:
                        d_name = (d.get('name') or '').strip()
                        if not d_name:
                            continue
                        cursor.execute("""
                            INSERT INTO meeting_attendance
                            (minutes_id, company_name, meeting_type, meeting_date, director_name, din, status)
                            VALUES (%s, %s, %s, %s, %s, %s, %s)
                        """, (minutes_id, request.companyName, request.meetingType, request.meetingDate,
                              d_name, (d.get('din') or '').strip(), d.get('status') or 'Present'))
                    conn.commit()
                    return minutes_id
                finally:
                    conn.close()
            minutes_id = await loop.run_in_executor(thread_pool, record_history)
        except Exception as hist_err:
            logger.warning(f"Failed to record history (non-fatal): {hist_err}")
            minutes_id = None

        # Return JSON with download URL (the frontend expects JSON, not FileResponse)
        return {
            "message": "Document generated as draft. Finalize it from Meeting Minutes when approved.",
            "filename": filename,
            "download_url": f"/api/generated-minutes/download/{filename}",
            "id": minutes_id,
            "status": "draft",
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generating minutes: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to generate minutes: {str(e)}")


# --- Resolution Endpoints ---

@router.get("/resolutions", response_model=ResolutionTemplatesList)
async def get_resolutions():
    try:
        def fetch():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn: raise RuntimeError("Database connection unavailable")
            cursor = get_pg_cursor(conn)
            try:
                cursor.execute("SELECT id, template_name, resolution_text, created_at FROM resolution_templates ORDER BY template_name")
                rows = cursor.fetchall()
                return [ResolutionTemplateResponse(id=r['id'], template_name=r['template_name'], resolution_text=r['resolution_text'], created_at=str(r['created_at'])) for r in rows]
            finally:
                conn.close()
        data = await asyncio.get_running_loop().run_in_executor(thread_pool, fetch)
        return ResolutionTemplatesList(data=data, count=len(data))
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/resolutions", response_model=ResolutionTemplateResponse)
async def create_resolution(payload: ResolutionCreate, user: dict = Depends(require_session)):
    try:
        def insert():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn: raise RuntimeError("DB connection failed")
            cursor = get_pg_cursor(conn)
            try:
                cursor.execute(
                    "INSERT INTO resolution_templates (template_name, resolution_text) VALUES (%s, %s) RETURNING id, template_name, resolution_text, created_at",
                    (payload.template_name, payload.resolution_text))
                row = cursor.fetchone()
                conn.commit()
                return ResolutionTemplateResponse(id=row['id'], template_name=row['template_name'], resolution_text=row['resolution_text'], created_at=str(row['created_at']))
            finally:
                conn.close()
        return await asyncio.get_event_loop().run_in_executor(thread_pool, insert)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.delete("/resolutions/{id}")
async def delete_resolution(id: int, user: dict = Depends(require_session)):
    try:
        def delete():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn: raise RuntimeError("DB connection failed")
            cursor = get_pg_cursor(conn)
            try:
                cursor.execute("DELETE FROM resolution_templates WHERE id = %s", (id,))
                conn.commit()
                return {"message": "Resolution template deleted successfully"}
            finally:
                conn.close()
        return await asyncio.get_event_loop().run_in_executor(thread_pool, delete)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# --- New Verticals, Companies and Sourced Directors API ---

class VerticalResponse(BaseModel):
    id: int
    name: str
    code: str

class CompanyResponse(BaseModel):
    id: int
    name: str
    code: Optional[str] = None
    cin: Optional[str] = None
    type: Optional[str] = None
    vertical_id: Optional[int] = None
    status: Optional[str] = None
    secretary_name: Optional[str] = None
    created_by: Optional[str] = None
    created_at: Optional[str] = None
    updated_by: Optional[str] = None
    updated_at: Optional[str] = None
    # Meeting statistics (when filtered by meeting type)
    total_meetings: Optional[int] = None
    last_meeting_date: Optional[str] = None
    last_meeting_number: Optional[str] = None
    next_meeting_number: Optional[str] = None

class VerticalsListResponse(BaseModel):
    data: List[VerticalResponse]
    count: int

class CompaniesListResponse(BaseModel):
    data: List[CompanyResponse]
    count: int

@router.get("/verticals", response_model=VerticalsListResponse)
async def get_verticals():
    """List all Verticals (Business Units)."""
    try:
        def fetch():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn: raise RuntimeError("Database connection unavailable")
            cursor = get_pg_cursor(conn)
            try:
                cursor.execute("SELECT id, name, code FROM verticals ORDER BY name")
                rows = cursor.fetchall()
                return [VerticalResponse(id=r['id'], name=r['name'], code=r['code']) for r in rows]
            finally:
                conn.close()
        data = await asyncio.get_running_loop().run_in_executor(thread_pool, fetch)
        return VerticalsListResponse(data=data, count=len(data))
    except Exception as e:
        logger.error(f"Error listing verticals: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/verticals/{vertical_id}/companies", response_model=CompaniesListResponse)
async def get_vertical_companies(
    vertical_id: int, 
    q: Optional[str] = None, 
    meeting_type_filter: Optional[str] = None,  # NEW: Filter by meeting type
    limit: int = 100, 
    offset: int = 0
):
    """
    List companies belonging to a Vertical, with query search, meeting type filtering, and pagination.
    
    Args:
        vertical_id: ID of the business vertical
        q: Search query for company name
        meeting_type_filter: Filter by meeting type (Board Meeting, Audit Committee, etc.)
                           Shows only companies that have this meeting type
        limit: Maximum number of results
        offset: Pagination offset
        
    Returns:
        List of companies with meeting statistics (if filtered)
    """
    try:
        def fetch():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn: raise RuntimeError("Database connection unavailable")
            cursor = get_pg_cursor(conn)
            try:
                # Helper function to extract numeric part from meeting number
                def extract_meeting_number(meeting_num: str) -> int:
                    """Extract numeric value from meeting number like '87TH' or 'EIGHTY SEVENTH'"""
                    if not meeting_num:
                        return 0
                    # Try to extract digits
                    import re
                    match = re.search(r'\d+', str(meeting_num))
                    if match:
                        return int(match.group())
                    return 0
                
                def format_next_meeting_number(last_num: str) -> str:
                    """Generate next meeting number"""
                    current = extract_meeting_number(last_num)
                    next_num = current + 1
                    
                    # Return in ordinal format like "88TH"
                    suffix = 'TH'
                    if next_num % 10 == 1 and next_num % 100 != 11:
                        suffix = 'ST'
                    elif next_num % 10 == 2 and next_num % 100 != 12:
                        suffix = 'ND'
                    elif next_num % 10 == 3 and next_num % 100 != 13:
                        suffix = 'RD'
                    
                    return f"{next_num}{suffix}"
                
                # If meeting type filter is applied, join with generated_minutes for stats.
                # Still return ALL companies under the BU — filter only scopes the stats
                # and what is shown after opening a company.
                if meeting_type_filter and meeting_type_filter.lower() != 'all':
                    query = """
                        SELECT 
                            c.id, 
                            c.name,
                            c.code,
                            c.cin, 
                            c.type, 
                            c.vertical_id, 
                            c.status, 
                            c.secretary_name,
                            c.created_by,
                            c.created_at,
                            c.updated_by,
                            c.updated_at,
                            COUNT(gm.id) as total_meetings,
                            MAX(gm.meeting_date) as last_meeting_date,
                            MAX(gm.meeting_number) as last_meeting_number
                        FROM companies c
                        LEFT JOIN generated_minutes gm 
                            ON gm.company_name = c.name 
                            AND UPPER(gm.meeting_type) = UPPER(%s)
                        WHERE c.vertical_id = %s
                    """
                    params = [meeting_type_filter, vertical_id]
                    
                    if q:
                        query += " AND UPPER(c.name) LIKE UPPER(%s)"
                        params.append(f"%{q}%")
                    
                    query += """
                        GROUP BY c.id, c.name, c.code, c.cin, c.type, c.vertical_id, 
                                 c.status, c.secretary_name, c.created_by, c.created_at, 
                                 c.updated_by, c.updated_at
                        ORDER BY c.name 
                        LIMIT %s OFFSET %s
                    """
                    params.extend([limit, offset])
                    
                    cursor.execute(query, tuple(params))
                    rows = cursor.fetchall()

                    # Resolve true last meeting number by numeric order (MAX on '9TH'/'90TH' is unreliable)
                    last_by_company: Dict[str, str] = {}
                    if rows:
                        company_names = [r['name'] for r in rows]
                        placeholders = ", ".join(["%s"] * len(company_names))
                        cursor.execute(
                            f"""
                            SELECT company_name, meeting_number
                            FROM generated_minutes
                            WHERE UPPER(meeting_type) = UPPER(%s)
                              AND company_name IN ({placeholders})
                            """,
                            tuple([meeting_type_filter] + company_names),
                        )
                        for mr in cursor.fetchall():
                            cname = mr['company_name']
                            mnum = mr.get('meeting_number') or ''
                            prev = last_by_company.get(cname, '')
                            if extract_meeting_number(mnum) >= extract_meeting_number(prev):
                                last_by_company[cname] = mnum
                    
                    result = []
                    for r in rows:
                        last_date = r.get('last_meeting_date')
                        last_num = last_by_company.get(r['name'], '') or (r.get('last_meeting_number') or '')
                        result.append(CompanyResponse(
                            id=r['id'],
                            name=r['name'],
                            code=r.get('code', ''),
                            cin=r['cin'],
                            type=r['type'],
                            vertical_id=r['vertical_id'],
                            status=r['status'],
                            secretary_name=r['secretary_name'],
                            created_by=r.get('created_by'),
                            created_at=str(r.get('created_at', '')),
                            updated_by=r.get('updated_by'),
                            updated_at=str(r.get('updated_at', '')),
                            total_meetings=r.get('total_meetings', 0) or 0,
                            last_meeting_date=str(last_date) if last_date else '',
                            last_meeting_number=last_num,
                            next_meeting_number=format_next_meeting_number(last_num)
                        ))
                    
                    # Count all companies in this vertical (search-aware), not only those with meetings
                    count_query = "SELECT COUNT(*) as count FROM companies WHERE vertical_id = %s"
                    count_params = [vertical_id]
                    
                    if q:
                        count_query += " AND UPPER(name) LIKE UPPER(%s)"
                        count_params.append(f"%{q}%")
                    
                    cursor.execute(count_query, tuple(count_params))
                    total = cursor.fetchone()['count'] or 0
                    
                    return result, total
                
                else:
                    # No filter - show all companies (existing behavior)
                    query = """
                        SELECT id, name, code, cin, type, vertical_id, status, secretary_name,
                               created_by, created_at, updated_by, updated_at
                        FROM companies 
                        WHERE vertical_id = %s
                    """
                    params = [vertical_id]
                    
                    if q:
                        query += " AND UPPER(name) LIKE UPPER(%s)"
                        params.append(f"%{q}%")
                    
                    query += " ORDER BY name LIMIT %s OFFSET %s"
                    params.extend([limit, offset])
                    
                    cursor.execute(query, tuple(params))
                    rows = cursor.fetchall()
                    
                    result = [CompanyResponse(
                        id=r['id'], 
                        name=r['name'],
                        code=r.get('code', ''),
                        cin=r['cin'], 
                        type=r['type'], 
                        vertical_id=r['vertical_id'], 
                        status=r['status'], 
                        secretary_name=r['secretary_name'],
                        created_by=r.get('created_by'),
                        created_at=str(r.get('created_at', '')),
                        updated_by=r.get('updated_by'),
                        updated_at=str(r.get('updated_at', ''))
                    ) for r in rows]
                    
                    # Get total count for pagination
                    count_query = "SELECT COUNT(*) as count FROM companies WHERE vertical_id = %s"
                    count_params = [vertical_id]
                    if q:
                        count_query += " AND UPPER(name) LIKE UPPER(%s)"
                        count_params.append(f"%{q}%")
                    
                    cursor.execute(count_query, tuple(count_params))
                    total = cursor.fetchone()['count'] or 0
                    
                    return result, total
                    
            finally:
                conn.close()
        
        data, count = await asyncio.get_running_loop().run_in_executor(thread_pool, fetch)
        return CompaniesListResponse(data=data, count=count)
    except Exception as e:
        logger.error(f"Error listing vertical companies: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# --- Company Management Endpoints ---

class CompanyCreateRequest(BaseModel):
    name: str
    code: Optional[str] = ""
    cin: Optional[str] = ""
    type: Optional[str] = "Public Limited"
    secretary_name: Optional[str] = ""
    status: str = "Active"


def generate_company_code(company_name: str) -> str:
    """
    Auto-generate company code from name
    Examples: 
        "Adani Green Energy Limited" -> "AGEL"
        "Adani Ports and SEZ" -> "APSZ"
    """
    # Remove common suffixes
    name = company_name.upper()
    for suffix in [' LIMITED', ' LTD', ' PVT', ' PRIVATE', ' LTD.', ' CORP', ' CORPORATION', ' INC']:
        name = name.replace(suffix, '')
    
    # Split into words and take first letter of significant words
    words = name.split()
    # Filter out common words
    significant_words = [w for w in words if w not in ['AND', 'THE', 'OF', 'FOR', 'TO', 'IN', 'ON', 'AT', 'BY']]
    
    if len(significant_words) >= 2:
        # Take first letter of each word (max 4 letters)
        code = ''.join(w[0] for w in significant_words[:4])
    elif len(significant_words) == 1:
        # Take first 3-4 letters of single word
        code = significant_words[0][:4]
    else:
        # Fallback: take first 3 letters of original name
        code = words[0][:3] if words else 'UNK'
    
    return code


@router.post("/verticals/{vertical_id}/companies", response_model=CompanyResponse)
async def add_company(
    vertical_id: int, 
    request: CompanyCreateRequest,
    req: Request,
    user: dict = Depends(require_session)
):
    """
    Add new company to a vertical with audit logging
    
    Requires authentication. Creates a new company record and logs the action.
    
    Args:
        vertical_id: ID of the business vertical
        request: Company details (name, CIN, type, secretary, etc.)
        req: FastAPI Request object for IP/user agent extraction
        user: Authenticated user from session
        
    Returns:
        Created company details
    """
    if not request.name.strip():
        raise HTTPException(status_code=400, detail="Company name is required")
    
    def insert():
        conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
        if not conn:
            raise HTTPException(status_code=503, detail="Database connection unavailable")
        
        cursor = get_pg_cursor(conn)
        try:
            # Check if company already exists
            cursor.execute(
                "SELECT id, name FROM companies WHERE UPPER(name) = UPPER(%s)",
                (request.name.strip(),)
            )
            existing = cursor.fetchone()
            if existing:
                raise HTTPException(
                    status_code=409, 
                    detail=f"Company '{existing['name']}' already exists with ID {existing['id']}"
                )
            
            # Auto-generate code if not provided
            company_code = request.code.strip() if request.code else generate_company_code(request.name)
            
            # Insert company
            cursor.execute("""
                INSERT INTO companies 
                (name, code, cin, type, vertical_id, secretary_name, status, created_by, created_at, updated_at)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)
                RETURNING id, name, code, cin, type, vertical_id, secretary_name, status, 
                          created_by, created_at, updated_by, updated_at
            """, (
                request.name.strip(),
                company_code,
                request.cin.strip() if request.cin else None,
                request.type,
                vertical_id,
                request.secretary_name.strip() if request.secretary_name else None,
                request.status,
                user['email']
            ))
            
            row = cursor.fetchone()
            if not row:
                raise HTTPException(status_code=500, detail="Failed to create company")
            
            # Prepare company data for audit log
            company_data = {
                "id": row['id'],
                "name": row['name'],
                "code": row['code'],
                "cin": row['cin'],
                "type": row['type'],
                "vertical_id": row['vertical_id'],
                "secretary_name": row['secretary_name'],
                "status": row['status'],
                "created_by": row['created_by']
            }
            
            # Extract IP and user agent
            ip_address = get_client_ip(req)
            user_agent = get_user_agent(req)
            
            # Log the audit entry
            AuditLogger.log_company_created(
                conn=conn,
                company_id=row['id'],
                company_name=row['name'],
                company_data=company_data,
                user_email=user['email'],
                vertical_id=vertical_id,
                ip_address=ip_address,
                user_agent=user_agent
            )
            
            # Commit transaction (includes both insert and audit log)
            conn.commit()
            
            logger.info(f"Company '{row['name']}' created by {user['email']} (ID: {row['id']})")
            
            return CompanyResponse(
                id=row['id'],
                name=row['name'],
                code=row['code'],
                cin=row['cin'],
                type=row['type'],
                vertical_id=row['vertical_id'],
                status=row['status'],
                secretary_name=row['secretary_name'],
                created_by=row['created_by'],
                created_at=str(row['created_at']),
                updated_by=row.get('updated_by'),
                updated_at=str(row.get('updated_at', ''))
            )
            
        except HTTPException:
            conn.rollback()
            raise
        except Exception as e:
            conn.rollback()
            logger.error(f"Error adding company: {e}")
            raise HTTPException(status_code=500, detail=f"Failed to add company: {str(e)}")
        finally:
            conn.close()
    
    return await asyncio.get_running_loop().run_in_executor(thread_pool, insert)


@router.delete("/companies/{company_id}")
async def delete_company(
    company_id: int,
    req: Request,
    confirm: bool = False,
    user: dict = Depends(require_session)
):
    """
    Delete company and all related records with audit logging
    
    DANGER: This is a destructive operation that cannot be undone.
    Deletes:
    - Company record
    - All meetings/minutes for this company
    - All agendas for this company
    - All director mappings for this company
    - All governance records for this company
    - All attendance records for this company
    
    Args:
        company_id: ID of the company to delete
        req: FastAPI Request object
        confirm: Must be True to proceed (safety check)
        user: Authenticated user from session
        
    Returns:
        Deletion confirmation with statistics
    """
    if not confirm:
        raise HTTPException(
            status_code=400, 
            detail="Confirmation required. Set confirm=true query parameter to proceed with deletion."
        )
    
    def delete():
        conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
        if not conn:
            raise HTTPException(status_code=503, detail="Database connection unavailable")
        
        cursor = get_pg_cursor(conn)
        try:
            # Get company details BEFORE deletion for audit log
            cursor.execute("""
                SELECT id, name, code, cin, type, vertical_id, secretary_name, status,
                       created_by, created_at, updated_by, updated_at
                FROM companies 
                WHERE id = %s
            """, (company_id,))
            
            company_row = cursor.fetchone()
            if not company_row:
                raise HTTPException(status_code=404, detail=f"Company with ID {company_id} not found")
            
            company_data = dict(company_row)
            company_name = company_row['name']
            
            logger.info(f"Deleting company '{company_name}' (ID: {company_id}) by user {user['email']}")
            
            # Count related records before deletion
            cursor.execute(
                "SELECT COUNT(*) as count FROM generated_minutes WHERE company_name = %s",
                (company_name,)
            )
            meetings_count = cursor.fetchone()['count'] or 0
            
            cursor.execute(
                "SELECT COUNT(*) as count FROM meeting_attendance WHERE company_name = %s",
                (company_name,)
            )
            attendance_count = cursor.fetchone()['count'] or 0
            
            cursor.execute(
                "SELECT COUNT(*) as count FROM company_directors WHERE company_name = %s",
                (company_name,)
            )
            directors_count = cursor.fetchone()['count'] or 0
            
            # Check if meeting_agendas table exists
            agendas_count = 0
            try:
                cursor.execute(
                    "SELECT COUNT(*) as count FROM meeting_agendas WHERE company_name = %s",
                    (company_name,)
                )
                agendas_count = cursor.fetchone()['count'] or 0
            except Exception:
                pass  # Table might not exist yet
            
            # Check if governance_records table exists
            governance_count = 0
            try:
                cursor.execute(
                    "SELECT COUNT(*) as count FROM governance_records WHERE company_name = %s",
                    (company_name,)
                )
                governance_count = cursor.fetchone()['count'] or 0
            except Exception:
                pass  # Table might not exist yet
            
            total_related_records = (
                meetings_count + attendance_count + directors_count + 
                agendas_count + governance_count
            )
            
            # Delete all related records (cascade delete)
            cursor.execute("DELETE FROM meeting_attendance WHERE company_name = %s", (company_name,))
            cursor.execute("DELETE FROM generated_minutes WHERE company_name = %s", (company_name,))
            cursor.execute("DELETE FROM company_directors WHERE company_name = %s", (company_name,))
            
            # Delete from optional tables if they exist
            try:
                cursor.execute("DELETE FROM meeting_agendas WHERE company_name = %s", (company_name,))
            except Exception:
                pass
            
            try:
                cursor.execute("DELETE FROM governance_records WHERE company_name = %s", (company_name,))
            except Exception:
                pass
            
            # Finally delete the company itself
            cursor.execute("DELETE FROM companies WHERE id = %s RETURNING id", (company_id,))
            deleted_row = cursor.fetchone()
            
            if not deleted_row:
                raise HTTPException(status_code=404, detail="Company not found or already deleted")
            
            # Extract IP and user agent
            ip_address = get_client_ip(req)
            user_agent = get_user_agent(req)
            
            # Log the deletion in audit log
            AuditLogger.log_company_deleted(
                conn=conn,
                company_id=company_id,
                company_name=company_name,
                company_data=company_data,
                user_email=user['email'],
                deleted_records_count=total_related_records,
                ip_address=ip_address,
                user_agent=user_agent
            )
            
            # Commit transaction (includes all deletes and audit log)
            conn.commit()
            
            logger.info(
                f"Company '{company_name}' deleted successfully by {user['email']}. "
                f"Deleted {total_related_records} related records."
            )
            
            return {
                "success": True,
                "message": f"Company '{company_name}' deleted successfully",
                "company_id": company_id,
                "company_name": company_name,
                "deleted_records": {
                    "meetings": meetings_count,
                    "attendance_records": attendance_count,
                    "directors": directors_count,
                    "agendas": agendas_count,
                    "governance_records": governance_count,
                    "total": total_related_records
                },
                "deleted_by": user['email'],
                "deleted_at": datetime.now().isoformat()
            }
            
        except HTTPException:
            conn.rollback()
            raise
        except Exception as e:
            conn.rollback()
            logger.error(f"Error deleting company {company_id}: {e}")
            raise HTTPException(status_code=500, detail=f"Failed to delete company: {str(e)}")
        finally:
            conn.close()
    
    return await asyncio.get_running_loop().run_in_executor(thread_pool, delete)


class CompanyUpdateRequest(BaseModel):
    name: Optional[str] = None
    code: Optional[str] = None
    cin: Optional[str] = None
    type: Optional[str] = None
    secretary_name: Optional[str] = None
    status: Optional[str] = None


@router.put("/companies/{company_id}", response_model=CompanyResponse)
async def update_company(
    company_id: int,
    request: CompanyUpdateRequest,
    req: Request,
    user: dict = Depends(require_session)
):
    """
    Update company details with audit logging
    
    Updates one or more fields of a company record. Only provided fields are updated.
    Tracks before/after state in audit log for accountability.
    
    Args:
        company_id: ID of the company to update
        request: Fields to update (only provided fields are changed)
        req: FastAPI Request object
        user: Authenticated user from session
        
    Returns:
        Updated company details
    """
    def update():
        conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
        if not conn:
            raise HTTPException(status_code=503, detail="Database connection unavailable")
        
        cursor = get_pg_cursor(conn)
        try:
            # Get current company data BEFORE update for audit log
            cursor.execute("""
                SELECT id, name, code, cin, type, vertical_id, secretary_name, status,
                       created_by, created_at, updated_by, updated_at
                FROM companies 
                WHERE id = %s
            """, (company_id,))
            
            old_row = cursor.fetchone()
            if not old_row:
                raise HTTPException(status_code=404, detail=f"Company with ID {company_id} not found")
            
            old_data = dict(old_row)
            company_name = old_row['name']
            
            # Build dynamic UPDATE query based on provided fields
            update_fields = []
            update_values = []
            
            if request.name is not None and request.name.strip():
                # Check if new name conflicts with existing company
                if request.name.strip().upper() != old_row['name'].upper():
                    cursor.execute(
                        "SELECT id, name FROM companies WHERE UPPER(name) = UPPER(%s) AND id != %s",
                        (request.name.strip(), company_id)
                    )
                    conflict = cursor.fetchone()
                    if conflict:
                        raise HTTPException(
                            status_code=409,
                            detail=f"Company name '{conflict['name']}' already exists (ID: {conflict['id']})"
                        )
                update_fields.append("name = %s")
                update_values.append(request.name.strip())
            
            if request.code is not None:
                update_fields.append("code = %s")
                update_values.append(request.code.strip() if request.code else None)
            
            if request.cin is not None:
                update_fields.append("cin = %s")
                update_values.append(request.cin.strip() if request.cin else None)
            
            if request.type is not None:
                update_fields.append("type = %s")
                update_values.append(request.type)
            
            if request.secretary_name is not None:
                update_fields.append("secretary_name = %s")
                update_values.append(request.secretary_name.strip() if request.secretary_name else None)
            
            if request.status is not None:
                update_fields.append("status = %s")
                update_values.append(request.status)
            
            if not update_fields:
                raise HTTPException(
                    status_code=400,
                    detail="No fields provided for update. At least one field must be specified."
                )
            
            # Add updated_by and updated_at
            update_fields.append("updated_by = %s")
            update_values.append(user['email'])
            
            update_fields.append("updated_at = CURRENT_TIMESTAMP")
            
            # Build and execute UPDATE query
            query = f"""
                UPDATE companies 
                SET {', '.join(update_fields)}
                WHERE id = %s
                RETURNING id, name, code, cin, type, vertical_id, secretary_name, status,
                          created_by, created_at, updated_by, updated_at
            """
            update_values.append(company_id)
            
            cursor.execute(query, tuple(update_values))
            updated_row = cursor.fetchone()
            
            if not updated_row:
                raise HTTPException(status_code=404, detail="Company not found or update failed")
            
            new_data = dict(updated_row)
            
            # Extract IP and user agent
            ip_address = get_client_ip(req)
            user_agent = get_user_agent(req)
            
            # Log the update in audit log
            AuditLogger.log_company_updated(
                conn=conn,
                company_id=company_id,
                company_name=company_name,
                old_data=old_data,
                new_data=new_data,
                user_email=user['email'],
                ip_address=ip_address,
                user_agent=user_agent
            )
            
            # Commit transaction (includes update and audit log)
            conn.commit()
            
            logger.info(f"Company '{company_name}' updated by {user['email']} (ID: {company_id})")
            
            return CompanyResponse(
                id=updated_row['id'],
                name=updated_row['name'],
                code=updated_row['code'],
                cin=updated_row['cin'],
                type=updated_row['type'],
                vertical_id=updated_row['vertical_id'],
                status=updated_row['status'],
                secretary_name=updated_row['secretary_name'],
                created_by=updated_row['created_by'],
                created_at=str(updated_row['created_at']),
                updated_by=updated_row['updated_by'],
                updated_at=str(updated_row['updated_at'])
            )
            
        except HTTPException:
            conn.rollback()
            raise
        except Exception as e:
            conn.rollback()
            logger.error(f"Error updating company {company_id}: {e}")
            raise HTTPException(status_code=500, detail=f"Failed to update company: {str(e)}")
        finally:
            conn.close()
    
    return await asyncio.get_running_loop().run_in_executor(thread_pool, update)


def _extract_meeting_number_int(meeting_num: Optional[str]) -> int:
    """Extract numeric value from meeting number like '87TH' or '87'."""
    if not meeting_num:
        return 0
    import re
    match = re.search(r'\d+', str(meeting_num))
    return int(match.group()) if match else 0


def _format_next_meeting_number(last_num: Optional[str]) -> str:
    """Generate next ordinal meeting number (e.g. 87TH -> 88TH)."""
    current = _extract_meeting_number_int(last_num)
    next_num = current + 1
    suffix = 'TH'
    if next_num % 10 == 1 and next_num % 100 != 11:
        suffix = 'ST'
    elif next_num % 10 == 2 and next_num % 100 != 12:
        suffix = 'ND'
    elif next_num % 10 == 3 and next_num % 100 != 13:
        suffix = 'RD'
    return f"{next_num}{suffix}"


@router.get("/companies/{company_id}/meetings", response_model=CompanyMeetingsListResponse)
async def get_company_meetings(
    company_id: int,
    meeting_type: Optional[str] = None,
):
    """
    List meetings/minutes for a company, optionally filtered by meeting type.
    Ordered by meeting number (numeric) ascending, then meeting date ascending.
    """
    try:
        def fetch():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn:
                raise RuntimeError("Database connection unavailable")
            cursor = get_pg_cursor(conn)
            try:
                cursor.execute(
                    "SELECT id, name FROM companies WHERE id = %s",
                    (company_id,),
                )
                company = cursor.fetchone()
                if not company:
                    return None

                company_name = company['name']
                params: list = [company_name]
                type_clause = ""
                if meeting_type and meeting_type.lower() != 'all':
                    type_clause = " AND UPPER(meeting_type) = UPPER(%s)"
                    params.append(meeting_type)

                cursor.execute(
                    f"""
                    SELECT id, company_name, meeting_type, meeting_date, meeting_number,
                           meeting_year, file_path, created_at, status, finalized_at, finalized_by,
                           is_signed, unsigned_file_path
                    FROM generated_minutes
                    WHERE company_name = %s
                    {type_clause}
                    """,
                    tuple(params),
                )
                rows = cursor.fetchall()

                # Sort by numeric meeting number, then date (oldest → newest)
                def sort_key(r):
                    num = _extract_meeting_number_int(r.get('meeting_number'))
                    date_str = str(r.get('meeting_date') or '')
                    return (num, date_str)

                rows_sorted = sorted(rows, key=sort_key)
                data = [_row_to_generated_minute(r) for r in rows_sorted]

                last_num = data[-1].meeting_number if data else None
                return data, len(data), _format_next_meeting_number(last_num)
            finally:
                conn.close()

        result = await asyncio.get_running_loop().run_in_executor(thread_pool, fetch)
        if result is None:
            raise HTTPException(status_code=404, detail=f"Company {company_id} not found")
        data, count, next_num = result
        return CompanyMeetingsListResponse(
            data=data,
            count=count,
            meeting_type=meeting_type if meeting_type and meeting_type.lower() != 'all' else None,
            next_meeting_number=next_num,
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error listing company meetings: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/companies/{company_name}/directors")
async def get_company_directors(company_name: str):
    """
    Fetch directors for a company dynamically (no hardcoded people).
    Sources (merged, de-duplicated by DIN/name):
      1) Director Disclosure DB external_board_members (when Postgres available)
      2) Minutes DB external_board_members (local/SQLite seed)
      3) Minutes DB company_directors overlay
    Company names match with Ltd./Limited normalization.
    """
    try:
        def fetch():
            term = (company_name or "").strip()
            if not term:
                return []
            results = []
            seen = set()

            def add_row(name, din, designation="Director", source="unknown", row_id=None):
                name = (name or "").strip()
                if not name:
                    return
                key = (din or "").strip() or name.upper()
                if key in seen:
                    for existing in results:
                        ek = (existing.get("din") or "").strip() or existing["name"].upper()
                        if ek == key:
                            cur_d = (existing.get("designation") or "Director")
                            new_d = designation or "Director"
                            if cur_d == "Director" and new_d and new_d != "Director":
                                existing["designation"] = new_d
                            break
                    return
                seen.add(key)
                results.append({
                    "name": name,
                    "din": (din or "").strip(),
                    "designation": (designation or "Director").strip() or "Director",
                    "source": source,
                    "id": row_id,
                })

            def rows_for_company(rows, name_field="company_name"):
                matched = []
                for r in rows or []:
                    try:
                        cname = r[name_field]
                    except Exception:
                        cname = r.get(name_field) if hasattr(r, 'get') else None
                    if _company_names_match(term, cname):
                        matched.append(r)
                return matched

            # 1. Director Disclosure Postgres (optional)
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_DIRECTOR'))
            if conn:
                cursor = get_pg_cursor(conn)
                try:
                    token = re.sub(r"[^A-Za-z0-9 ]", " ", term).split()
                    distinctive = [w for w in token if len(w) > 3 and w.lower() not in ("adani", "limited", "private", "energy", "green")]
                    like = f"%{distinctive[-1]}%" if distinctive else (f"%{token[-1]}%" if token else f"%{term}%")
                    try:
                        cursor.execute(
                            """
                            SELECT DISTINCT name, din, designation, company_name
                            FROM directors_master.external_board_members
                            WHERE company_name IS NOT NULL
                              AND UPPER(company_name) LIKE UPPER(%s)
                            """,
                            (like,),
                        )
                        disc_rows = cursor.fetchall()
                    except Exception:
                        cursor.execute(
                            """
                            SELECT DISTINCT name, din, company_name
                            FROM directors_master.external_board_members
                            WHERE company_name IS NOT NULL
                              AND UPPER(company_name) LIKE UPPER(%s)
                            """,
                            (like,),
                        )
                        disc_rows = cursor.fetchall()
                    for r in rows_for_company(disc_rows):
                        desig = r["designation"] if "designation" in r.keys() else "Director"
                        add_row(r["name"], r["din"], desig, "disclosure")
                except Exception as ex:
                    logger.warning(f"Disclosure directors query failed for {term}: {ex}")
                finally:
                    conn.close()

            # 2 + 3. Minutes DB: external_board_members + company_directors
            try:
                m_conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
                if m_conn:
                    m_cursor = get_pg_cursor(m_conn)
                    try:
                        token = re.sub(r"[^A-Za-z0-9 ]", " ", term).split()
                        distinctive = [w for w in token if len(w) > 3 and w.lower() not in ("adani", "limited", "private", "energy", "green")]
                        like = f"%{distinctive[-1]}%" if distinctive else (f"%{token[-1]}%" if token else f"%{term}%")

                        try:
                            m_cursor.execute(
                                """
                                SELECT name, din, designation, company_name
                                FROM external_board_members
                                WHERE company_name IS NOT NULL
                                  AND UPPER(company_name) LIKE UPPER(%s)
                                """,
                                (like,),
                            )
                        except Exception:
                            m_cursor.execute(
                                """
                                SELECT name, din, company_name
                                FROM external_board_members
                                WHERE company_name IS NOT NULL
                                  AND UPPER(company_name) LIKE UPPER(%s)
                                """,
                                (like,),
                            )
                        for r in rows_for_company(m_cursor.fetchall()):
                            desig = r["designation"] if "designation" in r.keys() else "Director"
                            add_row(r["name"], r["din"], desig, "minutes_external")

                        if not any(x["source"] == "minutes_external" for x in results):
                            try:
                                m_cursor.execute(
                                    "SELECT name, din, designation, company_name FROM external_board_members WHERE company_name IS NOT NULL"
                                )
                            except Exception:
                                m_cursor.execute(
                                    "SELECT name, din, company_name FROM external_board_members WHERE company_name IS NOT NULL"
                                )
                            for r in rows_for_company(m_cursor.fetchall()):
                                desig = r["designation"] if "designation" in r.keys() else "Director"
                                add_row(r["name"], r["din"], desig, "minutes_external")

                        try:
                            m_cursor.execute(
                                """
                                SELECT id, name, din, designation, company_name
                                FROM company_directors
                                WHERE company_name IS NOT NULL
                                  AND UPPER(company_name) LIKE UPPER(%s)
                                ORDER BY name
                                """,
                                (like,),
                            )
                        except Exception:
                            m_cursor.execute(
                                """
                                SELECT id, name, din, company_name
                                FROM company_directors
                                WHERE company_name IS NOT NULL
                                  AND UPPER(company_name) LIKE UPPER(%s)
                                ORDER BY name
                                """,
                                (like,),
                            )
                        local_rows = rows_for_company(m_cursor.fetchall())
                        if not local_rows:
                            try:
                                m_cursor.execute(
                                    "SELECT id, name, din, designation, company_name FROM company_directors WHERE company_name IS NOT NULL"
                                )
                            except Exception:
                                m_cursor.execute(
                                    "SELECT id, name, din, company_name FROM company_directors WHERE company_name IS NOT NULL"
                                )
                            local_rows = rows_for_company(m_cursor.fetchall())
                        for r in local_rows:
                            desig = r["designation"] if "designation" in r.keys() else "Director"
                            add_row(r["name"], r["din"], desig, "local", r["id"] if "id" in r.keys() else None)
                    finally:
                        m_conn.close()
            except Exception as ex:
                logger.warning(f"Minutes directors query failed for {term}: {ex}")

            def sort_key(d):
                desig = (d.get("designation") or "").lower()
                chair_rank = 0 if "chair" in desig else 1
                return (chair_rank, d.get("name") or "")

            results.sort(key=sort_key)
            return results

        loop = asyncio.get_running_loop()
        data = await loop.run_in_executor(thread_pool, fetch)
        default_chairman = ""
        for d in data:
            desig = (d.get("designation") or "").lower()
            if "chair" in desig:
                default_chairman = d["name"]
                break
        # Do not fall back to first director — seed data lists Gautam Adani on most companies
        return {
            "data": data,
            "count": len(data),
            "default_chairman": default_chairman,
            "company_name": company_name,
        }
    except Exception as e:
        logger.error(f"Error fetching company directors: {e}")
        return {"data": [], "count": 0, "default_chairman": "", "company_name": company_name}


@router.get("/directors")
@router.get("/directors-master")
async def get_all_master_directors(q: Optional[str] = None):
    """Fetch all master directors across all portfolio entities."""
    try:
        def fetch():
            results = []
            seen = set()
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if conn:
                cursor = get_pg_cursor(conn)
                try:
                    query = "SELECT DISTINCT name, din, company_name FROM external_board_members"
                    params = []
                    if q:
                        query += " WHERE UPPER(name) LIKE UPPER(%s) OR UPPER(din) LIKE UPPER(%s) OR UPPER(company_name) LIKE UPPER(%s)"
                        params = [f"%{q}%", f"%{q}%", f"%{q}%"]
                    query += " ORDER BY name LIMIT 500"
                    cursor.execute(query, params)
                    for r in cursor.fetchall():
                        k = (r["din"] or "").strip() or (r["name"] or "").strip().upper()
                        if k and k not in seen:
                            seen.add(k)
                            results.append({"name": r["name"], "din": r["din"], "company_name": r.get("company_name", ""), "source": "registry"})
                except Exception as ex:
                    logger.warning(f"Error fetching external_board_members: {ex}")

                try:
                    query = "SELECT id, name, din, company_name FROM company_directors"
                    params = []
                    if q:
                        query += " WHERE UPPER(name) LIKE UPPER(%s) OR UPPER(din) LIKE UPPER(%s) OR UPPER(company_name) LIKE UPPER(%s)"
                        params = [f"%{q}%", f"%{q}%", f"%{q}%"]
                    query += " ORDER BY name LIMIT 200"
                    cursor.execute(query, params)
                    for r in cursor.fetchall():
                        k = (r["din"] or "").strip() or (r["name"] or "").strip().upper()
                        if k and k not in seen:
                            seen.add(k)
                            results.append({"id": r["id"], "name": r["name"], "din": r["din"], "company_name": r.get("company_name", ""), "source": "local"})
                except Exception as ex:
                    logger.warning(f"Error fetching company_directors: {ex}")
                conn.close()
            return results

        loop = asyncio.get_running_loop()
        data = await loop.run_in_executor(thread_pool, fetch)
        return {"data": data, "count": len(data)}
    except Exception as e:
        logger.error(f"Error fetching all master directors: {e}")
        return {"data": [], "count": 0}


# --- Local Company Director Overlay CRUD (never writes to Director Disclosure DB) ---

class CompanyDirectorCreate(BaseModel):
    name: str
    din: Optional[str] = ""


@router.post("/companies/{company_name}/directors")
async def add_company_director(company_name: str, payload: CompanyDirectorCreate, user: dict = Depends(require_session)):
    """Add a manual director entry for a company in the local minutes DB only."""
    if not payload.name.strip():
        raise HTTPException(status_code=400, detail="Director name is required")
    try:
        def insert():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn: raise RuntimeError("Database connection unavailable")
            cursor = get_pg_cursor(conn)
            try:
                cursor.execute("""
                    INSERT INTO company_directors (company_name, name, din)
                    VALUES (%s, %s, %s)
                    RETURNING id, name, din, created_at
                """, (company_name.strip(), payload.name.strip(), (payload.din or "").strip()))
                row = cursor.fetchone()
                conn.commit()
                return {"id": row["id"], "name": row["name"], "din": row["din"],
                        "source": "local", "created_at": str(row["created_at"])}
            finally:
                conn.close()
        return await asyncio.get_running_loop().run_in_executor(thread_pool, insert)
    except Exception as e:
        logger.error(f"Error adding company director: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.put("/companies/{company_name}/directors/{director_id}")
async def update_company_director(company_name: str, director_id: int, payload: CompanyDirectorCreate, user: dict = Depends(require_session)):
    """Update a manual (local) director entry. Disclosure-DB directors cannot be edited from here."""
    try:
        def update():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn: raise RuntimeError("Database connection unavailable")
            cursor = get_pg_cursor(conn)
            try:
                cursor.execute("""
                    UPDATE company_directors SET name = %s, din = %s
                    WHERE id = %s AND UPPER(company_name) = UPPER(%s)
                    RETURNING id, name, din, created_at
                """, (payload.name.strip(), (payload.din or "").strip(), director_id, company_name.strip()))
                row = cursor.fetchone()
                conn.commit()
                if not row:
                    return None
                return {"id": row["id"], "name": row["name"], "din": row["din"],
                        "source": "local", "created_at": str(row["created_at"])}
            finally:
                conn.close()
        result = await asyncio.get_running_loop().run_in_executor(thread_pool, update)
        if result is None:
            raise HTTPException(status_code=404, detail="Local director entry not found (registry directors are read-only)")
        return result
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error updating company director: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/companies/{company_name}/directors/{director_id}")
async def delete_company_director(company_name: str, director_id: int, user: dict = Depends(require_session)):
    """Delete a manual (local) director entry. Disclosure-DB directors cannot be deleted from here."""
    try:
        def delete():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn: raise RuntimeError("Database connection unavailable")
            cursor = get_pg_cursor(conn)
            try:
                cursor.execute("""
                    DELETE FROM company_directors
                    WHERE id = %s AND UPPER(company_name) = UPPER(%s)
                    RETURNING id
                """, (director_id, company_name.strip()))
                row = cursor.fetchone()
                conn.commit()
                return row is not None
            finally:
                conn.close()
        deleted = await asyncio.get_running_loop().run_in_executor(thread_pool, delete)
        if not deleted:
            raise HTTPException(status_code=404, detail="Local director entry not found (registry directors are read-only)")
        return {"message": "Director removed", "id": director_id}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting company director: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/resolutions/search")
async def search_resolutions(q: str = ""):
    """Keyword search for resolution templates by title or content."""
    try:
        def fetch():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn: raise RuntimeError("Database connection unavailable")
            cursor = get_pg_cursor(conn)
            try:
                term = f"%{q}%"
                cursor.execute("""
                    SELECT id, template_name, resolution_text, created_at 
                    FROM resolution_templates 
                    WHERE UPPER(template_name) LIKE UPPER(%s) OR UPPER(resolution_text) LIKE UPPER(%s)
                    ORDER BY template_name
                """, (term, term))
                rows = cursor.fetchall()
                return [ResolutionTemplateResponse(id=r['id'], template_name=r['template_name'], resolution_text=r['resolution_text'], created_at=str(r['created_at'])) for r in rows]
            finally:
                conn.close()
        data = await asyncio.get_running_loop().run_in_executor(thread_pool, fetch)
        return {"data": data, "count": len(data)}
    except Exception as e:
        logger.error(f"Error searching resolutions: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# --- Structured Repository Explorer & Search APIs (BRD Requirements #7 & #10) ---

@router.get("/repository/tree")
async def get_repository_tree():
    """BRD Requirement #7: Return structured directory tree of generated minutes.
    Hierarchy: Vertical (BU) -> Company Name -> Meeting -> Type of Meeting -> Date(year)
    """
    repo_dir = os.path.join(os.path.dirname(__file__), "..", "public", "repository")
    
    def build_tree(path):
        tree = []
        if not os.path.exists(path):
            return tree
        for entry in sorted(os.listdir(path)):
            full_path = os.path.join(path, entry)
            rel_path = os.path.relpath(full_path, repo_dir).replace('\\', '/')
            if os.path.isdir(full_path):
                tree.append({
                    "name": entry,
                    "type": "folder",
                    "path": rel_path,
                    "children": build_tree(full_path)
                })
            elif entry.endswith('.docx') or entry.endswith('.pdf'):
                stats = os.stat(full_path)
                tree.append({
                    "name": entry,
                    "type": "file",
                    "path": rel_path,
                    "size": stats.st_size,
                    "lastModified": datetime.fromtimestamp(stats.st_mtime).strftime('%Y-%m-%d %H:%M:%S'),
                    "download_url": f"/api/generated-minutes/download/{entry}"
                })
        return tree

    try:
        loop = asyncio.get_running_loop()
        tree_data = await loop.run_in_executor(thread_pool, build_tree, repo_dir)
        return {"data": tree_data, "count": len(tree_data)}
    except Exception as e:
        logger.error(f"Error building repository tree: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/repository/search")
async def search_repository(q: str = ""):
    """BRD Requirement #10: Keyword & Title Search for generated minutes documents across DB and Repository."""
    if not q or not q.strip():
        return await get_history()

    search_term = q.strip().upper()
    results = []

    try:
        # Search PostgreSQL DB records
        def search_db():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn:
                return []
            cursor = get_pg_cursor(conn)
            try:
                cursor.execute("""
                    SELECT id, company_name, meeting_type, meeting_date, file_path, created_at 
                    FROM generated_minutes 
                    WHERE UPPER(company_name) LIKE %s 
                       OR UPPER(meeting_type) LIKE %s 
                       OR UPPER(file_path) LIKE %s 
                    ORDER BY id DESC
                """, (f"%{search_term}%", f"%{search_term}%", f"%{search_term}%"))
                rows = cursor.fetchall()
                return [GeneratedMinuteResponse(
                    id=r['id'],
                    company_name=r['company_name'],
                    meeting_type=r['meeting_type'],
                    meeting_date=str(r['meeting_date']),
                    file_path=r['file_path'],
                    created_at=str(r['created_at']),
                    download_url=f"/api/generated-minutes/download/{r['file_path']}"
                ) for r in rows]
            finally:
                conn.close()

        loop = asyncio.get_running_loop()
        db_results = await loop.run_in_executor(thread_pool, search_db)
        return {"data": db_results, "count": len(db_results)}
    except Exception as e:
        logger.error(f"Error searching repository: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/repository/upload")
async def upload_repository_document(
    file: UploadFile = File(...),
    vertical_name: str = Form("Energy"),
    company_name: str = Form("General_Company"),
    meeting_number: str = Form("1"),
    meeting_type: str = Form("Board_Meeting"),
    meeting_year: str = Form("2026")
):
    """BRD Requirement #8: Upload PDF/Word document into structured repository folder path.
    Now also extracts text + tables from uploaded documents and stores parsed content in DB."""
    if not file.filename or not (file.filename.endswith('.pdf') or file.filename.endswith('.docx') or file.filename.endswith('.pptx')):
        raise HTTPException(status_code=400, detail="Only .pdf, .docx and .pptx files are supported")

    def _clean_str(s: str) -> str:
        return "".join(c for c in s if c.isalnum() or c in (' ', '_', '-')).strip().replace(' ', '_')

    try:
        v_name = _clean_str(vertical_name)
        c_name = _clean_str(company_name)
        m_name = _clean_str(f"Meeting_{meeting_number}")
        t_name = _clean_str(meeting_type)
        y_name = _clean_str(meeting_year)

        target_dir = os.path.join(
            os.path.dirname(__file__), "..", "public", "repository",
            v_name, c_name, m_name, t_name, y_name
        )
        os.makedirs(target_dir, exist_ok=True)

        file_path = os.path.join(target_dir, file.filename)
        file_bytes = await file.read()
        with open(file_path, "wb") as f:
            f.write(file_bytes)

        # --- Extract text + tables from uploaded document ---
        extracted = {"text": "", "paragraph_count": 0, "tables": [], "table_count": 0}
        try:
            if file.filename.endswith('.docx') and DOCX_AVAILABLE:
                extracted = extract_text_from_docx(file_bytes)
            elif file.filename.endswith('.pdf'):
                extracted = extract_text_from_pdf(file_bytes)
        except Exception as parse_err:
            logger.warning(f"Content extraction failed for {file.filename}: {parse_err}")

        # Record in PostgreSQL history + store extracted content
        def record_db():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if conn:
                cursor = get_pg_cursor(conn)
                try:
                    # Ensure document_contents table exists
                    cursor.execute("""
                        CREATE TABLE IF NOT EXISTS document_contents (
                            id SERIAL PRIMARY KEY,
                            filename VARCHAR(500) NOT NULL,
                            company_name VARCHAR(500),
                            vertical_name VARCHAR(200),
                            meeting_type VARCHAR(200),
                            meeting_year VARCHAR(10),
                            file_type VARCHAR(10),
                            extracted_text TEXT,
                            tables_json JSONB,
                            paragraph_count INTEGER DEFAULT 0,
                            table_count INTEGER DEFAULT 0,
                            file_path VARCHAR(1000),
                            uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                        )
                    """)
                    conn.commit()

                    # Insert extracted content
                    file_ext = os.path.splitext(file.filename)[1].lstrip('.')
                    rel_path = f"{v_name}/{c_name}/{m_name}/{t_name}/{y_name}/{file.filename}"
                    cursor.execute("""
                        INSERT INTO document_contents 
                        (filename, company_name, vertical_name, meeting_type, meeting_year, file_type,
                         extracted_text, tables_json, paragraph_count, table_count, file_path)
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                        RETURNING id
                    """, (
                        file.filename, company_name, vertical_name, meeting_type, meeting_year,
                        file_ext, extracted["text"], json.dumps(extracted["tables"]),
                        extracted["paragraph_count"], extracted["table_count"], rel_path
                    ))
                    doc_id = cursor.fetchone()["id"]

                    # Also insert into generated_minutes for repository tree listing
                    cursor.execute("""
                        INSERT INTO generated_minutes (company_name, meeting_type, meeting_date, file_path)
                        VALUES (%s, %s, %s, %s)
                    """, (company_name, meeting_type, f"{meeting_year}-01-01", file.filename))
                    conn.commit()
                    return doc_id
                except Exception as db_err:
                    logger.error(f"DB error during document content save: {db_err}")
                    conn.rollback()
                    return None
                finally:
                    conn.close()
            return None

        loop = asyncio.get_running_loop()
        doc_id = await loop.run_in_executor(thread_pool, record_db)

        logger.info(f"Repository file uploaded and parsed: {file_path} (text: {extracted['paragraph_count']} paragraphs, tables: {extracted['table_count']})")
        return {
            "message": "File uploaded and parsed successfully",
            "filename": file.filename,
            "document_id": doc_id,
            "path": f"{v_name}/{c_name}/{m_name}/{t_name}/{y_name}/{file.filename}",
            "extraction_summary": {
                "paragraphs": extracted["paragraph_count"],
                "tables": extracted["table_count"],
                "ocr_used": extracted.get("ocr_used", False),
                "text_preview": extracted["text"][:300] + "..." if len(extracted["text"]) > 300 else extracted["text"]
            }
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error uploading repository document: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to upload document: {str(e)}")


@router.get("/repository/documents")
async def get_repository_documents(q: str = "", vertical: str = "", limit: int = 50):
    """List all uploaded documents with extracted content metadata. Supports keyword search."""
    try:
        def fetch():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn:
                return []
            cursor = get_pg_cursor(conn)
            try:
                if q:
                    cursor.execute("""
                        SELECT id, filename, company_name, vertical_name, meeting_type, meeting_year,
                               file_type, paragraph_count, table_count, file_path, uploaded_at
                        FROM document_contents
                        WHERE LOWER(filename) LIKE %s OR LOWER(extracted_text) LIKE %s 
                              OR LOWER(company_name) LIKE %s
                        ORDER BY uploaded_at DESC LIMIT %s
                    """, (f"%{q.lower()}%", f"%{q.lower()}%", f"%{q.lower()}%", limit))
                elif vertical:
                    cursor.execute("""
                        SELECT id, filename, company_name, vertical_name, meeting_type, meeting_year,
                               file_type, paragraph_count, table_count, file_path, uploaded_at
                        FROM document_contents
                        WHERE LOWER(vertical_name) = %s
                        ORDER BY uploaded_at DESC LIMIT %s
                    """, (vertical.lower(), limit))
                else:
                    cursor.execute("""
                        SELECT id, filename, company_name, vertical_name, meeting_type, meeting_year,
                               file_type, paragraph_count, table_count, file_path, uploaded_at
                        FROM document_contents
                        ORDER BY uploaded_at DESC LIMIT %s
                    """, (limit,))
                rows = cursor.fetchall()
                return [{
                    "id": r["id"],
                    "filename": r["filename"],
                    "company_name": r["company_name"],
                    "vertical_name": r["vertical_name"],
                    "meeting_type": r["meeting_type"],
                    "meeting_year": r["meeting_year"],
                    "file_type": r["file_type"],
                    "paragraph_count": r["paragraph_count"],
                    "table_count": r["table_count"],
                    "file_path": r["file_path"],
                    "uploaded_at": str(r["uploaded_at"])
                } for r in rows]
            finally:
                conn.close()

        loop = asyncio.get_running_loop()
        docs = await loop.run_in_executor(thread_pool, fetch)
        return {"data": docs, "count": len(docs)}
    except Exception as e:
        logger.error(f"Error fetching repository documents: {e}")
        return {"data": [], "count": 0}


@router.get("/repository/document-content/{doc_id}")
async def get_document_content(doc_id: int):
    """Retrieve full extracted text and tables for a document (queries document_contents DB or parses file on-the-fly)."""
    try:
        def fetch():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn:
                raise HTTPException(status_code=503, detail="Database unavailable")
            cursor = get_pg_cursor(conn)
            try:
                # 1. Try fetching from document_contents table (optional — may not exist in SQLite)
                try:
                    cursor.execute("""
                        SELECT id, filename, company_name, vertical_name, meeting_type, meeting_year,
                               file_type, extracted_text, tables_json, paragraph_count, table_count, 
                               file_path, uploaded_at
                        FROM document_contents WHERE id = %s
                    """, (doc_id,))
                    row = cursor.fetchone()
                    if row:
                        tables = row["tables_json"]
                        if isinstance(tables, str):
                            tables = json.loads(tables)
                        return {
                            "id": row["id"],
                            "filename": row["filename"],
                            "company_name": row["company_name"],
                            "vertical_name": row["vertical_name"],
                            "meeting_type": row["meeting_type"],
                            "meeting_year": row["meeting_year"],
                            "file_type": row["file_type"],
                            "extracted_text": row["extracted_text"],
                            "tables": tables,
                            "paragraph_count": row["paragraph_count"],
                            "table_count": row["table_count"],
                            "file_path": row["file_path"],
                            "uploaded_at": str(row["uploaded_at"])
                        }
                except Exception as table_err:
                    logger.debug(f"document_contents unavailable, falling back to generated_minutes: {table_err}")
                    try:
                        conn.rollback()
                    except Exception:
                        pass

                # 2. Fallback: Query generated_minutes table and parse file from disk on-the-fly
                cursor.execute("""
                    SELECT id, company_name, meeting_type, meeting_date, file_path, created_at
                    FROM generated_minutes WHERE id = %s
                """, (doc_id,))
                gm_row = cursor.fetchone()
                if not gm_row:
                    return None

                filename = gm_row["file_path"]
                found_path = _resolve_minutes_file_path(filename)

                # Also search repository tree for uploaded files
                if not found_path:
                    repo_base = os.path.join(os.path.dirname(__file__), "..", "public", "repository")
                    if os.path.exists(repo_base) and filename:
                        for root, _, files in os.walk(repo_base):
                            if filename in files:
                                found_path = os.path.join(root, filename)
                                break

                extracted = {"text": "Document file not found on server disk.", "paragraph_count": 0, "tables": [], "table_count": 0}
                file_ext = os.path.splitext(filename or "")[1].lstrip('.')

                if found_path:
                    try:
                        with open(found_path, "rb") as f:
                            file_bytes = f.read()
                        if (filename or "").lower().endswith('.docx') and DOCX_AVAILABLE:
                            extracted = extract_text_from_docx(file_bytes)
                        elif (filename or "").lower().endswith('.pdf'):
                            extracted = extract_text_from_pdf(file_bytes)
                    except Exception as p_err:
                        logger.warning(f"On-the-fly extraction error for {filename}: {p_err}")

                return {
                    "id": gm_row["id"],
                    "filename": filename,
                    "company_name": gm_row["company_name"],
                    "vertical_name": "General",
                    "meeting_type": gm_row["meeting_type"],
                    "meeting_year": str(gm_row["meeting_date"])[:4] if gm_row["meeting_date"] else "2026",
                    "file_type": file_ext,
                    "extracted_text": extracted["text"],
                    "tables": extracted["tables"],
                    "paragraph_count": extracted["paragraph_count"],
                    "table_count": extracted["table_count"],
                    "file_path": filename,
                    "uploaded_at": str(gm_row["created_at"]) if gm_row.get("created_at") else "N/A",
                    "view_url": f"/api/generated-minutes/view/{filename}" if filename else None,
                }
            finally:
                conn.close()

        loop = asyncio.get_running_loop()
        result = await loop.run_in_executor(thread_pool, fetch)
        if result is None:
            raise HTTPException(status_code=404, detail="Document not found")
        return result
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error fetching document content {doc_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))



@router.get("/compliances/kpis")
async def get_compliance_kpis():
    """Compute secretarial compliance KPI analytics from PostgreSQL."""
    try:
        def compute():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn:
                return {
                    "total_compliances": 6,
                    "completed": 3,
                    "pending": 3,
                    "critical": 1,
                    "completion_rate": 50.0
                }
            cursor = get_pg_cursor(conn)
            try:
                cursor.execute("SELECT COUNT(*) as count FROM compliances")
                total = cursor.fetchone()['count'] or 0

                cursor.execute("SELECT COUNT(*) as count FROM compliances WHERE LOWER(status) = 'completed'")
                completed = cursor.fetchone()['count'] or 0

                cursor.execute("SELECT COUNT(*) as count FROM compliances WHERE LOWER(status) = 'pending'")
                pending = cursor.fetchone()['count'] or 0

                cursor.execute("SELECT COUNT(*) as count FROM compliances WHERE LOWER(priority) = 'critical'")
                critical = cursor.fetchone()['count'] or 0

                rate = round((completed / total * 100), 1) if total > 0 else 0.0

                # MOM #17: compliance evidence derived from uploaded/OCR-extracted documents
                documents_analysis = {"documents_processed": 0, "meetings_detected": 0, "meeting_dates": [], "companies_covered": []}
                try:
                    cursor.execute("""
                        SELECT id, filename, company_name, meeting_type, meeting_year, extracted_text
                        FROM document_contents ORDER BY id DESC LIMIT 200
                    """)
                    docs = cursor.fetchall()
                    date_pattern = re.compile(
                        r'\b(?:\d{1,2}(?:st|nd|rd|th)?\s+(?:January|February|March|April|May|June|July|August|September|October|November|December)[,\s]+\d{4}'
                        r'|(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2}[,\s]+\d{4}'
                        r'|\d{1,2}[/-]\d{1,2}[/-]\d{2,4})\b', re.IGNORECASE)
                    all_dates = set()
                    companies = set()
                    meetings = 0
                    for d in docs:
                        text = d['extracted_text'] or ''
                        found = date_pattern.findall(text[:20000])
                        if found:
                            meetings += 1
                            all_dates.update(found[:5])
                        if d['company_name']:
                            companies.add(d['company_name'])
                    documents_analysis = {
                        "documents_processed": len(docs),
                        "meetings_detected": meetings,
                        "meeting_dates": sorted(all_dates)[:25],
                        "companies_covered": sorted(companies)
                    }
                except Exception as doc_err:
                    logger.warning(f"Document-based compliance analysis skipped: {doc_err}")

                return {
                    "total_compliances": total,
                    "completed": completed,
                    "pending": pending,
                    "critical": critical,
                    "completion_rate": rate,
                    "documents_analysis": documents_analysis
                }
            finally:
                conn.close()

        loop = asyncio.get_running_loop()
        kpi_data = await loop.run_in_executor(thread_pool, compute)
        return kpi_data
    except Exception as e:
        logger.error(f"Error computing compliance KPIs: {e}")
        return {
            "total_compliances": 6,
            "completed": 3,
            "pending": 3,
            "critical": 1,
            "completion_rate": 50.0
        }


@router.get("/reports/attendance")
async def get_attendance_report(company_name: Optional[str] = None):
    """BRD Requirement #4: Return person-wise and meeting-wise attendance tracking data (real data)."""
    try:
        def fetch_report():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn:
                return {"person_wise": [], "meeting_wise": []}
            cursor = get_pg_cursor(conn)
            try:
                params = []
                company_filter = ""
                if company_name:
                    company_filter = " WHERE UPPER(company_name) = UPPER(%s)"
                    params.append(company_name)

                # Meeting-wise: real attendee counts per generated meeting
                cursor.execute(f"""
                    SELECT gm.id, gm.company_name, gm.meeting_type, gm.meeting_date, gm.file_path,
                           COALESCE(SUM(CASE WHEN ma.status = 'Present' THEN 1 ELSE 0 END), 0) AS present_count,
                           COUNT(ma.id) AS total_directors
                    FROM generated_minutes gm
                    LEFT JOIN meeting_attendance ma ON ma.minutes_id = gm.id
                    {company_filter.replace('company_name', 'gm.company_name') if company_filter else ''}
                    GROUP BY gm.id, gm.company_name, gm.meeting_type, gm.meeting_date, gm.file_path
                    ORDER BY gm.id DESC LIMIT 100
                """, tuple(params))
                rows = cursor.fetchall()

                meeting_wise = [{
                    "company_name": r["company_name"],
                    "meeting_type": r["meeting_type"],
                    "meeting_date": str(r["meeting_date"]),
                    "total_attendees": r["present_count"],
                    "total_directors": r["total_directors"],
                    "file_path": r["file_path"]
                } for r in rows]

                # Person-wise: per-director attended vs invited across recorded meetings
                cursor.execute(f"""
                    SELECT director_name, MAX(din) AS din,
                           SUM(CASE WHEN status = 'Present' THEN 1 ELSE 0 END) AS attended,
                           COUNT(*) AS invited
                    FROM meeting_attendance
                    {company_filter}
                    GROUP BY director_name
                    ORDER BY attended DESC, director_name
                """, tuple(params))
                p_rows = cursor.fetchall()

                person_wise = [{
                    "director_name": p["director_name"],
                    "din": p["din"],
                    "meetings_attended": p["attended"],
                    "meetings_invited": p["invited"],
                    "attendance_rate": f"{round((p['attended'] / p['invited']) * 100)}%" if p["invited"] else "0%"
                } for p in p_rows]

                return {
                    "meeting_wise": meeting_wise,
                    "person_wise": person_wise
                }
            finally:
                conn.close()

        loop = asyncio.get_running_loop()
        report_data = await loop.run_in_executor(thread_pool, fetch_report)
        return report_data
    except Exception as e:
        logger.error(f"Error building attendance report: {e}")
        return {"person_wise": [], "meeting_wise": []}


# --- MOM Email Delivery ---

class EmailDeliveryRequest(BaseModel):
    to_emails: List[str]
    subject: str = "Meeting Minutes"
    body: str = "Please find the attached Meeting Minutes document."
    filename: str  # File name from public/templates/ directory


@router.post("/email/send-mom")
async def send_mom_email(request: EmailDeliveryRequest):
    """BRD Requirement: Send generated MOM as email attachment via SMTP."""
    import smtplib
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    from email.mime.base import MIMEBase
    from email import encoders

    # Locate the file (generated outputs first, templates for legacy files)
    if os.path.basename(request.filename) != request.filename:
        raise HTTPException(status_code=400, detail="Invalid filename")
    base = os.path.join(os.path.dirname(__file__), "..", "public")
    file_path = None
    for sub in ("generated", "templates"):
        candidate = os.path.join(base, sub, request.filename)
        if os.path.exists(candidate):
            file_path = candidate
            break
    if not file_path:
        raise HTTPException(status_code=404, detail=f"File not found: {request.filename}")

    # Read SMTP config from environment
    smtp_host = os.getenv("SMTP_HOST", "smtp.office365.com")
    smtp_port = int(os.getenv("SMTP_PORT", "587"))
    smtp_user = os.getenv("SMTP_USER", "")
    smtp_pass = os.getenv("SMTP_PASSWORD", "")
    smtp_from = os.getenv("SMTP_FROM", smtp_user)

    if not smtp_user or not smtp_pass:
        raise HTTPException(
            status_code=503,
            detail="SMTP credentials not configured. Set SMTP_HOST, SMTP_PORT, SMTP_USER, SMTP_PASSWORD, SMTP_FROM in .env"
        )

    try:
        def send_email():
            msg = MIMEMultipart()
            msg['From'] = smtp_from
            msg['To'] = ", ".join(request.to_emails)
            msg['Subject'] = request.subject

            msg.attach(MIMEText(request.body, 'html'))

            # Attach the file
            with open(file_path, "rb") as f:
                part = MIMEBase('application', 'octet-stream')
                part.set_payload(f.read())
                encoders.encode_base64(part)
                part.add_header('Content-Disposition', f'attachment; filename="{request.filename}"')
                msg.attach(part)

            with smtplib.SMTP(smtp_host, smtp_port) as server:
                server.starttls()
                server.login(smtp_user, smtp_pass)
                server.send_message(msg)

            return True

        loop = asyncio.get_running_loop()
        await loop.run_in_executor(thread_pool, send_email)

        logger.info(f"MOM email sent to {request.to_emails} with attachment {request.filename}")
        return {"message": "Email sent successfully", "recipients": request.to_emails}
    except smtplib.SMTPException as smtp_err:
        logger.error(f"SMTP error: {smtp_err}")
        raise HTTPException(status_code=502, detail=f"Failed to send email: {str(smtp_err)}")
    except Exception as e:
        logger.error(f"Email delivery error: {e}")
        raise HTTPException(status_code=500, detail=f"Email delivery failed: {str(e)}")


# --- Place Master CRUD (Update & Delete) ---

class PlaceUpdateRequest(BaseModel):
    name: str
    address: str
    is_default: bool = False


@router.put("/places/{place_id}")
async def update_place(place_id: int, request: PlaceUpdateRequest, user: dict = Depends(require_session)):
    """Update an existing meeting place."""
    try:
        def update():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn:
                raise HTTPException(status_code=503, detail="Database unavailable")
            cursor = get_pg_cursor(conn)
            try:
                cursor.execute("""
                    UPDATE places SET name = %s, address = %s, is_default = %s WHERE id = %s
                    RETURNING id, name, address, is_default, created_at
                """, (request.name, request.address, request.is_default, place_id))
                row = cursor.fetchone()
                conn.commit()
                if not row:
                    return None
                return {
                    "id": row["id"], "name": row["name"],
                    "address": row["address"], "is_default": row["is_default"],
                    "created_at": str(row["created_at"])
                }
            finally:
                conn.close()

        loop = asyncio.get_running_loop()
        result = await loop.run_in_executor(thread_pool, update)
        if result is None:
            raise HTTPException(status_code=404, detail="Place not found")
        return result
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error updating place {place_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/places/{place_id}")
async def delete_place(place_id: int, user: dict = Depends(require_session)):
    """Delete a meeting place by ID."""
    try:
        def delete():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn:
                raise HTTPException(status_code=503, detail="Database unavailable")
            cursor = get_pg_cursor(conn)
            try:
                cursor.execute("DELETE FROM places WHERE id = %s RETURNING id", (place_id,))
                row = cursor.fetchone()
                conn.commit()
                return row is not None
            finally:
                conn.close()

        loop = asyncio.get_running_loop()
        deleted = await loop.run_in_executor(thread_pool, delete)
        if not deleted:
            raise HTTPException(status_code=404, detail="Place not found")
        return {"message": "Place deleted successfully", "id": place_id}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting place {place_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# --- Attendance Report CSV Export ---

@router.get("/reports/attendance/export")
async def export_attendance_csv():
    """Export attendance report as downloadable CSV."""
    from fastapi.responses import StreamingResponse

    try:
        def build_csv():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn:
                return "No Data"
            cursor = get_pg_cursor(conn)
            try:
                lines = ["Section,Company Name,Meeting Type,Meeting Date,Director Name,DIN,Status,File Name"]

                # Detailed per-director attendance rows
                cursor.execute("""
                    SELECT ma.company_name, ma.meeting_type, ma.meeting_date,
                           ma.director_name, ma.din, ma.status, gm.file_path
                    FROM meeting_attendance ma
                    LEFT JOIN generated_minutes gm ON gm.id = ma.minutes_id
                    ORDER BY ma.id DESC LIMIT 1000
                """)
                for r in cursor.fetchall():
                    lines.append(f'"Attendance","{r["company_name"]}","{r["meeting_type"]}","{r["meeting_date"]}","{r["director_name"]}","{r["din"] or ""}","{r["status"]}","{r["file_path"] or ""}"')

                # Meeting list (also covers minutes generated before attendance tracking)
                cursor.execute("""
                    SELECT company_name, meeting_type, meeting_date, file_path 
                    FROM generated_minutes ORDER BY id DESC LIMIT 200
                """)
                for r in cursor.fetchall():
                    lines.append(f'"Meeting","{r["company_name"]}","{r["meeting_type"]}","{r["meeting_date"]}","","","","{r["file_path"]}"')
                return "\n".join(lines)
            finally:
                conn.close()

        loop = asyncio.get_running_loop()
        csv_content = await loop.run_in_executor(thread_pool, build_csv)

        return StreamingResponse(
            iter([csv_content]),
            media_type="text/csv",
            headers={"Content-Disposition": "attachment; filename=attendance_report.csv"}
        )
    except Exception as e:
        logger.error(f"Error exporting attendance CSV: {e}")
        raise HTTPException(status_code=500, detail=str(e))



# --- Audit Log Endpoints ---

class AuditLogResponse(BaseModel):
    id: int
    entity_type: str
    entity_id: Optional[int]
    entity_name: str
    action: str
    performed_by: str
    performed_at: str
    old_data: Optional[dict]
    new_data: Optional[dict]
    remarks: Optional[str]
    vertical_id: Optional[int]
    company_name: Optional[str]
    ip_address: Optional[str]
    user_agent: Optional[str]


class AuditLogsListResponse(BaseModel):
    data: List[AuditLogResponse]
    count: int
    total_count: int


@router.get("/audit-logs", response_model=AuditLogsListResponse)
async def get_audit_logs(
    entity_type: Optional[str] = None,
    entity_id: Optional[int] = None,
    action: Optional[str] = None,
    company_name: Optional[str] = None,
    performed_by: Optional[str] = None,
    date_from: Optional[str] = None,
    date_to: Optional[str] = None,
    limit: int = 100,
    offset: int = 0,
    user: dict = Depends(require_session)
):
    """
    Get audit logs with comprehensive filtering
    
    Returns audit trail of all system operations. Admin users can see all logs,
    regular users can only see their own actions (unless they have special permissions).
    
    Args:
        entity_type: Filter by entity type ('company', 'meeting', 'user', etc.)
        entity_id: Filter by specific entity ID
        action: Filter by action ('created', 'updated', 'deleted', 'finalized', etc.)
        company_name: Filter by company name
        performed_by: Filter by user email who performed the action
        date_from: Filter from date (YYYY-MM-DD format)
        date_to: Filter to date (YYYY-MM-DD format)
        limit: Maximum number of results (default: 100, max: 1000)
        offset: Pagination offset
        user: Authenticated user from session
        
    Returns:
        List of audit log entries with pagination info
    """
    # Limit max results
    if limit > 1000:
        limit = 1000
    
    try:
        def fetch():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn:
                raise RuntimeError("Database connection unavailable")
            
            cursor = get_pg_cursor(conn)
            try:
                # Build dynamic query with filters
                query = """
                    SELECT 
                        id, entity_type, entity_id, entity_name, action, 
                        performed_by, performed_at, old_data, new_data, 
                        remarks, vertical_id, company_name, ip_address, user_agent
                    FROM audit_logs
                    WHERE 1=1
                """
                params = []
                
                # Check if user has admin privileges
                # For now, assume non-admin users can only see their own logs
                # You can enhance this with proper RBAC later
                user_email = user.get('email', '')
                user_role = user.get('role', 'user')
                
                # Apply user-level filtering (non-admins see only their own actions)
                if user_role not in ['master_admin', 'admin']:
                    query += " AND performed_by = %s"
                    params.append(user_email)
                
                # Apply entity type filter
                if entity_type:
                    query += " AND entity_type = %s"
                    params.append(entity_type)
                
                # Apply entity ID filter
                if entity_id:
                    query += " AND entity_id = %s"
                    params.append(entity_id)
                
                # Apply action filter
                if action:
                    query += " AND action = %s"
                    params.append(action)
                
                # Apply company name filter
                if company_name:
                    query += " AND company_name = %s"
                    params.append(company_name)
                
                # Apply performed_by filter (admin only)
                if performed_by and user_role in ['master_admin', 'admin']:
                    query += " AND performed_by = %s"
                    params.append(performed_by)
                
                # Apply date range filters
                if date_from:
                    query += " AND DATE(performed_at) >= %s"
                    params.append(date_from)
                
                if date_to:
                    query += " AND DATE(performed_at) <= %s"
                    params.append(date_to)
                
                # Get total count before applying pagination
                count_query = query.replace(
                    "SELECT id, entity_type, entity_id, entity_name, action, performed_by, performed_at, old_data, new_data, remarks, vertical_id, company_name, ip_address, user_agent",
                    "SELECT COUNT(*) as total"
                )
                cursor.execute(count_query, tuple(params))
                total_count = cursor.fetchone()['total'] or 0
                
                # Apply ordering and pagination
                query += " ORDER BY performed_at DESC, id DESC LIMIT %s OFFSET %s"
                params.extend([limit, offset])
                
                cursor.execute(query, tuple(params))
                rows = cursor.fetchall()
                
                data = []
                for r in rows:
                    # Parse JSON fields
                    old_data = r['old_data']
                    new_data = r['new_data']
                    
                    # If they're strings, parse them
                    if isinstance(old_data, str):
                        try:
                            old_data = json.loads(old_data) if old_data else None
                        except:
                            old_data = None
                    
                    if isinstance(new_data, str):
                        try:
                            new_data = json.loads(new_data) if new_data else None
                        except:
                            new_data = None
                    
                    data.append(AuditLogResponse(
                        id=r['id'],
                        entity_type=r['entity_type'],
                        entity_id=r['entity_id'],
                        entity_name=r['entity_name'],
                        action=r['action'],
                        performed_by=r['performed_by'],
                        performed_at=str(r['performed_at']),
                        old_data=old_data,
                        new_data=new_data,
                        remarks=r['remarks'],
                        vertical_id=r['vertical_id'],
                        company_name=r['company_name'],
                        ip_address=r.get('ip_address'),
                        user_agent=r.get('user_agent')
                    ))
                
                return data, len(data), total_count
                
            finally:
                conn.close()
        
        data, count, total = await asyncio.get_running_loop().run_in_executor(thread_pool, fetch)
        return AuditLogsListResponse(data=data, count=count, total_count=total)
        
    except Exception as e:
        logger.error(f"Error fetching audit logs: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch audit logs: {str(e)}")


@router.get("/audit-logs/summary")
async def get_audit_logs_summary(
    user: dict = Depends(require_session)
):
    """
    Get audit log summary statistics
    
    Returns overview of audit activity including:
    - Total actions logged
    - Actions by type
    - Recent activity
    - Top users by activity
    
    Args:
        user: Authenticated user from session
        
    Returns:
        Summary statistics dictionary
    """
    try:
        def fetch():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn:
                raise RuntimeError("Database connection unavailable")
            
            cursor = get_pg_cursor(conn)
            try:
                user_role = user.get('role', 'user')
                user_email = user.get('email', '')
                
                # Base filter for non-admin users
                user_filter = ""
                if user_role not in ['master_admin', 'admin']:
                    user_filter = f" WHERE performed_by = '{user_email}'"
                
                # Total count
                cursor.execute(f"SELECT COUNT(*) as total FROM audit_logs{user_filter}")
                total_actions = cursor.fetchone()['total'] or 0
                
                # Actions by type
                cursor.execute(f"""
                    SELECT action, COUNT(*) as count 
                    FROM audit_logs{user_filter}
                    GROUP BY action 
                    ORDER BY count DESC 
                    LIMIT 10
                """)
                actions_by_type = [{"action": r['action'], "count": r['count']} for r in cursor.fetchall()]
                
                # Actions by entity type
                cursor.execute(f"""
                    SELECT entity_type, COUNT(*) as count 
                    FROM audit_logs{user_filter}
                    GROUP BY entity_type 
                    ORDER BY count DESC 
                    LIMIT 10
                """)
                actions_by_entity = [{"entity_type": r['entity_type'], "count": r['count']} for r in cursor.fetchall()]
                
                # Recent activity (last 24 hours)
                cursor.execute(f"""
                    SELECT COUNT(*) as count 
                    FROM audit_logs 
                    WHERE performed_at > NOW() - INTERVAL '24 hours'{' AND ' + user_filter.replace('WHERE ', '') if user_filter else ''}
                """)
                recent_24h = cursor.fetchone()['count'] or 0
                
                # Top users (admin only)
                top_users = []
                if user_role in ['master_admin', 'admin']:
                    cursor.execute("""
                        SELECT performed_by, COUNT(*) as count 
                        FROM audit_logs 
                        GROUP BY performed_by 
                        ORDER BY count DESC 
                        LIMIT 10
                    """)
                    top_users = [{"user": r['performed_by'], "actions": r['count']} for r in cursor.fetchall()]
                
                return {
                    "total_actions": total_actions,
                    "recent_24h": recent_24h,
                    "actions_by_type": actions_by_type,
                    "actions_by_entity": actions_by_entity,
                    "top_users": top_users
                }
                
            finally:
                conn.close()
        
        summary = await asyncio.get_running_loop().run_in_executor(thread_pool, fetch)
        return summary
        
    except Exception as e:
        logger.error(f"Error fetching audit log summary: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch summary: {str(e)}")



@router.get("/companies/{company_id}/audit-history", response_model=AuditLogsListResponse)
async def get_company_audit_history(
    company_id: int,
    limit: int = 50,
    offset: int = 0,
    user: dict = Depends(require_session)
):
    """
    Get complete audit history for a specific company
    
    Returns all audit log entries related to a company including:
    - Company creation, updates, and deletion
    - Related meeting actions (if company_name matches)
    - Any other operations involving this company
    
    Args:
        company_id: ID of the company
        limit: Maximum number of results (default: 50)
        offset: Pagination offset
        user: Authenticated user from session
        
    Returns:
        Chronological list of all audit entries for the company
    """
    try:
        def fetch():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn:
                raise RuntimeError("Database connection unavailable")
            
            cursor = get_pg_cursor(conn)
            try:
                # First, get the company name
                cursor.execute("SELECT name FROM companies WHERE id = %s", (company_id,))
                company_row = cursor.fetchone()
                
                if not company_row:
                    # Company might have been deleted, try to find it in audit logs
                    cursor.execute("""
                        SELECT entity_name 
                        FROM audit_logs 
                        WHERE entity_type = 'company' AND entity_id = %s 
                        ORDER BY performed_at DESC 
                        LIMIT 1
                    """, (company_id,))
                    deleted_row = cursor.fetchone()
                    if deleted_row:
                        company_name = deleted_row['entity_name']
                    else:
                        raise HTTPException(status_code=404, detail=f"Company with ID {company_id} not found")
                else:
                    company_name = company_row['name']
                
                # Get all audit logs related to this company
                # This includes:
                # 1. Direct company actions (entity_type='company' AND entity_id=company_id)
                # 2. Related actions (company_name field matches)
                query = """
                    SELECT 
                        id, entity_type, entity_id, entity_name, action, 
                        performed_by, performed_at, old_data, new_data, 
                        remarks, vertical_id, company_name, ip_address, user_agent
                    FROM audit_logs
                    WHERE (entity_type = 'company' AND entity_id = %s)
                       OR (company_name = %s)
                    ORDER BY performed_at DESC, id DESC
                    LIMIT %s OFFSET %s
                """
                
                cursor.execute(query, (company_id, company_name, limit, offset))
                rows = cursor.fetchall()
                
                # Get total count
                count_query = """
                    SELECT COUNT(*) as total
                    FROM audit_logs
                    WHERE (entity_type = 'company' AND entity_id = %s)
                       OR (company_name = %s)
                """
                cursor.execute(count_query, (company_id, company_name))
                total_count = cursor.fetchone()['total'] or 0
                
                data = []
                for r in rows:
                    # Parse JSON fields
                    old_data = r['old_data']
                    new_data = r['new_data']
                    
                    if isinstance(old_data, str):
                        try:
                            old_data = json.loads(old_data) if old_data else None
                        except:
                            old_data = None
                    
                    if isinstance(new_data, str):
                        try:
                            new_data = json.loads(new_data) if new_data else None
                        except:
                            new_data = None
                    
                    data.append(AuditLogResponse(
                        id=r['id'],
                        entity_type=r['entity_type'],
                        entity_id=r['entity_id'],
                        entity_name=r['entity_name'],
                        action=r['action'],
                        performed_by=r['performed_by'],
                        performed_at=str(r['performed_at']),
                        old_data=old_data,
                        new_data=new_data,
                        remarks=r['remarks'],
                        vertical_id=r['vertical_id'],
                        company_name=r['company_name'],
                        ip_address=r.get('ip_address'),
                        user_agent=r.get('user_agent')
                    ))
                
                return data, len(data), total_count
                
            finally:
                conn.close()
        
        data, count, total = await asyncio.get_running_loop().run_in_executor(thread_pool, fetch)
        return AuditLogsListResponse(data=data, count=count, total_count=total)
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error fetching company audit history for ID {company_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch audit history: {str(e)}")


@router.get("/companies/{company_id}/audit-timeline")
async def get_company_audit_timeline(
    company_id: int,
    user: dict = Depends(require_session)
):
    """
    Get a visual timeline of company history
    
    Returns a simplified timeline view of key events in the company's lifecycle:
    - Creation date and creator
    - All updates with what changed
    - Meeting milestones
    - Deletion (if applicable)
    
    Args:
        company_id: ID of the company
        user: Authenticated user from session
        
    Returns:
        Timeline data structure suitable for UI visualization
    """
    try:
        def fetch():
            conn = get_pg_connection(os.getenv('POSTGRES_DATABASE_MINUTES'))
            if not conn:
                raise RuntimeError("Database connection unavailable")
            
            cursor = get_pg_cursor(conn)
            try:
                # Get company name
                cursor.execute("SELECT name FROM companies WHERE id = %s", (company_id,))
                company_row = cursor.fetchone()
                
                if not company_row:
                    cursor.execute("""
                        SELECT entity_name 
                        FROM audit_logs 
                        WHERE entity_type = 'company' AND entity_id = %s 
                        ORDER BY performed_at DESC 
                        LIMIT 1
                    """, (company_id,))
                    deleted_row = cursor.fetchone()
                    if deleted_row:
                        company_name = deleted_row['entity_name']
                    else:
                        raise HTTPException(status_code=404, detail=f"Company with ID {company_id} not found")
                else:
                    company_name = company_row['name']
                
                # Get all relevant audit entries
                cursor.execute("""
                    SELECT 
                        id, entity_type, action, performed_by, performed_at, 
                        old_data, new_data, remarks
                    FROM audit_logs
                    WHERE (entity_type = 'company' AND entity_id = %s)
                       OR (company_name = %s)
                    ORDER BY performed_at ASC
                """, (company_id, company_name))
                
                rows = cursor.fetchall()
                
                timeline = []
                for r in rows:
                    event = {
                        "id": r['id'],
                        "timestamp": str(r['performed_at']),
                        "action": r['action'],
                        "entity_type": r['entity_type'],
                        "performed_by": r['performed_by'],
                        "description": r['remarks'] or f"{r['action']} {r['entity_type']}"
                    }
                    
                    # Add change details for updates
                    if r['action'] == 'updated' and r['old_data'] and r['new_data']:
                        try:
                            old = json.loads(r['old_data']) if isinstance(r['old_data'], str) else r['old_data']
                            new = json.loads(r['new_data']) if isinstance(r['new_data'], str) else r['new_data']
                            
                            changes = []
                            for key in new.keys():
                                if key in old and old[key] != new[key]:
                                    changes.append({
                                        "field": key,
                                        "from": str(old[key]),
                                        "to": str(new[key])
                                    })
                            event['changes'] = changes
                        except:
                            pass
                    
                    timeline.append(event)
                
                return {
                    "company_id": company_id,
                    "company_name": company_name,
                    "total_events": len(timeline),
                    "timeline": timeline
                }
                
            finally:
                conn.close()
        
        timeline_data = await asyncio.get_running_loop().run_in_executor(thread_pool, fetch)
        return timeline_data
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error fetching company timeline for ID {company_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch timeline: {str(e)}")
